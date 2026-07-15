import pymupdf as fitz
from docx import Document
import os
import re
import shutil
from pathlib import Path
import time
import logging
from datetime import datetime
import gc
import json
from itertools import groupby

try:
    from pdf2docx import Converter
except ImportError:
    Converter = None

# Настройки pdf2docx (передаются в convert()). Дефолты библиотеки разумны, но часть
# «потерь текста / кривой вёрстки» лечится их подкруткой ПОД ТИП ДОКУМЕНТОВ — здесь
# их удобно держать в одном месте. Самые влияющие:
#   parse_stream_table=False — НЕ искать «таблицы без рамок». Частая причина кривой
#       вёрстки текстовых документов: pdf2docx ошибочно сворачивает абзацы в таблицу.
#       Для юр-документов (сплошной текст, мало настоящих таблиц) выключение обычно
#       даёт более верную вёрстку. Если в ваших PDF есть таблицы без видимых рамок —
#       верните True.
#   ignore_page_error=True — не падать из-за ошибки на одной странице (но её текст
#       тогда теряется; это ловит контроль покрытия _check_text_coverage ниже).
PDF2DOCX_SETTINGS = {
    "parse_stream_table": False,
    "ignore_page_error":  True,
}

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('pdf_conversion.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

class PDFPipelineConverter:
    def __init__(self, input_folder, output_folder, backup_folder=None):
        """
        input_folder: папка с PDF файлами (будут удаляться после конвертации)
        output_folder: папка для сохранения DOCX файлов
        backup_folder: опциональная папка для бэкапа PDF перед удалением
        """
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.backup_folder = Path(backup_folder) if backup_folder else None
        
        # Статистика
        self.stats = {
            'processed': 0,
            'success': 0,
            'failed': 0,
            'skipped': 0,
            'total_size_saved_mb': 0,
            'start_time': None,
            'end_time': None
        }
        
        # Создаем папки
        self.output_folder.mkdir(parents=True, exist_ok=True)
        if self.backup_folder:
            self.backup_folder.mkdir(parents=True, exist_ok=True)
        
        # Файл для сохранения состояния
        self.checkpoint_file = self.input_folder / 'conversion_state.json'
        
    @staticmethod
    def _norm(s: str) -> str:
        """Нормализация для сверки: убираем ВСЕ пробелы и регистр — устойчиво к
        переносам/переформатированию строк, которое делает pdf2docx."""
        return re.sub(r"\s+", "", s).lower()

    _WORD_RE = re.compile(r"[A-Za-zА-Яа-яЁё0-9]{3,}")

    @classmethod
    def _words(cls, s: str) -> list[str]:
        """Значащие слова (3+ символа) в нижнем регистре — для пословной сверки."""
        return cls._WORD_RE.findall(s.lower())

    def _recover_missing_text(self, pdf_path, docx_path) -> None:
        """ГАРАНТИЯ ПОЛНОТЫ ТЕКСТА. pdf2docx иногда теряет строки. Здесь сверяем
        встроенный текст PDF (для нативных PDF он достоверен) с текстом DOCX и
        ДОПИСЫВАЕМ в конец DOCX строки, которых в нём не оказалось. Так ни одно слово
        не теряется — ценой того, что восстановленные строки идут отдельным блоком в
        конце (без исходной вёрстки), но текст присутствует полностью."""
        try:
            d = Document(str(docx_path))
        except Exception as exc:
            logging.warning("Полнота: не открыть DOCX %s: %s", docx_path.name, exc)
            return
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        docx_words = set(self._words(" ".join(parts)))

        try:
            pdf = fitz.open(pdf_path)
        except Exception as exc:
            logging.warning("Полнота: не открыть PDF %s: %s", pdf_path.name, exc)
            return
        missing: list[tuple[int, str]] = []   # (номер страницы, строка)
        seen: set[str] = set()
        src_chars = 0
        for i in range(pdf.page_count):
            page_text = pdf[i].get_text()
            src_chars += len(page_text)
            for line in page_text.splitlines():
                s = line.strip()
                lw = self._words(s)
                if len(lw) < 2:             # слишком коротко, чтобы судить (номер/одно слово)
                    continue
                # Строка ПОТЕРЯНА, только если БОЛЬШИНСТВО её слов отсутствуют в DOCX.
                # Иначе это та же строка, просто переразбитая pdf2docx (не дублируем).
                absent = sum(1 for w in lw if w not in docx_words)
                if absent >= max(2, int(0.6 * len(lw) + 0.999)):
                    key = self._norm(s)
                    if key not in seen:     # не дописываем одну и ту же строку дважды
                        seen.add(key)
                        missing.append((i + 1, s))
        pdf.close()

        if src_chars < 50:                  # PDF без текстового слоя (скан) — сверять нечего
            return
        if not missing:
            logging.info("Полнота текста %s: OK — потерь не найдено", pdf_path.name)
            return

        try:
            d.add_page_break()
            d.add_paragraph("[Восстановленный текст — pdf2docx не перенёс эти строки в вёрстку]")
            cur_page = None
            for page_no, s in missing:
                if page_no != cur_page:     # заголовок-ориентир: откуда взята строка
                    cur_page = page_no
                    d.add_paragraph(f"— страница {page_no} —")
                d.add_paragraph(s)
            d.save(str(docx_path))
            logging.warning("%s: pdf2docx потерял %d строк — дописаны в конец DOCX "
                            "(полнота восстановлена)", pdf_path.name, len(missing))
        except Exception as exc:
            logging.error("Полнота: не дописать восстановленный текст в %s: %s", docx_path.name, exc)

    # Длина окна n-грамм: сколько подряд идущих слов PDF склеиваем в один ключ.
    # 6 покрывает реальные склейки pdf2docx (обычно 2-4 слова) с запасом.
    _GLUE_MAX_N = 6

    @classmethod
    def _build_glue_map(cls, pdf) -> dict[str, list[int]]:
        """КАРТА РЕАЛЬНЫХ СКЛЕЕК. Берём слова PDF (PyMuPDF расставляет пробелы по
        координатам корректно), группируем по строкам (block, line) и генерируем
        конкатенации подряд идущих слов. Ключ — склеенная форма в нижнем регистре,
        значение — позиции границ (куда вставлять пробелы).

        Расклеиваем ТОЛЬКО то, что реально стояло рядом в исходнике, — никакой
        словарной сегментации, поэтому корректные слова не шинкуются."""
        glue: dict[str, list[int]] = {}
        singles: set[str] = set()           # самостоятельные слова PDF — их НЕ трогаем
        # Строки по блокам (в порядке следования) — для склеек на стыке соседних строк.
        by_block: dict[tuple, list[list[str]]] = {}
        for page in pdf:
            # (x0, y0, x1, y1, "word", block_no, line_no, word_no)
            words = page.get_text("words")
            words.sort(key=lambda w: (w[5], w[6], w[7]))
            for (blk, _ln), line in groupby(words, key=lambda w: (w[5], w[6])):
                toks = [w[4] for w in line]
                by_block.setdefault((page.number, blk), []).append(toks)
                for i in range(len(toks)):
                    acc, bounds = "", []
                    for n in range(1, cls._GLUE_MAX_N + 1):
                        if i + n > len(toks):
                            break
                        if n > 1:
                            bounds.append(len(acc))   # граница перед очередным словом
                        acc += toks[i + n - 1]
                        if n == 1:
                            singles.add(acc.lower())
                        else:
                            # setdefault: при коллизии оставляем самую короткую (раннюю)
                            # разбивку — она однозначнее
                            glue.setdefault(acc.lower(), list(bounds))
        # СКЛЕЙКА НА СТЫКЕ СОСЕДНИХ СТРОК ОДНОГО БЛОКА: слово/словосочетание перенесено
        # на след. строку, а pdf2docx убрал и перенос, и пробел («Центрального»⏎
        # «федерального» → «Центральногофедерального»). Обычный перенос → вставляем
        # пробел на стыке (граница = длина левого слова). Дефис-перенос («об-»/«ласти»)
        # здесь НЕ трогаем — это снятие дефиса, иная операция (отдельный шаг).
        for lines in by_block.values():
            for k in range(len(lines) - 1):
                left, right = lines[k], lines[k + 1]
                if not left or not right:
                    continue
                a, b2 = left[-1], right[0]
                if not a or not b2 or a.endswith("-"):
                    continue
                glue.setdefault((a + b2).lower(), [len(a)])
        # Защита от ложных разрезов: если склеенная форма совпадает с реальным
        # самостоятельным словом (напр. «дело» = «дел» + «о»), не расклеиваем его.
        for key in singles:
            glue.pop(key, None)
        return glue

    # Обрамляющая пунктуация, которую pdf2docx может приклеить к токену иначе,
    # чем PyMuPDF разбил слова (кавычки, скобки, знаки препинания по краям).
    _EDGE_PUNCT = "«»\"'()[]{}.,;:!?—–-…"

    @classmethod
    def _unglue(cls, token: str, glue: dict[str, list[int]]) -> str:
        """Вставляет пробелы в склеенный токен по известным из PDF границам.
        Символы DOCX не заменяются — сохраняется точный регистр/написание.
        При промахе повторяет поиск без обрамляющей пунктуации: pdf2docx и
        PyMuPDF по-разному приклеивают кавычки/скобки к словам."""
        bounds = glue.get(token.lower())
        head = tail = ""
        if not bounds:
            core = token.strip(cls._EDGE_PUNCT)
            if not core or core == token:
                return token
            i = token.find(core)
            head, tail = token[:i], token[i + len(core):]
            bounds = glue.get(core.lower())
            if not bounds:
                return token
            token = core
        parts, prev = [], 0
        for b in bounds:
            parts.append(token[prev:b])
            prev = b
        parts.append(token[prev:])
        return head + " ".join(p for p in parts if p) + tail

    def _fix_paragraph(self, paragraph, glue: dict[str, list[int]]) -> int:
        """Чинит склейки в одном абзаце на двух уровнях:
        1) внутри run (pdf2docx склеил слова в одном фрагменте текста);
        2) на стыке соседних runs (слова в разных фрагментах без пробела между ними
           — частый случай: «1.» + «Включить» → визуально «1.Включить»).
        Возвращает число изменённых мест."""
        changed = 0
        runs = paragraph.runs

        # (1) внутрирунные склейки
        for run in runs:
            toks = run.text.split(" ")
            new = [self._unglue(t, glue) if t else t for t in toks]
            if new != toks:
                run.text = " ".join(new)
                changed += 1

        # (2) склейки на границе соседних runs — вставляем пробел в стык, только
        # если конкатенация хвоста и головы совпадает с известной границей из PDF
        for i in range(len(runs) - 1):
            left, right = runs[i].text, runs[i + 1].text
            if not left or not right or left[-1].isspace() or right[0].isspace():
                continue
            a = left.rsplit(" ", 1)[-1]     # последний токен левого run
            b = right.split(" ", 1)[0]      # первый токен правого run
            bounds = glue.get((a + b).lower())
            if bounds and len(a) in bounds:
                runs[i + 1].text = " " + right
                changed += 1

        return changed

    def _fix_glued_words(self, pdf_path, docx_path) -> None:
        """ИСПРАВЛЕНИЕ СКЛЕЕННЫХ СЛОВ. pdf2docx иногда теряет пробелы между словами
        (особенность реконструкции текста по координатам глифов). Строим карту
        реальных склеек из PDF и точечно вставляем недостающие пробелы в DOCX."""
        try:
            pdf = fitz.open(pdf_path)
        except Exception as exc:
            logging.warning("Расклейка: не открыть PDF %s: %s", pdf_path.name, exc)
            return
        try:
            glue = self._build_glue_map(pdf)
        finally:
            pdf.close()
        if not glue:                        # нет текстового слоя (скан) — расклеивать нечего
            return

        try:
            d = Document(str(docx_path))
        except Exception as exc:
            logging.warning("Расклейка: не открыть DOCX %s: %s", docx_path.name, exc)
            return

        changed = 0

        for paragraph in d.paragraphs:
            changed += self._fix_paragraph(paragraph, glue)
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        changed += self._fix_paragraph(paragraph, glue)

        if changed:
            try:
                d.save(str(docx_path))
                logging.info("Расклейка %s: исправлено runs с пробелами — %d",
                             pdf_path.name, changed)
            except Exception as exc:
                logging.error("Расклейка: не сохранить %s: %s", docx_path.name, exc)

    # Закрывающая кавычка », у которой ПЕРЕД ней пробел, а ПОСЛЕ — буква: это на деле
    # ОТКРЫВАЮЩАЯ кавычка, неверно повёрнутая pdf2docx при конвертации прямых кавычек
    # («… "ДЕЛО"» распозналось как «… »ДЕЛО»»). Диапазоны кириллицы — unicode-escape'ами,
    # чтобы не зависеть от кодировки исходника. Настоящая закрывающая » идёт вплотную
    # к слову (без пробела перед), поэтому такой случай сюда не попадает.
    _MISPLACED_OPEN_QUOTE_RE = re.compile(
        "(?<=\\s)»(?=[А-Яа-яЁёA-Za-z])"
    )

    def _fix_quote_direction(self, docx_path) -> None:
        """Чинит неверно повёрнутую кавычку »→« (см. _MISPLACED_OPEN_QUOTE_RE)."""
        try:
            d = Document(str(docx_path))
        except Exception as exc:
            logging.warning("Кавычки: не открыть DOCX %s: %s", docx_path.name, exc)
            return

        def fix_paragraph(paragraph) -> int:
            n = 0
            for run in paragraph.runs:
                if "»" in run.text:
                    new = self._MISPLACED_OPEN_QUOTE_RE.sub("«", run.text)
                    if new != run.text:
                        run.text = new
                        n += 1
            return n

        changed = 0
        for paragraph in d.paragraphs:
            changed += fix_paragraph(paragraph)
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        changed += fix_paragraph(paragraph)
        if changed:
            try:
                d.save(str(docx_path))
                logging.info("Кавычки %s: повёрнуто »→« — %d", docx_path.name, changed)
            except Exception as exc:
                logging.error("Кавычки: не сохранить %s: %s", docx_path.name, exc)

    def convert_single_pdf(self, pdf_path, docx_path):
        """Конвертация одного PDF в редактируемый DOCX"""
        converter = None
        try:
            if Converter is None:
                raise ImportError(
                    "Не установлена библиотека pdf2docx. Установите ее командой: python -m pip install pdf2docx"
                )

            converter = Converter(str(pdf_path))
            # Настройки pdf2docx — в PDF2DOCX_SETTINGS (см. начало файла, можно крутить).
            # end не указываем — дефолт pdf2docx (None = до последней страницы).
            converter.convert(str(docx_path), start=0, **PDF2DOCX_SETTINGS)

            # Расклейка: возвращаем пробелы, потерянные pdf2docx между словами.
            # ВАЖНО до _recover_missing_text — расклеенный текст лучше сверяется
            # пословно, поэтому ложных «потерь» меньше.
            self._fix_glued_words(pdf_path, docx_path)

            # Кавычки: pdf2docx мог неверно повернуть открывающую кавычку (»ДЕЛО → «ДЕЛО).
            self._fix_quote_direction(docx_path)

            # Гарантия полноты: дописываем текст, который pdf2docx потерял
            self._recover_missing_text(pdf_path, docx_path)

            # Получаем размер файла для статистики
            file_size = os.path.getsize(pdf_path) / (1024 * 1024)

            return True, file_size
            
        except MemoryError:
            logging.error(f"Memory error while converting {pdf_path.name}")
            return False, 0
        except Exception as e:
            logging.error(f"Conversion error for {pdf_path.name}: {str(e)}")
            return False, 0
        finally:
            if converter:
                converter.close()
            gc.collect()
    
    def safe_delete_pdf(self, pdf_path, move_to_backup=True):
        """Безопасное удаление PDF файла (с возможностью перемещения в бэкап)"""
        try:
            if move_to_backup and self.backup_folder:
                # Перемещаем в папку бэкапа вместо удаления
                backup_path = self.backup_folder / pdf_path.name
                shutil.move(str(pdf_path), str(backup_path))
                logging.debug(f"Moved to backup: {pdf_path.name}")
            else:
                # Удаляем файл
                os.remove(pdf_path)
                logging.debug(f"Deleted: {pdf_path.name}")
            return True
        except Exception as e:
            logging.error(f"Failed to delete/move {pdf_path.name}: {e}")
            return False
    
    def save_checkpoint(self):
        """Сохранение состояния обработки"""
        checkpoint_data = {
            'processed': self.stats['processed'],
            'success': self.stats['success'],
            'failed': self.stats['failed'],
            'total_size_saved_mb': self.stats['total_size_saved_mb'],
            'last_update': datetime.now().isoformat()
        }
        try:
            with open(self.checkpoint_file, 'w') as f:
                json.dump(checkpoint_data, f, indent=2)
        except Exception as e:
            logging.warning(f"Failed to save checkpoint: {e}")
    
    def load_checkpoint(self):
        """Загрузка состояния обработки"""
        if self.checkpoint_file.exists():
            try:
                with open(self.checkpoint_file, 'r') as f:
                    data = json.load(f)
                    self.stats.update(data)
                logging.info(f"Loaded checkpoint: {self.stats['success']} files already processed")
                return True
            except Exception as e:
                logging.warning(f"Failed to load checkpoint: {e}")
        return False
    
    def process_pipeline(self, delete_after_processing=True, move_to_backup=False):
        """
        Основной конвейер обработки
        delete_after_processing: удалять PDF после конвертации
        move_to_backup: перемещать PDF в папку бэкапа перед удалением
        """
        self.stats['start_time'] = time.time()
        
        # Загружаем чекпоинт
        self.load_checkpoint()
        
        # Получаем список PDF файлов
        pdf_files = sorted(self.input_folder.glob("*.pdf"))
        
        # Фильтруем уже обработанные (если есть чекпоинт)
        total_found = len(pdf_files)
        logging.info(f"Found {total_found} PDF files in {self.input_folder}")
        
        # Мониторинг дискового пространства
        def check_disk_space(folder, required_gb: float = 1):
            free_space = shutil.disk_usage(folder).free / (1024**3)
            if free_space < required_gb:
                logging.warning(f"Low disk space on {folder}: only {free_space:.1f} GB free")
                return False
            return True
        
        # Основной цикл обработки
        for i, pdf_path in enumerate(pdf_files, 1):
            # Проверяем, не обработан ли уже файл (по имени)
            docx_path = self.output_folder / f"{pdf_path.stem}.docx"
            if docx_path.exists():
                logging.info(f"[{i}/{total_found}] Skipping (already converted): {pdf_path.name}")
                self.stats['skipped'] += 1
                if delete_after_processing:
                    self.safe_delete_pdf(pdf_path, move_to_backup=move_to_backup)
                continue
            
            logging.info(f"[{i}/{total_found}] Processing: {pdf_path.name} ({pdf_path.stat().st_size / (1024*1024):.2f} MB)")
            
            # Конвертируем PDF в DOCX
            success, file_size = self.convert_single_pdf(pdf_path, docx_path)
            
            if success:
                self.stats['success'] += 1
                self.stats['total_size_saved_mb'] += file_size
                
                # Удаляем или перемещаем исходный PDF
                if delete_after_processing:
                    if self.safe_delete_pdf(pdf_path, move_to_backup=move_to_backup):
                        logging.info(f"✓ Converted and removed: {pdf_path.name}")
                    else:
                        logging.warning(f"✓ Converted but failed to remove: {pdf_path.name}")
                else:
                    logging.info(f"✓ Converted successfully: {pdf_path.name}")
            else:
                self.stats['failed'] += 1
                logging.error(f"✗ Failed to convert: {pdf_path.name}")
            
            self.stats['processed'] = i
            
            # Каждые 10 файлов сохраняем чекпоинт и выводим статистику
            if i % 10 == 0:
                self.save_checkpoint()
                self.print_stats(i, total_found)
                check_disk_space(self.output_folder, 0.5)
            
            # Очищаем память каждые 50 файлов
            if i % 50 == 0:
                gc.collect()
                logging.info("Memory cleanup performed")
        
        self.stats['end_time'] = time.time()
        self.print_final_stats()
        
        # Удаляем чекпоинт после успешного завершения
        if self.stats['failed'] == 0:
            try:
                self.checkpoint_file.unlink()
                logging.info("Checkpoint file removed (all files processed successfully)")
            except:
                pass
    
    def print_stats(self, current, total):
        """Вывод промежуточной статистики"""
        elapsed = time.time() - self.stats['start_time']
        rate = current / elapsed if elapsed > 0 else 0
        remaining_files = total - current
        eta = remaining_files / rate if rate > 0 else 0
        
        logging.info(f"Progress: {current}/{total} | "
                    f"{self.stats['success']} | "
                    f"{self.stats['failed']} | "
                    f"{self.stats['skipped']} | "
                    f"Saved: {self.stats['total_size_saved_mb']:.1f} MB | "
                    f"{rate:.2f} files/sec | "
                    f"ETA: {eta/60:.1f} min")
    
    def print_final_stats(self):
        """Вывод финальной статистики"""
        total_time = self.stats['end_time'] - self.stats['start_time']
        logging.info("CONVERSION PIPELINE COMPLETED")
        logging.info(f"Total time: {total_time/60:.1f} minutes ({total_time/3600:.2f} hours)")
        logging.info(f"Successfully converted: {self.stats['success']}")
        logging.info(f"Failed: {self.stats['failed']}")
        logging.info(f"Skipped: {self.stats['skipped']}")
        logging.info(f"Total disk space saved: {self.stats['total_size_saved_mb']:.1f} MB")
        logging.info(f"Average speed: {self.stats['success']/(total_time/3600):.1f} files/hour")
        
        if self.stats['failed'] > 0:
            logging.warning(f"{self.stats['failed']} files failed. Check log for details.")
        else:
            logging.info("All files processed successfully")
        logging.info("="*60)

# Запуск конвейера
def run_conversion_pipeline():
    """
    Пример настройки и запуска конвейера
    """
    # Настройка путей
    INPUT_FOLDER = "./pdf_to_convert"      # Папка с PDF (файлы будут удаляться)
    OUTPUT_FOLDER = "./converted_word"     # Папка для DOCX файлов
    
    # Создаем экземпляр конвертера
    converter = PDFPipelineConverter(
        input_folder=INPUT_FOLDER,
        output_folder=OUTPUT_FOLDER,
        backup_folder=None
    )
    
    # Запускаем обработку
    # delete_after_processing=True - удалять PDF после конвертации
    # move_to_backup=False - удалять PDF сразу, без папки бэкапа
    converter.process_pipeline(
        delete_after_processing=True,  # Удалять оригиналы
        move_to_backup=False            # Удалять сразу, без бэкапа
    )

# Альтернативный простой вариант (минимальная настройка)
def simple_pipeline():
    """Простой вариант - только конвертация и удаление без бэкапов"""
    
    INPUT_FOLDER = "./pdf_files"
    OUTPUT_FOLDER = "./word_files"
    
    converter = PDFPipelineConverter(
        input_folder=INPUT_FOLDER,
        output_folder=OUTPUT_FOLDER,
        backup_folder=None  # Без бэкапа - сразу удаляем
    )
    
    converter.process_pipeline(
        delete_after_processing=True,
        move_to_backup=False  # Просто удаляем
    )

if __name__ == "__main__":
    # 2 варианта запуска
    
    # Вариант 1: Полный с бэкапом (рекомендуется для важных файлов)
    run_conversion_pipeline()
    
    # Вариант 2: Простой без бэкапа (максимальная экономия места)
    # simple_pipeline()