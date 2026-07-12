"""
docx_builder.py
───────────────
Все функции для сборки Word-документа: стили, таблицы, шапка side-by-side,
блоки МЕТКА:содержимое, анализ страниц.
"""
from __future__ import annotations

import statistics
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .config import (
    BODY_PT, EMU_PER_INCH, FONT_NAME,
    LABEL_COL_RATIO, LABEL_LINE_RE, SIDE_LABEL_RE,
    MARGIN_INCH, PAGE_H_INCH, PAGE_W_INCH,
)
from .geometry import bbox_h, bbox_mid_y, bbox_x0, coplanar
from .ocr_fixes import postprocess

_BODY_LABELS = frozenset({"paragraph", "text", "list_item"})


# ── Инициализация документа ───────────────────────────────────────────────────

def init_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = int(PAGE_W_INCH * EMU_PER_INCH)
    sec.page_height = int(PAGE_H_INCH * EMU_PER_INCH)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, int(MARGIN_INCH * EMU_PER_INCH))

    normal           = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.widow_control     = False
    normal.paragraph_format.keep_together     = False
    normal.paragraph_format.keep_with_next    = False

    for level in (1, 2, 3):
        h                = doc.styles[f"Heading {level}"]
        h.font.name      = FONT_NAME
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.font.bold      = True
        h.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.SINGLE
        h.paragraph_format.keep_with_next     = False
        h.paragraph_format.keep_together      = False
        # Сбрасываем отступы заголовков чтобы RIGHT/CENTER-выравнивание
        # не сдвигалось из-за дефолтного left_indent Word (Heading 2 ≈ 0.25in)
        h.paragraph_format.left_indent        = Pt(0)
        h.paragraph_format.first_line_indent  = Pt(0)

    # Нумерованный список: явный висячий отступ совпадает с "-" подпунктами.
    # number at 17.7pt, text at 35.4pt — одинаково для "1." и "–" items.
    try:
        ln = doc.styles["List Number"]
        ln.font.name = FONT_NAME
        ln.font.size = Pt(BODY_PT)
        ln.paragraph_format.left_indent         = Pt(35.4)
        ln.paragraph_format.first_line_indent   = Pt(-17.7)
        ln.paragraph_format.line_spacing_rule   = WD_LINE_SPACING.SINGLE
        ln.paragraph_format.widow_control       = False
    except KeyError:
        pass

    return doc


# ── Выравнивание ──────────────────────────────────────────────────────────────

def detect_alignment(bbox, page_width: float) -> WD_ALIGN_PARAGRAPH:
    if page_width <= 0:
        return WD_ALIGN_PARAGRAPH.LEFT
    x0      = float(getattr(bbox, "l", getattr(bbox, "x0", 0)))
    x1      = float(getattr(bbox, "r", getattr(bbox, "x1", page_width)))
    block_w = x1 - x0
    cx      = (x0 + x1) / 2
    ratio   = block_w / page_width
    if ratio > 0.75:
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if abs(cx - page_width / 2) < page_width * 0.08 and ratio < 0.68:
        return WD_ALIGN_PARAGRAPH.CENTER
    # RIGHT только для КОРОТКИХ блоков, реально прижатых к правому краю
    # (дата, подпись, номер). Длинный текст в правой колонке (адреса, реквизиты) —
    # это ЛЕВОЕ чтение в колонке: делаем LEFT, позицию задаёт left_indent.
    if x0 > page_width * 0.12 and x1 > page_width * 0.72 and ratio < 0.45:
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


# ── Вспомогательные функции для таблиц ───────────────────────────────────────

def _set_cell_borders(cell) -> None:
    tc      = cell._tc
    tcPr    = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)


def _make_borderless_table(doc: Document, n_cols: int,
                           total_width_dxa: int | None = None) -> object:
    """Создаёт таблицу без видимых границ с фиксированной раскладкой ячеек."""
    tbl       = doc.add_table(rows=1, cols=n_cols)
    tbl.style = "Normal Table"
    tbl_el    = tbl._tbl
    tbl_pr    = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    # Фиксированная раскладка — Word не пересчитывает ширины
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    # Явная ширина таблицы
    if total_width_dxa is not None:
        tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), str(total_width_dxa))
        tbl_w.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_w)
    tbl_brd = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        tbl_brd.append(el)
    tbl_pr.append(tbl_brd)
    return tbl


def _set_cell_width(cell, width_inch: float) -> None:
    tc  = cell._tc
    tcp = tc.find(qn("w:tcPr"))
    if tcp is None:
        tcp = OxmlElement("w:tcPr")
        tc.insert(0, tcp)
    # Удаляем все существующие w:tcW чтобы наш был единственным
    for _old in tcp.findall(qn("w:tcW")):
        tcp.remove(_old)
    tcw = OxmlElement("w:tcW")
    tcw.set(qn("w:w"), str(int(width_inch * 1440)))
    tcw.set(qn("w:type"), "dxa")
    tcp.insert(0, tcw)
    tcb = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "nil")
        tcb.append(el)
    tcp.append(tcb)


# ── Таблицы документа ─────────────────────────────────────────────────────────

def add_table_from_grid(doc: Document, grid: list) -> None:
    if not grid:
        return
    n_rows = len(grid)
    n_cols = max((len(row) for row in grid), default=0)
    if n_cols == 0:
        return
    tbl       = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    for r_idx, row in enumerate(grid):
        for c_idx in range(n_cols):
            cell      = tbl.rows[r_idx].cells[c_idx]
            cell_data = row[c_idx] if c_idx < len(row) else None
            text      = postprocess((getattr(cell_data, "text", "") or "") if cell_data else "")
            _set_cell_borders(cell)
            for p in cell.paragraphs:
                p.clear()
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)
            run           = para.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = Pt(9.0)
            run.bold      = (r_idx == 0)
    doc.add_paragraph()


def _table_col_widths_inch(data, n_cols: int, text_w_inch: float) -> list[float] | None:
    """Ширины колонок из bbox ячеек TableFormer, отмасштабированные на ширину
    текстовой зоны. None если у ячеек нет пригодных координат."""
    starts: dict[int, list[float]] = {}
    right = 0.0
    for cell in getattr(data, "table_cells", []) or []:
        bb = getattr(cell, "bbox", None)
        if bb is None:
            continue
        c0 = int(getattr(cell, "start_col_offset_idx", 0))
        l  = float(getattr(bb, "l", 0.0))
        r  = float(getattr(bb, "r", 0.0))
        starts.setdefault(c0, []).append(l)
        right = max(right, r)
    if len(starts) < n_cols or right <= 0:
        return None
    xs = []
    for c in range(n_cols):
        vals = sorted(starts.get(c, []))
        if not vals:
            return None
        xs.append(vals[len(vals) // 2])
    if any(b <= a for a, b in zip(xs, xs[1:])) or right <= xs[-1]:
        return None
    bounds = xs + [right]
    widths_pt = [bounds[i + 1] - bounds[i] for i in range(n_cols)]
    total = sum(widths_pt)
    if total <= 0:
        return None
    return [w / total * text_w_inch for w in widths_pt]


def add_table_from_cells(doc: Document, table_item,
                         text_w_inch: float | None = None) -> None:
    """Таблица из TableFormer-ячеек: с объединёнными ячейками (span'ы) и
    ширинами колонок из координат. Раньше span'ы игнорировались (писался
    только start-offset) — сетка «съезжала», Word давил колонки поровну."""
    data   = table_item.data
    n_rows = getattr(data, "num_rows", 0)
    n_cols = getattr(data, "num_cols", 0)
    if n_rows == 0 or n_cols == 0:
        return
    tbl       = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"

    # Ширины колонок: фиксированная раскладка, чтобы Word не пересчитывал
    widths = (_table_col_widths_inch(data, n_cols, text_w_inch)
              if text_w_inch else None)
    if widths:
        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl_el.insert(0, tbl_pr)
        _layout = OxmlElement("w:tblLayout")
        _layout.set(qn("w:type"), "fixed")
        tbl_pr.append(_layout)
        tbl_grid = tbl_el.find(qn("w:tblGrid"))
        if tbl_grid is not None:
            for gc in tbl_grid.findall(qn("w:gridCol")):
                tbl_grid.remove(gc)
            for w in widths:
                gc = OxmlElement("w:gridCol")
                gc.set(qn("w:w"), str(int(w * 1440)))
                tbl_grid.append(gc)

    filled: set[tuple[int, int]] = set()
    for cell in getattr(data, "table_cells", []) or []:
        r0 = max(int(getattr(cell, "start_row_offset_idx", 0)), 0)
        c0 = max(int(getattr(cell, "start_col_offset_idx", 0)), 0)
        r1 = min(int(getattr(cell, "end_row_offset_idx", r0 + 1)), n_rows)
        c1 = min(int(getattr(cell, "end_col_offset_idx", c0 + 1)), n_cols)
        if r0 >= n_rows or c0 >= n_cols or r1 <= r0 or c1 <= c0:
            continue
        target = tbl.cell(r0, c0)
        if (r1 - r0 > 1 or c1 - c0 > 1) and (r0, c0) not in filled:
            try:
                target = target.merge(tbl.cell(r1 - 1, c1 - 1))
            except Exception:
                pass                       # уже слитые/пересекающиеся span'ы
        for rr in range(r0, r1):
            for cc in range(c0, c1):
                filled.add((rr, cc))
        text = postprocess(getattr(cell, "text", "") or "")
        if not text:
            continue
        for p in target.paragraphs:
            p.clear()
        para = target.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run           = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(9.0)
        run.bold      = bool(getattr(cell, "column_header", False))

    for row in tbl.rows:
        for cell in row.cells:
            _set_cell_borders(cell)
    doc.add_paragraph()


# ── Шапка: логотип + текст рядом ─────────────────────────────────────────────

def add_sidebyside(
    doc: Document,
    pil_img,
    img_w_inch: float,
    right_blocks: list[dict],
    text_zone_inch: float,
) -> None:
    """Безрамочная 2-колоночная таблица: логотип слева, текст реквизитов справа."""
    img_col_w = min(img_w_inch + 0.15, text_zone_inch * 0.48)
    txt_col_w = max(text_zone_inch - img_col_w, 1.0)

    tbl    = _make_borderless_table(doc, 2)
    cell_l = tbl.cell(0, 0)
    cell_r = tbl.cell(0, 1)
    _set_cell_width(cell_l, img_col_w)
    _set_cell_width(cell_r, txt_col_w)

    p_img = cell_l.paragraphs[0]
    p_img.alignment                     = WD_ALIGN_PARAGRAPH.LEFT
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after  = Pt(0)
    buf = BytesIO()
    pil_img.save(buf, format="PNG")
    buf.seek(0)
    p_img.add_run().add_picture(buf, width=Inches(min(img_w_inch, img_col_w - 0.05)))

    first = True
    for blk in right_blocks:
        para = cell_r.paragraphs[0] if first else cell_r.add_paragraph()
        first = False
        para.alignment                     = blk.get("alignment", WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(1)
        run           = para.add_run(blk["text"])
        run.font.name = FONT_NAME
        run.font.size = Pt(blk.get("font_pt", BODY_PT))
        run.bold      = blk.get("bold", False)
        run.italic    = blk.get("italic", False)

    doc.add_paragraph()


# ── Блоки МЕТКА:содержимое ───────────────────────────────────────────────────

def add_header_row(
    doc: Document,
    cells: list[dict],
    text_zone_inch: float,
) -> None:
    """Render a first-page letterhead row with image/text cells in visual order."""
    if not cells:
        return

    # Суммарная ширина всех ячеек в dxa для фиксированной раскладки
    col_dxa_list = [
        int(max(float(c.get("width_pt", 72.0)), 1.0) / 72.0 * 1440)
        for c in cells
    ]
    total_dxa = sum(col_dxa_list)
    tbl = _make_borderless_table(doc, len(cells), total_width_dxa=total_dxa)

    # Перезаписываем tblGrid правильными ширинами колонок — иначе Word игнорирует tcW
    tbl_el = tbl._tbl
    tbl_grid = tbl_el.find(qn("w:tblGrid"))
    if tbl_grid is not None:
        for gc in tbl_grid.findall(qn("w:gridCol")):
            tbl_grid.remove(gc)
        for dxa in col_dxa_list:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(dxa))
            tbl_grid.append(gc)

    for idx, cell_data in enumerate(cells):
        cell = tbl.cell(0, idx)
        # Используем абсолютную ширину в pt чтобы ячейки встали точно
        col_w_inch = max(float(cell_data.get("width_pt", 72.0)), 1.0) / 72.0
        _set_cell_width(cell, col_w_inch)

        # Убираем внутренние поля ячейки чтобы текст занимал максимальную ширину
        _tc = cell._tc
        _tcPr = _tc.get_or_add_tcPr()
        _mar = OxmlElement("w:tcMar")
        for _side in ("top", "left", "bottom", "right"):
            _el = OxmlElement(f"w:{_side}")
            _el.set(qn("w:w"), "0")
            _el.set(qn("w:type"), "dxa")
            _mar.append(_el)
        _tcPr.append(_mar)

        para = cell.paragraphs[0]
        para.alignment = cell_data.get("alignment", WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        if cell_data.get("kind") == "image":
            pil_img = cell_data.get("image")
            if pil_img is None:
                continue
            buf = BytesIO()
            pil_img.save(buf, format="PNG")
            buf.seek(0)
            img_w = min(float(cell_data.get("image_width_inch", col_w_inch)), col_w_inch - 0.05)
            if img_w > 0:
                para.add_run().add_picture(buf, width=Inches(img_w))
            continue

        # Tab-stop для label\tcontent блоков: 5cm от левого края ячейки
        _tab_stop_dxa = int(5.0 * 567)  # 5cm в dxa

        first = True
        for block in cell_data.get("blocks", []):
            p = para if first else cell.add_paragraph()
            first = False
            p.alignment = cell_data.get("alignment", WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_before = Pt(block.get("space_before", 0.0))
            p.paragraph_format.space_after = Pt(0)
            text = block.get("text", "")
            font_pt = block.get("font_pt", 8.5)
            blk_bold = block.get("bold", False)
            blk_italic = block.get("italic", False)
            # Поддержка \t (label\tcontent) и \n (перенос строки)
            # \t разбивает на [label_bold, tab, content_normal]
            if "\t" in text:
                # Устанавливаем tab-stop на 5cm от края ячейки
                pPr = p._p.get_or_add_pPr()
                tabs_el = OxmlElement("w:tabs")
                tab_el = OxmlElement("w:tab")
                tab_el.set(qn("w:val"), "left")
                tab_el.set(qn("w:pos"), str(_tab_stop_dxa))
                tabs_el.append(tab_el)
                pPr.append(tabs_el)

                label_part, _, content_part = text.partition("\t")
                # Метка — жирный
                r_label = p.add_run(label_part)
                r_label.font.name = FONT_NAME
                r_label.font.size = Pt(font_pt)
                r_label.bold = True
                r_label.italic = blk_italic
                # TAB run
                r_tab = p.add_run("\t")
                r_tab.font.name = FONT_NAME
                r_tab.font.size = Pt(font_pt)
                # Содержимое после \t — может содержать \n
                sub_parts = content_part.split("\n")
                for k, sp in enumerate(sub_parts):
                    if k > 0:
                        p.add_run().add_break()
                    r = p.add_run(sp)
                    r.font.name = FONT_NAME
                    r.font.size = Pt(font_pt)
                    r.bold = False
                    r.italic = blk_italic
            else:
                parts = text.split("\n")
                for k, part in enumerate(parts):
                    if k > 0:
                        p.add_run().add_break()
                    run = p.add_run(part)
                    run.font.name = FONT_NAME
                    run.font.size = Pt(font_pt)
                    run.bold = blk_bold
                    run.italic = blk_italic

    doc.add_paragraph()


def add_signature_row(
    doc: Document,
    left_text: str,
    right_text: str,
    text_w_inch: float,
    space_before: float = 6.0,
    sig_image=None,
) -> None:
    """Строка подписи: текст (bold) | картинка подписи | инициалы (bold)."""
    n_cols = 3 if (sig_image is not None or right_text) else 2
    tbl = _make_borderless_table(doc, n_cols)

    if n_cols == 3:
        col_l = text_w_inch * 0.38
        col_m = text_w_inch * 0.38
        col_r = max(text_w_inch - col_l - col_m, 0.5)
        for i, w in enumerate((col_l, col_m, col_r)):
            _set_cell_width(tbl.cell(0, i), w)
    else:
        col_l = text_w_inch * 0.70
        col_r = max(text_w_inch - col_l, 0.5)
        for i, w in enumerate((col_l, col_r)):
            _set_cell_width(tbl.cell(0, i), w)

    def _vcenter(cell) -> None:
        """Вертикальное выравнивание по центру ячейки."""
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        va   = OxmlElement("w:vAlign")
        va.set(qn("w:val"), "center")
        tcPr.append(va)

    # Левая колонка — жирный текст должности
    _vcenter(tbl.cell(0, 0))
    p = tbl.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(left_text)
    run.font.name = FONT_NAME
    run.font.size = Pt(BODY_PT)
    run.bold      = True

    if n_cols == 3:
        # Средняя колонка — картинка подписи (если есть)
        _vcenter(tbl.cell(0, 1))
        p = tbl.cell(0, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(0)
        if sig_image is not None:
            buf = BytesIO()
            sig_image.save(buf, format="PNG")
            buf.seek(0)
            iw, ih = sig_image.size
            aspect = ih / iw if iw > 0 else 1.0
            target_w = min(col_m - 0.1, 1.5)
            if target_w * aspect > 1.2:
                target_w = 1.2 / aspect
            p.add_run().add_picture(buf, width=Inches(max(target_w, 0.3)))

        # Правая колонка — жирные инициалы, прижаты к правому краю
        _vcenter(tbl.cell(0, 2))
        p_right = tbl.cell(0, 2).paragraphs[0]
    else:
        _vcenter(tbl.cell(0, 1))
        p_right = tbl.cell(0, 1).paragraphs[0]
    p = p_right

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Л.В.Халиль к правому краю
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(0)
    if right_text:
        run = p.add_run(right_text)
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_PT)
        run.bold      = True


def add_es_stamp(doc: Document, lines: list[str], width_inch: float) -> None:
    """Штамп электронной подписи в виде рамки (одна ячейка с границами).

    Воспроизводит визуальный штамп ЭП из исходного PDF: каждая строка штампа
    («Электронная подпись действительна», «Данные ЭП: …», «Дата …», «Кому
    выдана: …») рендерится отдельным абзацем внутри одной обрамлённой ячейки.
    """
    lines = [ln for ln in (l.strip() for l in lines) if ln]
    if not lines:
        return
    box_w = max(min(width_inch * 0.55, width_inch), 2.0)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Normal Table"
    # Штамп ЭП в оригинале — справа внизу: выравниваем рамку по правому краю.
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # Фиксированная раскладка + явная ширина
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    _layout = OxmlElement("w:tblLayout")
    _layout.set(qn("w:type"), "fixed")
    tbl_pr.append(_layout)
    _tw = OxmlElement("w:tblW")
    _tw.set(qn("w:w"), str(int(box_w * 1440)))
    _tw.set(qn("w:type"), "dxa")
    tbl_pr.append(_tw)

    cell = tbl.cell(0, 0)
    # Ширина ячейки задаём напрямую (без nil-границ, которые ставит _set_cell_width,
    # иначе они конфликтуют с видимой рамкой ниже).
    _tc  = cell._tc
    _tcp = _tc.get_or_add_tcPr()
    for _old in _tcp.findall(qn("w:tcW")):
        _tcp.remove(_old)
    _tcw = OxmlElement("w:tcW")
    _tcw.set(qn("w:w"), str(int(box_w * 1440)))
    _tcw.set(qn("w:type"), "dxa")
    _tcp.insert(0, _tcw)
    _set_cell_borders(cell)        # видимая рамка штампа

    for p in cell.paragraphs:
        p.clear()
    for i, ln in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(ln)
        run.font.name = FONT_NAME
        run.font.size = Pt(9.0)
        run.italic    = True
    doc.add_paragraph()


def split_label_content(text: str) -> tuple[str, str] | None:
    """Возвращает (метка, содержимое), если текст начинается с метки блока:
    ALL-CAPS метки (ДОЛЖНИК:/КРЕДИТОР:) ИЛИ якорной метки стороны в любом регистре
    (Должник:/Кредитор:/Адрес:… — RapidOCR даёт их Titlecase, см. SIDE_LABEL_RE)."""
    t = text.strip()
    m = LABEL_LINE_RE.match(t)
    if m:
        label = m.group(1).strip()
        alpha = [c for c in label if c.isalpha()]
        # Минимум 4 буквы: исключаем ИНН, БИК, КПП — они реквизиты, не метки блоков
        if alpha and all(c.isupper() for c in alpha) and len(alpha) >= 4:
            return label, m.group(2).strip()
    # Якорные Titlecase-метки сторон (Должник:/Кредитор:/Финансовый управляющий:…)
    ms = SIDE_LABEL_RE.match(t)
    if ms:
        return ms.group(1).strip(), ms.group(2).strip()
    return None


def add_label_content_table(
    doc: Document,
    label: str,
    content_items: list[dict],
    text_w_inch: float,
    space_before: float = 0.0,
    indent_inch: float = 0.0,
    col_ratio: float | None = None,
) -> None:
    """Безрамочная 2-колоночная строка: метка слева, содержимое справа.

    indent_inch — сдвиг таблицы от левого поля (для правоколоночных блоков).
    col_ratio   — доля ширины для колонки метки (None → LABEL_COL_RATIO).
    """
    ratio = col_ratio if col_ratio is not None else LABEL_COL_RATIO
    col_l = text_w_inch * ratio
    col_r = max(text_w_inch - col_l, 1.5)

    tbl    = _make_borderless_table(doc, 2)
    cl, cr = tbl.cell(0, 0), tbl.cell(0, 1)
    _set_cell_width(cl, col_l)
    _set_cell_width(cr, col_r)

    # Убираем внутренние отступы ячеек (default=0.08in) — контент ближе к метке
    for _cell in (cl, cr):
        _tc   = _cell._tc
        _tcPr = _tc.get_or_add_tcPr()
        _mar  = OxmlElement("w:tcMar")
        for _side in ("top", "left", "bottom", "right"):
            _el = OxmlElement(f"w:{_side}")
            _el.set(qn("w:w"),    "0")
            _el.set(qn("w:type"), "dxa")
            _mar.append(_el)
        _tcPr.append(_mar)

    # Сдвиг таблицы от левого поля (tblInd, единица — twips = 1/1440 дюйма)
    if indent_inch > 0.01:
        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(qn("w:tblPr"))
        if tbl_pr is not None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_ind.set(qn("w:w"),    str(int(indent_inch * 1440)))
            tbl_ind.set(qn("w:type"), "dxa")
            tbl_pr.append(tbl_ind)

    p           = cl.paragraphs[0]
    p.alignment                     = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(0)
    r           = p.add_run(label)
    r.font.name = FONT_NAME
    r.font.size = Pt(BODY_PT)
    r.bold      = True

    first = True
    for blk in content_items:
        p = cr.paragraphs[0] if first else cr.add_paragraph()
        first = False
        p.alignment                     = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r           = p.add_run(blk["text"])
        r.font.name = FONT_NAME
        r.font.size = Pt(blk.get("font_pt", BODY_PT))
        r.bold      = blk.get("bold", False)
        r.italic    = blk.get("italic", False)


# ── Анализ страниц ────────────────────────────────────────────────────────────

def analyse_pages(
    items: list,
) -> tuple[dict[int, float], dict[int, float]]:
    """
    Однопроходный анализ всех элементов:
      page_medians  — медианная высота bbox body-text по каждой странице
      page_left_min — 10-й перцентиль левого края (эффективное поле)
    """
    heights_by_page: dict[int, list[float]] = {}
    lefts_by_page:   dict[int, list[float]] = {}

    for item, _ in items:
        raw = getattr(item, "label", None)
        if raw is None:
            continue
        label = (raw.value if hasattr(raw, "value") else str(raw)).lower()
        if label not in _BODY_LABELS:
            continue
        prov_list = getattr(item, "prov", None) or []
        if not prov_list:
            continue
        bbox = getattr(prov_list[0], "bbox", None)
        if bbox is None:
            continue
        h = bbox_h(bbox)
        l = float(getattr(bbox, "l", 0))
        if h < 2:
            continue
        page = int(getattr(prov_list[0], "page_no", 1))
        heights_by_page.setdefault(page, []).append(h)
        if l >= 0:
            lefts_by_page.setdefault(page, []).append(l)

    page_medians: dict[int, float] = {
        p: statistics.median(hs) for p, hs in heights_by_page.items() if hs
    }
    page_left_min: dict[int, float] = {}
    for p, ls in lefts_by_page.items():
        if ls:
            ls_s = sorted(ls)
            page_left_min[p] = ls_s[max(0, len(ls_s) // 10)]

    return page_medians, page_left_min
