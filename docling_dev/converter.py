"""
converter.py
Основная логика конвертации DoclingDocument → DOCX.
word_order.py используется только для ПЕРЕУПОРЯДОЧИВАНИЯ блоков Docling —
текст из Docling остаётся нетронутым (качество OCR Docling лучше сырого EasyOCR).
"""
from __future__ import annotations

import gc
import logging
import re
import shutil
import statistics
import time
from collections import defaultdict
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:
    from docx.document import Document as _Document
    from docx.text.paragraph import Paragraph

from .config import (
    BODY_PT, FONT_NAME, LABEL_HEADING, LABEL_MARGIN_THRESHOLD, LABEL_PT,
    LINE_SPACING, MARGIN_INCH, SIDE_LABEL_RE, SKIP_LABELS,
)
from .docx_builder import (
    add_label_content_table, add_sidebyside, add_es_stamp,
    add_header_row, add_signature_row,
    add_table_from_cells, add_table_from_grid,
    analyse_pages, detect_alignment, init_document, split_label_content,
)
from .geometry import (
    bbox_h, bbox_mid_y, bbox_x0, coplanar,
    detect_pdf_native, sort_reading_order,
)
from .highlight import is_junk_text
from .ocr_fixes import clean_body_text, postprocess
from .ocr_preprocess import is_junk_image
from .page_analyser import (
    analyse_pages as analyse_page_infos,
    build_indent_levels, snap_indent,
)
from .pipeline import build_converter
from .word_order import TextBlock, blocks_for_page

log = logging.getLogger(__name__)


#Переупорядочивание через word_order

def _reorder_by_word_order(
    all_items: list,
    word_blocks_map: dict[int, tuple[list, float]],
    page_sizes: dict[int, tuple[float, float]],
) -> list:
    """
    Переставляет ТОЛЬКО body-текст (paragraph/text) по позициям word_order-блоков.

    Картинки, таблицы, заголовки остаются на своих местах — это критично для
    корректной работы add_sidebyside (логотип + текст реквизитов рядом).

    Алгоритм:
      1. Собираем paragraph/text элементы каждой страницы, пропуская:
         - элементы в верхней зоне (топ 15% страницы) — шапка с логотипом
         - страницы без word_order данных
      2. Сортируем по word_order-блокам (block_idx, x0).
      3. Вставляем отсортированные элементы обратно на те же позиции.
         Нетекстовые элементы и элементы шапки не трогаем.

    Константа HEADER_ZONE: топ 15% страницы не переупорядочивается, чтобы
    не сломать логику add_sidebyside (логотип + текст реквизитов рядом).
    """
    _BODY        = frozenset({"paragraph", "text"})
    HEADER_ZONE  = 0.15   # топ 15% — шапка страницы (логотип, реквизиты)
    MIN_MOVE     = 2      # минимум элементов должны сдвинуться (снижено с 3-х до 2-х: ловит swap двух блоков)
    MAX_Y_DIFF   = 0.04   # макс. разница Y [0–1] для надёжного совпадения блока

    result: list = list(all_items)
    page_body: defaultdict[int, list] = defaultdict(list)

    for i, (item, _level) in enumerate(all_items):
        raw = getattr(item, "label", None)
        if raw is None:
            continue
        label = (raw.value if hasattr(raw, "value") else str(raw)).lower()
        if label not in _BODY:
            continue

        prov_list = getattr(item, "prov", None) or []
        if not prov_list:
            continue
        page_no = int(getattr(prov_list[0], "page_no", 9999))
        if page_no not in word_blocks_map:
            continue

        bbox = getattr(prov_list[0], "bbox", None)
        if bbox is None:
            continue

        blocks, img_h = word_blocks_map[page_no]
        _, ph = page_sizes.get(page_no, (595.0, 842.0))
        if not blocks or img_h <= 0 or ph <= 0:
            continue

        # Нормализуем Y к [0,1]: Docling (y=0 снизу) - (0=верх, 1=низ)
        y_norm = (ph - bbox_mid_y(bbox)) / ph

        # Пропускаем элементы из шапки страницы (логотип, реквизиты)
        if y_norm < HEADER_ZONE:
            continue

        best_bi  = min(range(len(blocks)),
                       key=lambda bi: abs(blocks[bi].mid_y / img_h - y_norm))
        y_diff = abs(blocks[best_bi].mid_y / img_h - y_norm)
        if y_diff > MAX_Y_DIFF:
            raw_text = (getattr(all_items[i][0], "text", "") or "")[:40]
            log.debug("word_order стр.%d: нет совпадения (y_diff=%.3f>%.3f) %r",
                      page_no, y_diff, MAX_Y_DIFF, raw_text)
            continue
        page_body[page_no].append((i, best_bi, bbox_x0(bbox)))

    for page_no, body in page_body.items():
        if len(body) < 2:
            log.debug("word_order стр.%d: только %d блоков — пропуск", page_no, len(body))
            continue

        slots     = sorted(x[0] for x in body)
        wo_sorted = sorted(body, key=lambda x: (x[1], x[2]))

        if [x[0] for x in body] == [x[0] for x in wo_sorted]:
            log.debug("word_order стр.%d: %d блоков — порядок уже верный", page_no, len(body))
            continue

        moved = sum(1 for s, (oi, _, _) in zip(slots, wo_sorted) if s != oi)
        log.info("word_order стр.%d: %d/%d блоков требуют перестановки (MIN_MOVE=%d)",
                 page_no, moved, len(body), MIN_MOVE)
        if moved < MIN_MOVE:
            for s, (oi, bi, x0) in zip(slots, wo_sorted):
                if s != oi:
                    txt = (getattr(all_items[oi][0], "text", "") or "")[:50]
                    log.info("  word_order пропущен (moved=%d<MIN=%d): slot=%d←orig=%d bi=%d %r",
                             moved, MIN_MOVE, s, oi, bi, txt)
            continue

        # Если все перемещаемые блоки указывают на один и тот же bi,
        # сортировка по x0 ненадёжна — пропускаем (оба блока на одной Y-позиции)
        moved_bis = {bi for s, (oi, bi, _) in zip(slots, wo_sorted) if s != oi}
        if len(moved_bis) == 1:
            sole_bi = next(iter(moved_bis))
            log.info("word_order стр.%d: пропуск — %d блоков конкурируют за bi=%d (x0-сортировка ненадёжна)",
                     page_no, moved, sole_bi)
            continue

        log.info("word_order стр.%d: применяем перестановку %d блоков:", page_no, moved)
        for s, (oi, bi, _) in zip(slots, wo_sorted):
            if s != oi:
                txt = (getattr(all_items[oi][0], "text", "") or "")[:50]
                log.info("  слот %d ← orig=%d (bi=%d) %r", s, oi, bi, txt)

        for slot, (orig_idx, _bi, _x0) in zip(slots, wo_sorted):
            result[slot] = all_items[orig_idx]

        log.info("word_order стр.%d: скорректировано %d блоков", page_no, moved)

    return result


# Вспомогательные функции

def _label_str(item) -> str:
    raw = getattr(item, "label", None)
    if raw is None:
        return ""
    return (raw.value if hasattr(raw, "value") else str(raw)).lower()


def _item_bbox_page(item, current_page: int):
    prov_list = getattr(item, "prov", None) or []
    if not prov_list:
        return None, current_page
    bbox   = getattr(prov_list[0], "bbox", None)
    page_no = int(getattr(prov_list[0], "page_no", current_page))
    return bbox, page_no


def _get_item_image(item, dl_doc):
    """Возвращает PIL Image для item.

    Docling кэширует page images при scale=images_scale (например 2.0), но
    page.image обращается к scale=1.0 (default). Если get_image() вернул None,
    пробуем напрямую через кэш страницы — берём первый доступный scale.
    """
    try:
        img = item.get_image(dl_doc)
        if img is not None:
            return img
    except Exception:
        pass

    # Fallback: crop из кэша страницы (любой доступный scale)
    try:
        prov_list = getattr(item, "prov", None) or []
        if not prov_list:
            return None
        prov = prov_list[0]
        bbox    = getattr(prov, "bbox", None)
        page_no = int(getattr(prov, "page_no", 0))
        page_obj = dl_doc.pages.get(page_no)
        if page_obj is None or bbox is None:
            return None
        cache = getattr(page_obj, "_image_cache", {})
        if not cache:
            return None
        page_size = getattr(page_obj, "size", None)
        if page_size is None:
            return None
        page_h = float(getattr(page_size, "height", 0) or 0)
        if page_h <= 0:
            return None
        # Берём изображение с наибольшим scale (лучшее качество)
        for _, img_ref in sorted(cache.items(), key=lambda kv: kv[0], reverse=True):
            pil = getattr(img_ref, "pil_image", None)
            if pil is None:
                continue
            img_size = getattr(img_ref, "size", None)
            if img_size is None:
                continue
            crop_bb = (
                bbox.to_top_left_origin(page_height=page_h)
                .scale_to_size(old_size=page_size, new_size=img_size)
            )
            return pil.crop(crop_bb.as_tuple())
    except Exception as exc:
        log.debug("_get_item_image fallback failed: %s", exc)
    return None


# Исправление порядка чтения

"""
# Номер пункта: «1 », «1. », «1) » (OCR часто теряет точку). 1–2 цифры —
# списки не длиннее 99 пунктов; защищает от срабатывания на суммах («1 356 000»
# в этом контексте не list_item, но ограничение цифр — доп. страховка).
"""
_NUM_ITEM_RE  = re.compile(r'^(\d{1,2})[.)]?\s+')
_DATE_START_RE = re.compile(r'^\d{2}\.\d{2}\.\d{4}')  # DD.MM.YYYY в начале блока


class _TextPrefixItem:
    """Обёртка элемента Docling: добавляет числовой префикс к тексту при рендере.

    Используется в _fix_reading_order когда ненумерованный item вставляется
    в пробел нумерованного списка (OCR пропустил число).
    """
    __slots__ = ("_item", "_prefix")

    def __init__(self, item, prefix: str) -> None:
        self._item   = item
        self._prefix = prefix

    @property
    def text(self) -> str:
        return self._prefix + (getattr(self._item, "text", None) or "")

    def __getattr__(self, name: str):
        if name in ("_item", "_prefix"):
            raise AttributeError(name)
        return getattr(self._item, name)


class _TextAppendItem:
    """Добавляет суффикс к тексту элемента (для патроним-фикса)."""
    __slots__ = ("_item", "_suffix")

    def __init__(self, item, suffix: str) -> None:
        self._item   = item
        self._suffix = suffix

    @property
    def text(self) -> str:
        return (getattr(self._item, "text", None) or "") + self._suffix

    def __getattr__(self, name: str):
        if name in ("_item", "_suffix"):
            raise AttributeError(name)
        return getattr(self._item, name)


class _TextTruncateItem:
    """Возвращает текст элемента ДО позиции end (для разбивки блоков на абзацы)."""
    __slots__ = ("_item", "_end")

    def __init__(self, item, end: int) -> None:
        self._item = item
        self._end  = end

    @property
    def text(self) -> str:
        return (getattr(self._item, "text", None) or "")[:self._end]

    def __getattr__(self, name: str):
        if name in ("_item", "_end"):
            raise AttributeError(name)
        return getattr(self._item, name)


class _TextSliceItem:
    """Возвращает текст элемента начиная с указанной позиции (для патроним-фикса)."""
    __slots__ = ("_item", "_start")

    def __init__(self, item, start: int) -> None:
        self._item  = item
        self._start = start

    @property
    def text(self) -> str:
        return (getattr(self._item, "text", None) or "")[self._start:]

    def __getattr__(self, name: str):
        if name in ("_item", "_start"):
            raise AttributeError(name)
        return getattr(self._item, name)


class _TextInsertItem:
    """Вставляет строку в позицию pos, опционально усекает суффикс (truncate_at >= 0)."""
    __slots__ = ("_item", "_pos", "_insertion", "_truncate_at")

    def __init__(self, item, pos: int, insertion: str, truncate_at: int = -1) -> None:
        self._item        = item
        self._pos         = pos
        self._insertion   = insertion
        self._truncate_at = truncate_at

    @property
    def text(self) -> str:
        t = getattr(self._item, "text", None) or ""
        if 0 <= self._truncate_at < len(t):
            t = t[:self._truncate_at]
        return t[:self._pos] + self._insertion + t[self._pos:]

    def __getattr__(self, name: str):
        if name in ("_item", "_pos", "_insertion", "_truncate_at"):
            raise AttributeError(name)
        return getattr(self._item, name)


class _ForceListItem:
    """Форсирует рендер text/paragraph как нумерованного list_item (для ПРОСИТ СУД)."""
    __slots__ = ("_item", "_prefix")

    def __init__(self, item, prefix: str) -> None:
        self._item   = item
        self._prefix = prefix

    @property
    def text(self) -> str:
        return self._prefix + (getattr(self._item, "text", None) or "")

    @property
    def label(self):
        class _L:
            value = "list_item"
        return _L()

    def __getattr__(self, name: str):
        if name in ("_item", "_prefix"):
            raise AttributeError(name)
        return getattr(self._item, name)


class _TextSliceAndInsertItem:
    """Slice от start + вставка insertion после insert_pos (в sliced-координатах)."""
    __slots__ = ("_item", "_start", "_insert_pos", "_insertion")

    def __init__(self, item, start: int, insert_pos: int, insertion: str) -> None:
        self._item       = item
        self._start      = start
        self._insert_pos = insert_pos
        self._insertion  = insertion

    @property
    def text(self) -> str:
        sliced = (getattr(self._item, "text", None) or "")[self._start:]
        p = self._insert_pos
        return sliced[:p] + self._insertion + sliced[p:]

    def __getattr__(self, name: str):
        if name in ("_item", "_start", "_insert_pos", "_insertion"):
            raise AttributeError(name)
        return getattr(self._item, name)


"""
Паттерн: блок начинается с отчества в тв. падеже (Игоревичем, Петровичем).
Используем /S+? вместо [А-ЯЁ][а-яё]*? — EasyOCR иногда подмешивает Latin-символы
(Latin 'e' вместо Cyrillic 'е'), поэтому кириллический char-class ненадёжен.
[еe] в суффиксе покрывает оба варианта написания «е».
"""
_PATRON_CONT_RE = re.compile(
    r'^(\S+?(?:ич[еe]м|вич[еe]м)(?:\s*\([^)]{2,30}\))?)\s+(.+)',
    re.DOTALL,
)

"""
Паттерн для поиска позиции вставки патронима в блоке A:
имя в тв. падеже (оканчивается на «ем»/«еем» с возможными Latin-символами)
непосредственно перед глаголом «был»/«была»/«были»/«заключен».
Захват: group(1) = слово с именем, позиция вставки = m.end(1).
"""
_FN_BEFORE_VERB_RE = re.compile(
    r'(\S+?[еe]{1,2}м)\s+(?=(?:был[аи]?\b|заключен\b))',
    re.IGNORECASE,
)

"""
Паттерн: блок A заканчивается "лимитом выдачи в размере" — OCR разорвал фразу.
Используем /S*keyword/S* для устойчивости к OCR-артефактам (Latin B вместо Cyrillic в и т.п.).
"""
_LIMIT_SUFFIX_RE  = re.compile(
    r'\s+\S*лимит\S*\s+\S*выдач\S*\s+\S+\s+\S*размер\S*\s*$',
    re.IGNORECASE,
)
_CREDIT_LINE_RE   = re.compile(r'\S*кредит\S*\s+\S*лини\S+', re.IGNORECASE)


def _fix_reading_order(all_items: list) -> tuple[list, set]:
    """
    Пост-обработка порядка элементов после word_order.

    Возвращает (result, continuation_ids):
      result          — переупорядоченный список элементов
      continuation_ids — set id() элементов, которые нужно безусловно слить
                         с предыдущим параграфом при рендере (Fix 2 date-start)
    """
    result           = list(all_items)
    n                = len(result)
    continuation_ids: set = set()

    def _lbl(i):
        return _label_str(result[i][0])

    def _pg(i):
        pv = (getattr(result[i][0], "prov", None) or [None])[0]
        return int(getattr(pv, "page_no", -1)) if pv else -1

    def _txt(i):
        return (getattr(result[i][0], "text", None) or "").strip()

    # Fix 1: ALL-CAPS section_header перед строчным subtitle
    for i in range(n - 1):
        if _lbl(i) != "section_header" or _lbl(i + 1) != "section_header":
            continue
        if _pg(i) != _pg(i + 1):
            continue
        ti, tnext  = _txt(i), _txt(i + 1)
        ai    = [c for c in ti    if c.isalpha()]
        anext = [c for c in tnext if c.isalpha()]
        if ai and ai[0].islower() and anext and all(c.isupper() for c in anext):
            result[i], result[i + 1] = result[i + 1], result[i]
            log.info("fix_order: swap '%s' ↔ '%s'", ti[:40], tnext[:40])

    """
    section_header «по делу №» стоит ПОСЛЕ label:content (КРЕДИТОР:/ДОЛЖНИК:)
    В судебных заявлениях «по делу» всегда предшествует блокам сторон.
    Docling иногда путает порядок из-за близости по Y-координате.
    """
    _DELO_RE = re.compile(r'^\s*по\s+делу\b', re.IGNORECASE)
    for i in range(n - 1):
        if _lbl(i) != "text" or _lbl(i + 1) not in ("section_header", "text"):
            continue
        if _pg(i) != _pg(i + 1):
            continue
        if not _DELO_RE.match(_txt(i + 1)):
            continue
        if split_label_content(_txt(i)) is None:
            continue
        result[i], result[i + 1] = result[i + 1], result[i]
        log.info("fix_order: по_делу swap %r ↔ %r", _txt(i)[:40], _txt(i + 1)[:40])
    """
    текстовый блок с датой DD.MM.YYYY стоит ПОСЛЕ незаконченного блока
    Docling иногда разбивает один абзац на части и переставляет их:
      "Игоревичем (Заемщик) Заемщику..." → "20.12.2019 между..."
    Корректный порядок: дата → продолжение.
    Признак неправильного порядка: предыдущий ТЕКСТОВЫЙ блок не заканчивается на «.»/«!»/«?»
    Важно: между двумя текстовыми блоками могут быть section_header / picture —
    ищем не строго i-1, а ближайший предшествующий текстовый блок на той же странице.
    """
    _TEXT_LABELS = frozenset({"text", "paragraph"})
    for i in range(1, n):
        if _lbl(i) not in _TEXT_LABELS:
            continue
        curr = _txt(i)
        date_m = _DATE_START_RE.match(curr)
        if not date_m:
            continue
        # Ищем ближайший предшествующий текстовый блок на той же странице
        prev_i = None
        for k in range(i - 1, max(i - 6, -1), -1):
            if _lbl(k) in _TEXT_LABELS and _pg(k) == _pg(i):
                prev_i = k
                break
            if _pg(k) != _pg(i):
                break
        if prev_i is None:
            log.debug("fix_order: date-start не нашёл текстового предшественника для %r", curr[:40])
            continue
        prev = _txt(prev_i)

        # «17.09.1987, место рождения:» — дата рождения, за датой стоит запятая.
        # Это НЕ начало предложения о договоре → не переставляем.
        date_match_end = date_m.end()
        rest = curr[date_match_end:].lstrip()
        if rest.startswith(','):
            log.debug("fix_order: date-start пропуск (дата рождения) %r", curr[:50])
            continue

        # Предыдущий блок завершает ПОЛНОЕ предложение → не переставляем.
        # «:» НЕ включаем: OCR часто ставит «:» вместо «.», а «средств:» перед
        # датой договора — это как раз наш случай (нужно переставить).
        if prev.endswith(('.', '!', '?', '»', ';')):
            log.debug("fix_order: date-start пропуск (предш. заканчивается на терминатор) %r", prev[-20:])
            continue

        """
        Предыдущий блок — предложение-ВВЕДЕНИЕ списка: «…подтверждается
        нижеследующим», «…в следующем размере:», «…следующим образом:».
        Блок с датой и есть это «последующее» содержимое — он должен идти
        ПОСЛЕ введения, поэтому НЕ переставляем. Двоеточие необязательно:
        OCR часто его теряет («подтверждается нижеследующим» без «:»).
        """
        if re.search(r'(следующ\w*|образом)\s*[:.]?\s*$', prev, re.IGNORECASE):
            log.debug("fix_order: date-start пропуск (введение списка) %r", prev[-40:])
            continue

        # Предыдущий блок — короткое ALL-CAPS ФИО (2-4 слова, только буквы и пробелы).
        # «ХЛИЯН ЕЛЕНА ОГАНОВНА» — это значение поля «Должник:», а не незаконченная фраза.
        _prev_words = prev.split()
        _prev_alpha = [c for c in prev if c.isalpha()]
        if (2 <= len(_prev_words) <= 4
                and _prev_alpha
                and all(c.isupper() for c in _prev_alpha)
                and all(c.isalpha() or c.isspace() or c in '-' for c in prev)):
            log.debug("fix_order: date-start пропуск (ALL-CAPS ФИО) %r", prev[:50])
            continue

        log.debug("fix_order: date-start кандидат: prev=%r curr=%r (prev_i=%d, i=%d)",
                  prev[:50], curr[:50], prev_i, i)

        # Переставляем: все элементы между prev_i и i сдвигаются на 1 вперёд
        saved = result[i]
        for k in range(i, prev_i, -1):
            result[k] = result[k - 1]
        result[prev_i] = saved
        log.info("fix_order: date-start swap %r ↔ %r", prev[:50], curr[:50])
        # Элемент на позиции prev_i+1 — прямое продолжение вставленного блока
        # (это и есть тот «Игоревичем...», что шёл ДО даты в Docling).
        # Помечаем его для безусловного слияния в renderer.
        if prev_i + 1 < n and _lbl(prev_i + 1) in _TEXT_LABELS:
            b_raw = _txt(prev_i + 1)   # raw text из Docling (до OCR-фиксов)
            pm = _PATRON_CONT_RE.match(b_raw)
            if pm:
                # Блок B начинается с отчества в тв. падеже (Игоревичем (Заемщик)).
                # Переносим его в конец блока A — восстанавливаем полное ФИО.
                patron     = pm.group(1)   # "Игоревичем (Заемщик)"
                b_start    = pm.start(2)   # позиция в b_raw (stripped)
                a_orig     = result[prev_i]
                b_orig     = result[prev_i + 1]
                # Корректируем b_start до позиции в сыром тексте оригинального элемента
                # (b_raw = strip(original.text), поэтому start может сдвинуться).
                orig_text  = getattr(b_orig[0], "text", None) or ""
                g2_prefix  = pm.group(2)[:20]   # первые 20 символов группы 2
                real_start = orig_text.find(g2_prefix)
                if real_start >= 0:
                    b_start = real_start
                a_text    = getattr(a_orig[0], "text", None) or ""

                # ищем имя перед "был" - вставляем патроним после имени.
                fn_match   = _FN_BEFORE_VERB_RE.search(a_text)
                # ищем «лимитом выдачи в размере» в конце блока A -
                # переносим в блок B после «кредитной линии».
                lim_match  = _LIMIT_SUFFIX_RE.search(a_text)
                lim_text   = a_text[lim_match.start():].strip() if lim_match else None
                truncate_a = lim_match.start() if lim_match else -1

                if fn_match:
                    ins_pos = fn_match.end(1)
                    result[prev_i] = (
                        _TextInsertItem(a_orig[0], ins_pos, " " + patron, truncate_at=truncate_a),
                        a_orig[1],
                    )
                    log.info("fix_order: patron-insert после %r (pos=%d) truncate=%d в %r",
                             fn_match.group(1), ins_pos, truncate_a, a_text[:60])
                else:
                    result[prev_i] = (_TextAppendItem(a_orig[0], " " + patron), a_orig[1])
                    log.info("fix_order: patron-append (имя не найдено) в %r", a_text[:60])

                # если нашли «лимитом» - вставляем его в блок B после «кредитной линии»
                b_sliced_text = orig_text[b_start:]
                cl_match = _CREDIT_LINE_RE.search(b_sliced_text) if lim_text else None
                if lim_text and cl_match:
                    result[prev_i + 1] = (
                        _TextSliceAndInsertItem(
                            b_orig[0], b_start,
                            cl_match.end(),
                            " с " + lim_text,
                        ),
                        b_orig[1],
                    )
                    log.info("fix_order: лимит-move %r → в блок B после 'кредитной линии'", lim_text)
                else:
                    result[prev_i + 1] = (_TextSliceItem(b_orig[0], b_start), b_orig[1])
                log.info("fix_order: B теперь %r", b_sliced_text[:50])
            else:
                log.debug("fix_order: mark continuation %r", b_raw[:40])
            continuation_ids.add(id(result[prev_i + 1][0]))

    # сортировка numbered list_items + заполнение пробелов ненумерованными
    i = 0
    while i < n:
        if _lbl(i) != "list_item":
            i += 1
            continue
        j, pg = i + 1, _pg(i)
        while j < n and _lbl(j) == "list_item" and _pg(j) == pg:
            j += 1
        run_len = j - i
        if run_len < 3:
            i = j
            continue

        txts       = [_txt(k) for k in range(i, j)]
        numbered   = [(int(m.group(1)), k)
                      for k, t in enumerate(txts)
                      for m in [_NUM_ITEM_RE.match(t)] if m]
        unnumbered = [k for k in range(run_len) if not _NUM_ITEM_RE.match(txts[k])]

        if len(numbered) < max(3, run_len // 2):
            i = j
            continue

        numbered_sorted = sorted(numbered)
        nums_only       = [num for num, _ in numbered]
        sorted_nums     = sorted(nums_only)

        # Есть ли пробелы в нумерации?
        has_gap = any(sorted_nums[k + 1] > sorted_nums[k] + 1
                      for k in range(len(sorted_nums) - 1))

        already_sorted = nums_only == sorted_nums
        if already_sorted and not (has_gap and unnumbered):
            i = j
            continue  # порядок верный, пробелов нет

        # Строим новый порядок: вставляем ненумерованные в пробелы нумерации
        final_order:  list[int]      = []
        gap_prefixes: dict[int, str] = {}   # gap_k → "N " для заполнителей пробелов
        unnum_used  = 0
        prev_num    = 0
        for num, k in numbered_sorted:
            gap_size = num - prev_num - 1
            for _ in range(gap_size):
                if unnum_used < len(unnumbered):
                    gap_k   = unnumbered[unnum_used]
                    gap_num = prev_num + 1
                    final_order.append(gap_k)
                    gap_prefixes[gap_k] = f"{gap_num} "
                    log.info("fix_order: пункт %d пропущен OCR, заполняем %r",
                             gap_num, txts[gap_k][:60])
                    unnum_used += 1
            final_order.append(k)
            prev_num = num
        # Оставшиеся ненумерованные — автонумерация продолжением
        _remaining_unnum = unnumbered[unnum_used:]
        if _remaining_unnum and numbered:
            _next_auto = max(n for n, _ in numbered) + 1
            for _gap_k in _remaining_unnum:
                final_order.append(_gap_k)
                gap_prefixes[_gap_k] = f"{_next_auto} "
                log.info("fix_order: автонумер %d %r", _next_auto, txts[_gap_k][:60])
                _next_auto += 1
        else:
            final_order.extend(_remaining_unnum)

        reordered = [
            (_TextPrefixItem(result[i + k][0], gap_prefixes[k]), result[i + k][1])
            if k in gap_prefixes else result[i + k]
            for k in final_order
        ]
        for k, item_lvl in enumerate(reordered):
            result[i + k] = item_lvl

        if not already_sorted:
            log.info("fix_order: sorted %d numbered list items on page %d",
                     len(numbered), pg)
        if has_gap and unnum_used > 0:
            log.info("fix_order: filled %d gap(s) with unnumbered items on page %d",
                     unnum_used, pg)
        i = j

    # list_item стоит ДО heading/text "ПРОСИТ СУД" — меняем местами
    # и нумеруем следующие заглавные пункты (1, 2, ...).
    _PROSIT_RE = re.compile(r'\bПРОСИТ\b', re.IGNORECASE)
    for i in range(n - 1):
        if _lbl(i) != "list_item":
            continue
        if _lbl(i + 1) not in ("section_header", "text", "title"):
            continue
        if _pg(i) != _pg(i + 1):
            continue
        if not _PROSIT_RE.search(_txt(i + 1)):
            continue
        result[i], result[i + 1] = result[i + 1], result[i]
        log.info("fix_order: ПРОСИТ_СУД swap %r ↔ %r", _txt(i)[:40], _txt(i + 1)[:40])
        # Нумеруем следующие заглавные пункты (list_item И text/paragraph)
        _PROSIT_LBLS = frozenset({"list_item", "text", "paragraph"})
        num_counter = 1
        for j in range(i + 1, min(i + 10, n)):
            cur_lbl = _lbl(j)
            if cur_lbl not in _PROSIT_LBLS:
                break
            t_j = _txt(j)
            a_j = [c for c in t_j if c.isalpha()]
            if not a_j or a_j[0].islower():
                break   # подпункт — стоп
            if not _NUM_ITEM_RE.match(t_j):
                if cur_lbl == "list_item":
                    result[j] = (_TextPrefixItem(result[j][0], f"{num_counter} "), result[j][1])
                else:
                    result[j] = (_ForceListItem(result[j][0], f"{num_counter} "), result[j][1])
                log.info("fix_order: ПРОСИТ_СУД нумерует %d %r (lbl=%s)",
                         num_counter, t_j[:40], cur_lbl)
                num_counter += 1
        break  # один ПРОСИТ СУД на документ

    # разбиваем блоки где OCR объединил несколько абзацев.
    # Цепочечная разбивка: каждый паттерн применяется к результатам предыдущего.
    _PARA_SPLIT_RES = [
        # «...оборотных средств: Срок возврата кредита...» - два абзаца
        re.compile(r'[;:,.]\s+((?:Срок\s+)?возврата\s+кредита)', re.IGNORECASE),
        # «...не позднее 10.05.2028. За пользование...» - два абзаца
        re.compile(r'\.\s+(За\s+пользование)', re.IGNORECASE),
    ]
    new_result = []
    for _item, _level in result:
        segs = [_item]
        for _pat in _PARA_SPLIT_RES:
            new_segs: list = []
            for _seg in segs:
                _raw = getattr(_seg, "text", None) or ""
                _m   = _pat.search(_raw)
                if _m:
                    _p1 = _m.start(0) + 1   # конец первого абзаца (включая знак)
                    _p2 = _m.start(1)        # начало следующего абзаца
                    new_segs.append(_TextTruncateItem(_seg, _p1))
                    new_segs.append(_TextSliceItem(_seg, _p2))
                    log.info("fix_order: split at %r (pos %d/%d)",
                             _pat.pattern[:35], _p1, _p2)
                else:
                    new_segs.append(_seg)
            segs = new_segs
        for _seg in segs:
            new_result.append((_seg, _level))
    result = new_result
    n = len(result)

    return result, continuation_ids


def _join_fragments(all_items: list, continuation_ids: set) -> set:
    """
    Склейка фрагментированных текстовых блоков для любых документов.

    Docling при плохом качестве скана или нестандартном layout (угловые штампы,
    ФНС-формы, газетные колонки) разбивает один абзац на множество коротких блоков
    (1-2 слова каждый). Эта функция детектирует такие страницы и добавляет блоки
    в continuation_ids — рендер склеит их в один параграф.

    Активируется только на страницах, где >50% text/paragraph блоков «короткие»
    (≤ 4 слова). Реальное разделение на предложения контролируется _prev_unfinished
    в рендере: блок после «.» / «!» / «?» всегда начинает новый параграф.
    """
    
    
    """
    Порог 2 слова (а не 4): реквизитные строки «ИНН: 7728168971» были бы склеены 
    при пороге 4, хотя это структурные блоки. При пороге 2 склеиваем только однословные
    / двусловные OCR-фрагменты (характерные для штампов и ФНС-форм), не трогая реквизиты
    и короткие, но осмысленные блоки.
    """
    MAX_FRAG_WORDS = 2      # блок «короткий» если <= 2 слов
    FRAG_THRESHOLD = 0.50   # >50% коротких блоков на странице - "фрагментированная"
    MIN_BLOCKS     = 5      # минимум блоков для принятия решения
    SAFE_LABELS    = frozenset({"text", "paragraph"})

    def _pg(i: int) -> int:
        pv = (getattr(all_items[i][0], "prov", None) or [None])[0]
        return int(getattr(pv, "page_no", -1)) if pv else -1

    def _lbl(i: int) -> str:
        return _label_str(all_items[i][0])

    def _txt(i: int) -> str:
        return (getattr(all_items[i][0], "text", None) or "").strip()

    # Собираем статистику text/paragraph-блоков по страницам
    from collections import defaultdict
    page_idxs:  defaultdict[int, list[int]] = defaultdict(list)
    page_short: defaultdict[int, int]       = defaultdict(int)

    for i in range(len(all_items)):
        if _lbl(i) not in SAFE_LABELS:
            continue
        pg = _pg(i)
        if pg < 0:
            continue
        page_idxs[pg].append(i)
        if len(_txt(i).split()) <= MAX_FRAG_WORDS:
            page_short[pg] += 1

    # Определяем страницы с экстремальной фрагментацией
    fragmented: set[int] = set()
    for pg, idxs in page_idxs.items():
        if len(idxs) < MIN_BLOCKS:
            continue
        ratio = page_short[pg] / len(idxs)
        if ratio > FRAG_THRESHOLD:
            fragmented.add(pg)
            log.info("fragment-join: стр.%d — %.0f%% коротких блоков (%d/%d) → склейка",
                     pg, ratio * 100, page_short[pg], len(idxs))

    if not fragmented:
        return continuation_ids

    # Маркируем блоки: КАЖДЫЙ кроме первого на странице → кандидат на склейку.
    # Финальное решение принимает рендер через _prev_unfinished (не склеивает
    # после ".", "!", "?" — так границы предложений сохраняются).
    seen_page: set[int] = set()
    for i in range(len(all_items)):
        if _lbl(i) not in SAFE_LABELS:
            continue
        pg = _pg(i)
        if pg not in fragmented:
            continue
        if pg in seen_page:
            continuation_ids.add(id(all_items[i][0]))
        else:
            seen_page.add(pg)

    return continuation_ids


def _hdr_mergeable(text: str) -> bool:
    """Фрагмент можно вклеивать в строку шапки, только если он несёт содержание
    (кириллица или цифры) либо это пунктуация/№. Чисто-латинский буквенный токен
    (OCR-мусор «CRPAAMNPYNS», «WY») в кириллическую шапку не вклеиваем — иначе
    мусор прилипает к заголовку («CRPAAMNPYNS Приложение:»)."""
    has_cyr = bool(re.search(r'[А-Яа-яЁё]', text))
    has_dig = any(c.isdigit() for c in text)
    has_lat = bool(re.search(r'[A-Za-z]', text))
    return has_cyr or has_dig or not has_lat


def _merge_header_lines(all_items: list, page_sizes: dict, pdf_native: bool,
                        skip_indices: set) -> int:
    """Склейка пословно-разорванной шапки в цельные строки (убирает «лесенку»).

    Docling на письмах ФНС дробит строку адресата/штампа на блоки-слова с ОДНИМ
    top и растущим x0 («в / лице / Межрайонной / ИФНС / России / № / 26»); рендер
    даёт каждому свой left_indent → визуальная лесенка. Здесь для зоны ШАПКИ
    (сверху страницы до заголовка тела `_DOC_TITLE_RE`) блоки группируются в строки
    по вертикали, внутри строки склеиваются соседи ОДНОЙ колонки (малый гор. зазор)
    в левый блок; хвостовые фрагменты уходят в skip_indices. Границу колонок держит
    порог зазора (левый штамп ≠ правый адресат), а слияние — только для КОРОТКИХ
    фрагментов (≤3 слов): длинные абзацы тела не трогаются. Мутирует item.text и
    bbox.r левого блока. Возвращает число склеенных фрагментов."""
    from collections import defaultdict
    MAX_FRAG_WORDS = 3
    page_items: dict[int, list[int]] = defaultdict(list)
    stopped: set[int] = set()
    for idx, (item, _lvl) in enumerate(all_items):
        prov = getattr(item, "prov", None) or []
        if not prov:
            continue
        pg = int(getattr(prov[0], "page_no", -1))
        if pg < 0 or pg in stopped:
            continue
        text = (getattr(item, "text", None) or "").strip()
        if text and _DOC_TITLE_RE.match(text):
            stopped.add(pg)                  # заголовок тела — конец зоны шапки
            continue
        if idx in skip_indices or not text:
            continue
        if _label_str(item) not in ("text", "paragraph", "section_header",
                                     "title", "caption"):
            continue
        if getattr(prov[0], "bbox", None) is None:
            continue
        page_items[pg].append(idx)

    merged_total = 0
    for pg, idxs in page_items.items():
        if len(idxs) < 2:
            continue
        pw, ph = page_sizes.get(pg, (595.0, 842.0))
        recs = []                            # (idx, l, r, top, bottom)
        for idx in idxs:
            bbox = all_items[idx][0].prov[0].bbox
            l = float(getattr(bbox, "l", 0)); r = float(getattr(bbox, "r", 0))
            top, bot = _bbox_top_bottom(bbox, ph, pdf_native)
            recs.append((idx, l, r, top, bot))
        med_h = sorted(b - t for _, _, _, t, b in recs)[len(recs) // 2] or 8.0
        row_tol = max(3.0, 0.5 * med_h)
        col_gap = max(20.0, 2.5 * med_h)
        recs.sort(key=lambda x: x[3])        # по top (0 = верх страницы)
        rows: list[list] = []
        for rec in recs:
            if rows and abs(rec[3] - rows[-1][0][3]) <= row_tol:
                rows[-1].append(rec)
            else:
                rows.append([rec])
        for row in rows:
            if len(row) < 2:
                continue
            row.sort(key=lambda x: x[1])     # по левому краю
            groups: list[list] = [[row[0]]]
            for rec in row[1:]:
                if rec[1] - groups[-1][-1][2] <= col_gap:   # this.l - prev.r
                    groups[-1].append(rec)
                else:
                    groups.append([rec])
            for g in groups:
                # чисто-латинский мусор в строку шапки не вклеиваем (см. _hdr_mergeable)
                gm = [gr for gr in g
                      if _hdr_mergeable(all_items[gr[0]][0].text or "")]
                if len(gm) < 2:
                    continue
                frags = [all_items[gr[0]][0].text.strip() for gr in gm]
                # слияние только коротких фрагментов — тело (длинные блоки) не трогаем
                if any(len(f.split()) > MAX_FRAG_WORDS for f in frags):
                    continue
                lead_item = all_items[gm[0][0]][0]
                lead_item.text = " ".join(frags)
                try:
                    lead_item.prov[0].bbox.r = max(gr[2] for gr in gm)
                except Exception:
                    pass
                for gr in gm[1:]:
                    skip_indices.add(gr[0])
                merged_total += len(gm) - 1
    if merged_total:
        log.info("merge_header_lines: склеено %d фрагментов шапки", merged_total)
    return merged_total


# Основная функция построения DOCX

_LETTERHEAD_STOP_RE = re.compile(
    r"\b("
    r"кредитор|должник|заявлени[ея]|исковое|госпошлин\w*|"
    r"арбитражн\w*\s+управляющ|арбитражн\w*\s+суд|по\s+делу"
    r")\b",
    re.IGNORECASE,
)

"""
Название юрлица (продолжение значения "Заявитель/Кредитор: ..."). OCR нестабилен
и иногда метит такую строку как section_header — её НЕ нужно считать началом тела
"ОПРЕДЕЛЕНИЕ"/"МИНФИН..." сюда не попадают — они останутся стоп-маркерами тела. 
/w* ловит OCR-склейку "АКЦИОНЕРНОЕОБЩЕСТВО".
"""

_ENTITY_NAME_RE = re.compile(
    r'^\s*(акционерн\w*|публичн\w*|обществ\w*|ооо|пао|оао|зао|ао|банк)\b',
    re.IGNORECASE)

# Стоп-паттерны шапки: блоки ниже этой точки — не шапка
_HEADER_BODY_START_RE = re.compile(
    r"^\s*(заявлени[ея]|исковое\s+заявлени[ея]|требовани[ея]|"
    r"ходатайство|возражени[ея]|отзыв)\b"
    r"|^\s*(должник|финансовый\s+управляющий)\s*:",
    re.IGNORECASE,
)


def _render_native_two_column_header(
    doc,
    all_items: list,
    page_sizes: dict[int, tuple[float, float]],
    pdf_native: bool,
    page_medians: dict,
    page_left_min: dict,
    already_skipped: set | None = None,
    prepend_right_blocks: list | None = None,
    logo_data: tuple | None = None,  # (pil_img, pic_w_pt, img_w_inch) от letterhead
) -> set[int]:
    """
    Рендер двухколоночной шапки для нативных PDF (без логотипа).

    Шапка включает все элементы первой страницы до первого заголовка
    типа "ЗАЯВЛЕНИЕ" / "ПРОСИТ СУД" / "ТРЕБОВАНИЕ" (которые переходят
    на следующую страницу или стоят в середине страницы как section_header).

    Возвращает set индексов элементов которые были включены в шапку.
    """
    first_page = min(page_sizes.keys()) if page_sizes else 1
    pw, ph = page_sizes.get(first_page, (595.0, 842.0))
    if pw <= 0 or ph <= 0:
        return set()

    _already_skipped = already_skipped or set()

    # Собираем все элементы первой страницы до стоп-маркера
    header_items: list[tuple[int, object, object, float, float]] = []
    for idx, (item, level) in enumerate(all_items):
        if idx in _already_skipped:
            continue
        lbl = _label_str(item)
        if not lbl or lbl in SKIP_LABELS:
            continue
        if lbl in ("table", "picture", "figure", "image"):
            continue
        bbox, page_no = _item_bbox_page(item, first_page)
        if page_no != first_page or bbox is None:
            continue
        text = postprocess((getattr(item, "text", None) or "").strip())
        if not text:
            continue
        """
        section_header / title = начало тела документа. ИСКЛЮЧЕНИЕ — когда OCR
        ошибочно пометил как section_header название юрлица (продолжение значения
        «Заявитель/Кредитор: …»): такой блок НЕ начало тела, иначе хвост значения
        сыплется из шапки в тело слева. "ОПРЕДЕЛЕНИЕ"/"МИНФИН..." под исключение не
        подпадают и корректно ломают сбор шапки.
        """
        
        if (lbl in ("section_header", "title")
                and not _ENTITY_NAME_RE.match(text)):
            break
        # Стоп: явный маркер начала тела
        if _HEADER_BODY_START_RE.match(text):
            break
        top, _ = _bbox_top_bottom(bbox, ph, pdf_native)
        x0 = float(getattr(bbox, "l", 0))
        header_items.append((idx, item, level, x0, top))

    if len(header_items) < 2:
        return set()

    # Вертикальный фильтр: блоки с top_nat > 0.68 не являются частью
    # шапки — это Место/Дата рождения, которые word_order ставит
    # раньше стоп-маркера «Должник:», хотя физически ниже него.
    _max_top_nat = 0.68
    header_items = [(idx, item, lvl, x0, top)
                    for idx, item, lvl, x0, top in header_items
                    if top <= ph * _max_top_nat]

    if len(header_items) < 2:
        return set()

    # Определяем границу колонок через k-means с k=2.
    # Разбиваем все x0 на два кластера и берём точку между ними как границу.
    x0_raw = sorted(x0 for _, _, _, x0, _ in header_items)
    if len(x0_raw) < 2:
        return set()
    # Инициализация: левый центр = медиана нижней половины, правый = верхней
    mid = len(x0_raw) // 2
    c_left  = sum(x0_raw[:mid]) / mid
    c_right = sum(x0_raw[mid:]) / max(len(x0_raw) - mid, 1)
    for _ in range(10):  # итерации k-means
        left_pts  = [x for x in x0_raw if abs(x - c_left) <= abs(x - c_right)]
        right_pts = [x for x in x0_raw if abs(x - c_right) < abs(x - c_left)]
        if not left_pts or not right_pts:
            break
        c_left  = sum(left_pts)  / len(left_pts)
        c_right = sum(right_pts) / len(right_pts)
    col_boundary = (c_left + c_right) / 2.0
    separation   = c_right - c_left
    log.debug("native_header: c_left=%.0f c_right=%.0f separation=%.0f boundary=%.0f",
              c_left, c_right, separation, col_boundary)
    # Колонки должны быть достаточно разделены (> 10% ширины страницы)
    if separation < pw * 0.10 or col_boundary <= 0:
        log.info("native_header: колонки не разделены (separation=%.0f < %.0f) — пропуск",
                 separation, pw * 0.10)
        return set()

    left_items  = [(idx, item, lvl, x0, top)
                   for idx, item, lvl, x0, top in header_items if x0 < col_boundary]
    right_items = [(idx, item, lvl, x0, top)
                   for idx, item, lvl, x0, top in header_items if x0 >= col_boundary]

    # Если нет разделения — не рендерим как двухколоночную
    if not left_items or not right_items:
        log.info("native_header: нет двух колонок (left=%d right=%d) — пропуск",
                 len(left_items), len(right_items))
        return set()

    log.info("native_header: left=%d right=%d блоков boundary=%.0fpt sep=%.0fpt",
             len(left_items), len(right_items), col_boundary, separation)

    skip_set = {idx for idx, _, _, _, _ in header_items}

    # Режим анализа (doc=None) — только возвращаем skip_set без рендера
    if doc is None:
        return skip_set

    text_w_inch = (pw - 2 * MARGIN_INCH * 72) / 72

    def _make_block(item, top: float) -> tuple[float, dict]:
        text = postprocess((getattr(item, "text", None) or "").strip())
        if not text:
            return top, {}
        lbl = _label_str(item)
        # Жирные только структурные заголовки. НЕ жирним ALL-CAPS аббревиатуры/
        # реквизиты (ИНН/БИК/КПП) и имена — это давало ложную жирность в шапке.
        bold = lbl in ("title", "section_header")
        return top, {"text": text, "font_pt": BODY_PT, "bold": bold, "italic": False}

    # Собираем блоки правой колонки по top.
    # Блоки из left_items ("Заявитель (кредитор):") встраиваем по реальному top —
    # они должны стоять inline с соответствующим right-блоком.
    topped_blocks: list[tuple[float, dict]] = []

    if prepend_right_blocks:
        top_first = min((top for _, _, _, _, top in right_items), default=0.0)
        for i, blk in enumerate(prepend_right_blocks):
            topped_blocks.append((top_first - len(prepend_right_blocks) + i, blk))

    for _, item, _, _, top in sorted(right_items, key=lambda e: e[4]):
        _, blk = _make_block(item, top)
        if blk:
            topped_blocks.append((top, blk))

    for _, item, _, _, top in left_items:
        _, blk = _make_block(item, top)
        if blk:
            topped_blocks.append((top, blk))

    topped_blocks.sort(key=lambda x: x[0])


    """
    coplanar merge (+-8pt) — объединяем блоки на одной строке.
    2 случая:
      1) label, заканчивающийся ':', + content → label + "\n" + content
      2) осколок (<=2 слова, без пунктуации) + следующий coplanar → prefix + " " + content
    Для 1 случая: используем "\n" чтобы они рендерились на двух строках.
    Для 2 случая: пробел — они читались как одна строка.
    """
    _RE_FRAG = re.compile(r'^[А-ЯЁа-яёA-Za-z0-9\-]+(?:\s+[А-ЯЁа-яёA-Za-z0-9\-]+)?$')
    merged: list[tuple[float, dict]] = []
    used = set()
    for i, (t_i, blk_i) in enumerate(topped_blocks):
        if i in used:
            continue
        text_i = blk_i.get("text", "").strip()
        words_i = text_i.split()
        # 1) label + ':' — объединяем через \t для рендера label | value в строку
        # Ищем coplanar (+-8pt) как вперёд, так и назад по списку.
        if text_i.rstrip().endswith(":"):
            partner_j = None
            # Ищем вперёд
            for j in range(i + 1, len(topped_blocks)):
                if j in used:
                    continue
                t_j, blk_j = topped_blocks[j]
                if abs(t_j - t_i) <= 8.0:
                    partner_j = j
                    break
                if t_j - t_i > 8.0:
                    break
                
            """
            # Ищем назад если вперёд не нашли.
            # Партнёр назад — первая строка текста не является label (нет "Буква:").
            # НЕ проверяем used — партнёр мог уже попасть в merged как самостоятельный
            # блок (АКЦИОНЕРНОЕ ОБЩЕСТВО добавляется раньше Заявителя при сортировке).
            """
            if partner_j is None:
                for j in range(i - 1, -1, -1):
                    t_j, blk_j = topped_blocks[j]
                    blk_j_text = blk_j.get("text", "")
                    first_line = blk_j_text.split("\n")[0]
                    has_label_colon = bool(re.search(r'[А-ЯЁа-яёA-Za-z]\s*:', first_line))
                    if abs(t_j - t_i) <= 8.0 and not has_label_colon:
                        partner_j = j
                        break
                    if t_i - t_j > 8.0:
                        break
            if partner_j is not None:
                t_j, blk_j = topped_blocks[partner_j]
                combined = text_i + "\t" + blk_j.get("text", "")
                # Если партнёр уже в merged как самостоятельный блок — удаляем его
                partner_text = blk_j.get("text", "")
                for mi in range(len(merged) - 1, -1, -1):
                    if merged[mi][1].get("text", "") == partner_text:
                        merged.pop(mi)
                        break
                merged.append((t_i, {**blk_i, "text": combined, "bold": True}))
                used.add(i); used.add(partner_j)
                continue
        # 2) осколок — coplanar с СЛЕДУЮЩИМ (+-4pt), приклеить как префикс
        elif (len(words_i) <= 2 and _RE_FRAG.match(text_i)
              and text_i and text_i[-1] not in ".,:;»"):
            for j in range(i + 1, len(topped_blocks)):
                if j in used:
                    continue
                t_j, blk_j = topped_blocks[j]
                if abs(t_j - t_i) <= 4.0:
                    combined = text_i + " " + blk_j.get("text", "")
                    merged.append((t_i, {**blk_j, "text": combined}))
                    used.add(i); used.add(j)
                    break
        if i not in used:
            merged.append((t_i, blk_i))
            used.add(i)


    """
    # исправляем два OCR-артефакта в последовательности блоков:
    #   1) "по" идёт ПОСЛЕ "доверенности: Аракелян..." (разница top 2pt)
    #      - переставляем: "по доверенности: Аракелян..."
    #   2) "по доверенности: Аракелян..." стоит сразу после "Представитель АО..."
    #      - объединяем в одну строку через пробел
    """
    _RE_DOVERENNOSTI_START = re.compile(r'(?i)^доверенности\s*:.+')
    _RE_PRED_START = re.compile(r'(?i)^представитель\b')
    fused: list[tuple[float, dict]] = []
    for i, (t_i, blk_i) in enumerate(merged):
        text_i = blk_i.get("text", "").strip()
        # 1) "по" + предыдущий "доверенности: ..." - "по доверенности: ..."
        if (fused
                and text_i.lower() == "по"
                and abs(t_i - fused[-1][0]) <= 15.0):
            prev_t, prev_blk = fused[-1]
            prev_txt = prev_blk.get("text", "")
            if _RE_DOVERENNOSTI_START.match(prev_txt):
                fused[-1] = (prev_t, {**prev_blk, "text": "по " + prev_txt})
                continue
        # 2) "по доверенности: ..." + предыдущий "Представитель АО..." - один блок
        if (fused
                and re.match(r'(?i)^по\s+доверенности\b', text_i)
                and abs(t_i - fused[-1][0]) <= 25.0):
            prev_t, prev_blk = fused[-1]
            prev_txt = prev_blk.get("text", "")
            if _RE_PRED_START.match(prev_txt):
                fused[-1] = (prev_t, {**prev_blk, "text": prev_txt + " " + text_i})
                continue
        fused.append((t_i, blk_i))

    # "Представитель АО..." + следующий "по доверенности: Аракелян...""
    fused2: list[tuple[float, dict]] = []
    for t_i, blk_i in fused:
        text_i = blk_i.get("text", "").strip()
        if (fused2
                and re.match(r'(?i)^по\s+доверенности\b', text_i)
                and _RE_PRED_START.match(fused2[-1][1].get("text", ""))
                and abs(t_i - fused2[-1][0]) <= 30.0):
            prev_t, prev_blk = fused2[-1]
            fused2[-1] = (prev_t, {**prev_blk,
                "text": prev_blk.get("text", "") + " " + text_i})
        else:
            fused2.append((t_i, blk_i))
    fused = fused2

    # Склейка конкретного шаблона «должника) управляющим»» с предыдущим.
    # Этот блок — хвост фразы "Назначение платежа: ... по согласованию финансовым
    # управляющим" — оторванный потому что OCR разбил длинную строку на два блока.
    fused3: list[tuple[float, dict]] = []
    for t_i, blk_i in fused:
        text_i = blk_i.get("text", "").strip()
        if (fused3
                and re.match(r'(?i)^должника\)', text_i)
                and abs(t_i - fused3[-1][0]) <= 100.0):
            prev_t, prev_blk = fused3[-1]
            fused3[-1] = (prev_t, {**prev_blk,
                "text": prev_blk.get("text", "") + " " + text_i})
        else:
            fused3.append((t_i, blk_i))
    fused = fused3

    # исправляем OCR-артефакты в ОБЪЕДИНЁННЫХ блоках.
    # некоторые блоки склеены из нескольких частей — фиксы
    # которые не сработали на отдельных частях могут сработать на склейке.
    fused = [(t, {**blk, "text": postprocess(blk.get("text", ""))}) for t, blk in fused]

    # Разбивка блока "ИНН представителя ... Адрес представителя ..." на строки.
    # Аналогично OCR-фиксам для реквизитов юрлица, но специфично для шапки.
    _RE_ADDR_REP = re.compile(
        r'\s+(Адрес\s+представителя)', re.IGNORECASE)
    expanded: list[tuple[float, dict]] = []
    for t, blk in fused:
        txt = blk.get("text", "")
        txt2 = _RE_ADDR_REP.sub(r'\n\1', txt)
        if txt2 != txt:
            expanded.append((t, {**blk, "text": txt2}))
        else:
            expanded.append((t, blk))

    # space_before: ставим 6pt только при вертикальном разрыве > 40pt
    # (реквизиты банка идут подряд с 12pt разрывом — не нужен отступ)
    final_pairs: list[tuple[float, dict]] = []
    prev_top: float | None = None
    for t, blk in expanded:
        b = dict(blk)
        b["space_before"] = 6.0 if (prev_top is not None and t - prev_top > 40) else 0.0
        final_pairs.append((t, b))
        prev_top = t

    """
    "Заявитель (кредитор):" — полноширинная строка "метка | значение". Он должен
    стоять в СВОЕЙ Y-позиции (сразу после "ДЕЛО №"), а не в конце шапки. Поэтому
    делим правоколоночные блоки на above (выше Заявителя) и below (ниже) и рендерим
    тремя сегментами: [логотип | above] - "Заявитель: значение" - [пусто | below].
    """
    _ZAYAV_RE = re.compile(r'^\s*Заявитель\s*\(кредитор\)\s*:', re.IGNORECASE)
    _zayav_idx: int | None = None
    for i, (_t, blk) in enumerate(final_pairs):
        txt = blk.get("text", "")
        label_part = txt.split("\t")[0] if "\t" in txt else txt
        if _ZAYAV_RE.match(label_part):
            _zayav_idx = i
            break

    if _zayav_idx is not None:
        above_blocks = [b for _t, b in final_pairs[:_zayav_idx]]
        zayav_block: dict | None = final_pairs[_zayav_idx][1]
        below_blocks = [b for _t, b in final_pairs[_zayav_idx + 1:]]
    else:
        above_blocks = [b for _t, b in final_pairs]
        zayav_block = None
        below_blocks = []

    # Канонический порядок реквизитов: Р/с - КПП - ИНН - Кор/сч - БИК.
    # OCR/Y-сортировка иногда ставит КПП перед Р/с. Реордерим только СМЕЖНЫЙ прогон
    # реквизитных блоков (не трогая юрлицо-ИНН/ОГРН, отделённые адресом/местом).
    above_blocks = _reorder_requisite_blocks(above_blocks)
    below_blocks = _reorder_requisite_blocks(below_blocks)

    # Геометрия колонок
    if logo_data is not None:
        pil_img, pic_w_pt, img_w_inch = logo_data
        logo_col_pt = max(7.0 / 2.54 * 72, pic_w_pt)
        text_zone_pt = pw - 2 * MARGIN_INCH * 72
        content_col_pt = max(text_zone_pt - logo_col_pt, text_zone_pt * 0.55)
    else:
        logo_col_pt = 0.0
        content_col_pt = pw - 2 * MARGIN_INCH * 72

    def _right_col_cells(blocks: list[dict], with_logo: bool) -> list[dict]:
        """Ячейки строки: (логотип|пустая лого-колонка) + текст правой колонки."""
        out: list[dict] = []
        if logo_data is not None:
            out.append({
                "kind": "image",
                "image": pil_img if with_logo else None,
                "width_pt": logo_col_pt,
                "image_width_inch": img_w_inch if with_logo else 0.0,
                "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            })
        out.append({
            "kind": "text",
            "blocks": blocks,
            "width_pt": content_col_pt,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        })
        return out

    # логотип + above (адрес суда, "ДЕЛО №")
    add_header_row(doc, _right_col_cells(above_blocks, with_logo=True), text_w_inch)
    log.info("native_header: летерхед above=%d блоков, content=%.0fpt",
             len(above_blocks), content_col_pt)

    """
    "Заявитель (кредитор):" — в оригинале это МЕТКА у левого края (под лого-
       колонкой), а её ЗНАЧЕНИЕ (наименование юрлица, ИНН/ОГРН, дата, место
       нахождения) — в ПРАВОЙ колонке, выровнено с "Арбитражный суд". Поэтому при
       наличии логотипа рендерим строкой [метка в лого-колонке | значение справа];
       без логотипа — обычным полноширинным label-content.
    """
    if zayav_block is not None:
        txt = zayav_block.get("text", "")
        if "\t" in txt:
            lbl, val = txt.split("\t", 1)
        else:
            # В реальном пайплайне блок приходит без \t: "Заявитель (кредитор):
            # АКЦИОНЕРНОЕ ОБЩЕСТВО...". Отделяем метку по _ZAYAV_RE, иначе весь текст
            # ушёл бы в lbl (левая лого-колонка, жирным) — что и ломало вёрстку.
            m = _ZAYAV_RE.match(txt)
            if m:
                lbl, val = txt[:m.end()], txt[m.end():]
            else:
                lbl, val = txt, ""
        lbl = lbl.rstrip()
        val = val.lstrip()
        if not lbl.endswith(":"):
            lbl += ":"
        if logo_data is not None:
            add_header_row(doc, [
                {"kind": "text",
                 "blocks": [{"text": lbl, "font_pt": BODY_PT, "bold": True,
                             "italic": False, "space_before": 6.0}],
                 "width_pt": logo_col_pt,
                 "alignment": WD_ALIGN_PARAGRAPH.LEFT},
                {"kind": "text",
                 "blocks": [{"text": val, "font_pt": BODY_PT, "bold": False,
                             "italic": False, "space_before": 6.0}],
                 "width_pt": content_col_pt,
                 "alignment": WD_ALIGN_PARAGRAPH.LEFT},
            ], text_w_inch)
        else:
            content_items = [{"text": val, "font_pt": BODY_PT, "bold": False, "italic": False}]
            add_label_content_table(doc, lbl, content_items, text_w_inch, space_before=6.0)
        log.info("native_header: Заявитель в позиции %d: %r → %r", _zayav_idx, lbl, val[:50])

    # below (Адрес для отправки, Реквизиты, "При перечислении...", Представитель) —
    #    продолжение ЗНАЧЕНИЯ заявителя: в ПРАВОЙ колонке (пустая лого-колонка слева),
    #    выровнено с наименованием юрлица. Под логотипом — пусто (так в оригинале).
    if below_blocks:
        add_header_row(doc, _right_col_cells(below_blocks, with_logo=False), text_w_inch)
        log.info("native_header: below=%d блоков (правая колонка)", len(below_blocks))

    log.info("native_header: %d элементов в шапке", len(skip_set))
    return skip_set


# Паттерны для build_docx

# Детектор "Госпошлина" для склейки соседнего bbox суммы
_GOSPOSHLINA_RE = re.compile(r'^Госпошлин', re.IGNORECASE)

# Реквизитные строки — всегда выравниваются по левому краю (это маркированные
# данные банка/получателя, а не право-/центро-выключенный текст).
_REQ_LABEL_RE = re.compile(
    r'^\s*(Р\s*/\s*с|Р\s*/\s*сч|К\s*/\s*с|Кор\s*/\s*сч|ИНН|КПП|БИК|ОГРН|ОКПО|'
    r'Получатель|Расч[её]тный\s+счет|Корреспондентский\s+счет)\b',
    re.IGNORECASE)

# Канонический приоритет реквизитных строк: Р/с - КПП - ИНН - Кор/сч - БИК.
# "Реквизиты для перечисления..."/"Plc:" — это строка Р/с.
_REQ_PRIO_RES = [
    (re.compile(r'^\s*(?:Реквизиты\s+для\s+перечисл|Р\s*/\s*с|Plc|Расч[её]тный)', re.IGNORECASE), 0),
    (re.compile(r'^\s*КПП', re.IGNORECASE), 1),
    (re.compile(r'^\s*ИНН', re.IGNORECASE), 2),
    (re.compile(r'^\s*(?:Кор|Корреспондентский)', re.IGNORECASE), 3),
    (re.compile(r'^\s*БИК', re.IGNORECASE), 4),
]


def _req_prio(text: str):
    for rx, p in _REQ_PRIO_RES:
        if rx.match(text or ""):
            return p
    return None


def _reorder_requisite_blocks(blocks: list[dict]) -> list[dict]:
    """Реордерит СМЕЖНЫЕ прогоны реквизитных блоков в канонический порядок
    (Р/с, КПП, ИНН, Кор/сч, БИК). Несмежные (юрлицо-ИНН, отделённый адресом) и
    нереквизитные блоки не трогаем."""
    out = list(blocks)
    i = 0
    n = len(out)
    while i < n:
        if _req_prio(out[i].get("text", "")) is None:
            i += 1
            continue
        j = i
        while j < n and _req_prio(out[j].get("text", "")) is not None:
            j += 1
        if j - i >= 2:
            out[i:j] = sorted(out[i:j], key=lambda b: _req_prio(b.get("text", "")))
        i = j
    return out


# Маркер начала штампа электронной подписи
_ES_STAMP_MARKER_RE = re.compile(
    r'Электронн\w*\s+подпис\w*\s+действительн', re.IGNORECASE)

# Детекторы строки подписи
# "«Представитель по доверенности" / "Представитель КРЕДИТОРА по доверенности" /
# "Представитель АО ... по доверенности" — допускаем слова между.
_REPR_RE     = re.compile(r'представитель\b.{0,40}?\bпо\s+доверенности', re.IGNORECASE)
# Инициалы: "И.О.Фамилия" ИЛИ "Фамилия И.О."/"Фамилия И О:"",
# OCR теряет точки/ставит двоеточие).
_INITIALS_RE = re.compile(
    r'^(?:[А-ЯЁA-Z]\.\s*[А-ЯЁA-Z]\.\s*\S+'
    r'|[А-ЯЁ][а-яё]+\s+[А-ЯЁ][.\s]\s*[А-ЯЁ][.:\s]?)')

# Слова-стартеры новых абзацев — блокируют ложное _is_justify_cont слияние
_PARA_STARTERS_RE = re.compile(
    r'^(возврата\b|срок\s+возврата\b|за\s+пользование\b|'
    r'оценочная\b|расчет\b|включить\b|признать\b|'
    r'в\s+соответствии\s+с\b|согласно\b|исполнение\b)',
    re.IGNORECASE,
)


def _bbox_top_bottom(bbox, page_height: float, pdf_native: bool) -> tuple[float, float]:
    a = float(getattr(bbox, "t", 0.0))
    b = float(getattr(bbox, "b", 0.0))
    if pdf_native:
        top = page_height - max(a, b)
        bottom = page_height - min(a, b)
    else:
        top = min(a, b)
        bottom = max(a, b)
    return top, bottom


def _is_letterhead_stop(text: str) -> bool:
    return bool(_LETTERHEAD_STOP_RE.search(text.casefold()))


# ── Угловой штамп (письма ФНС и т.п.) ────────────────────────────────────────
# Дата/№ и «На №» — печатные токены, которые оставляем текстом; всё прочее
# без ≥3 кириллических букв подряд в зоне штампа — OCR-мусор от рукописи.
_CORNER_KEEP_RE = re.compile(r'\d{2}\.\d{2}\.\d{4}|№|\bНа\b', re.IGNORECASE)
_CYR3_RE = re.compile(r'[А-Яа-яЁё]{3}')
# Идентичность письма ФНС: угловой штамп детектится только если левый кластер
# содержит маркеры налоговой шапки. Якорь против ложных срабатываний на чужих
# двухколоночных шапках (ПСБ и т.п.) — режим построен исключительно для ФНС.
_FNS_ID_RE = re.compile(r'НАЛОГОВ|ИФНС|\bФНС\b|ИНСПЕКЦИЯ\s+ФЕДЕРАЛЬН', re.IGNORECASE)
# Заголовок тела документа (центрированный, ниже шапки). Если такой блок из-за
# рваной геометрии OCR попал в зону/правый кластер — его нельзя класть в колонку
# адресата: рендерим ниже обычным центрированным абзацем.
_DOC_TITLE_RE = re.compile(
    r'^(ЗАЯВЛЕНИЕ|ХОДАТАЙСТВО|ОТЗЫВ|ВОЗРАЖЕНИ|ЗАМЕЧАНИ|ОБЪЯСНЕНИ|ПОЯСНЕНИ|ЖАЛОБА)',
    re.IGNORECASE)


def _detect_corner_letterhead(
    text_zone: list[tuple[int, object, object, str]],  # (idx, item, bbox, text)
    pic_bbox,
    pw: float,
) -> tuple[list, list, object] | None:
    """Детект «углового штампа» письма ФНС: герб в ЛЕВОЙ половине, левый кластер
    содержит маркеры налоговой шапки, а текст верхней зоны образует два
    непересекающихся x-кластера — узкая левая колонка штампа (все r <= 0.55pw)
    и правая колонка адресата (все x0 >= 0.45pw), когерентная как колонка
    (право- либо левовыровнена).

    Возвращает (left_cluster, right_cluster, right_align) — списки элементов
    text_zone и выравнивание правой колонки (WD_ALIGN_PARAGRAPH.RIGHT/LEFT) —
    или None, если раскладка не похожа на угловой штамп ФНС."""
    if pic_bbox is None or pw <= 0:
        return None
    pic_cx = (float(getattr(pic_bbox, "l", 0)) + float(getattr(pic_bbox, "r", 0))) / 2
    if pic_cx >= pw * 0.5:
        return None
    left, right = [], []
    for entry in text_zone:
        bbox = entry[2]
        x0 = float(getattr(bbox, "l", 0))
        x1 = float(getattr(bbox, "r", 0))
        if x1 <= pw * 0.55:
            left.append(entry)
        elif x0 >= pw * 0.45:
            right.append(entry)
        else:
            return None        # блок пересекает обе половины — не двухколонник
    if len(left) < 3 or len(right) < 3:
        return None
    # Якорь идентичности: режим только для писем ФНС. Левый кластер обязан
    # содержать маркеры налоговой шапки — иначе это чужая двухколоночная
    # раскладка (ПСБ и т.п.), которую corner-рендер бы испортил.
    if not any(_FNS_ID_RE.search(e[3]) for e in left):
        return None
    # Правая колонка адресата должна быть КОГЕРЕНТНОЙ колонкой — либо
    # правовыровнена (адресат прижат к полю, как у Зайцева), либо
    # левовыровнена по единому x0 (Артемов/Геворгян/6л). Разброс обоих краёв
    # (метки+значения двумя подколонками) — не адресат, отклоняем.
    at_right = sum(1 for e in right
                   if float(getattr(e[2], "r", 0)) >= pw * 0.86)
    # Левовыровненность: наибольшая группа блоков с общим левым краем (±16pt).
    # Именно группа, а не min(x0): центрированный заголовок «ЗАЯВЛЕНИЕ» или
    # строка-вставка с меньшим x0 не должна ломать кластер адресата.
    x0s = [float(getattr(e[2], "l", 0)) for e in right]
    at_left = max(sum(1 for b in x0s if abs(b - a) <= 16.0) for a in x0s)
    if at_right < 0.6 * len(right) and at_left < 0.6 * len(right):
        return None
    # Выравнивание правой колонки для рендера: приоритет правовыровненности.
    right_align = (WD_ALIGN_PARAGRAPH.RIGHT if at_right >= at_left
                   else WD_ALIGN_PARAGRAPH.LEFT)
    return left, right, right_align


def _crop_page_image(dl_doc, page_no: int, l_pt: float, top_pt: float,
                     r_pt: float, bottom_pt: float, pw: float, ph: float):
    """Вырезает прямоугольник из page image (экранные pt-координаты).
    None если изображение страницы недоступно."""
    try:
        page = dl_doc.pages.get(page_no)
        pil = page.image.pil_image if page is not None and page.image else None
        if pil is None:
            return None
        sx = pil.width / pw
        sy = pil.height / ph
        box = (max(int(l_pt * sx), 0), max(int(top_pt * sy), 0),
               min(int(r_pt * sx), pil.width), min(int(bottom_pt * sy), pil.height))
        if box[2] - box[0] < 4 or box[3] - box[1] < 4:
            return None
        return pil.crop(box)
    except Exception as exc:
        log.debug("crop_page_image: %s", exc)
        return None


def _render_corner_letterhead(
    doc, dl_doc, page_no: int, pw: float, ph: float, pdf_native: bool,
    pil_img, pic_w_pt: float, img_w_inch: float, pic_idx: int,
    left: list, right: list, right_align=WD_ALIGN_PARAGRAPH.RIGHT,
) -> set[int]:
    """Рендер углового штампа: [герб + центрированная колонка | правая колонка
    адресата]. Выравнивание правой колонки — right_align (RIGHT у Зайцева,
    LEFT у Артемова/Геворгяна/6л). Печатный текст остаётся ТЕКСТОМ; рукописные
    дата/№ (OCR-мусор без кириллицы) заменяются картинкой-кропом строки из
    скана. Возвращает skip-индексы (пусто = corner-рендер не удался)."""
    if doc is None:
        return set()

    def _top(e):
        return _bbox_top_bottom(e[2], ph, pdf_native)[0]

    def _bot(e):
        return _bbox_top_bottom(e[2], ph, pdf_native)[1]

    left = sorted(left, key=_top)
    right = sorted(right, key=_top)

    # Границы левой колонки (для ширины ячеек и кропов рукописи)
    l_min = min(float(getattr(e[2], "l", 0)) for e in left)
    l_max = max(float(getattr(e[2], "r", 0)) for e in left)
    r_min = min(float(getattr(e[2], "l", 0)) for e in right)

    # ── Рукописные строки: блоки без кириллицы и без дата/№-паттернов ──
    hand = [e for e in left
            if not _CYR3_RE.search(e[3]) and not _CORNER_KEEP_RE.search(e[3])]
    hand_rows: list[tuple[float, float]] = []     # (y0, y1) в экранных pt
    for e in sorted(hand, key=_top):
        t, b = _top(e), _bot(e)
        if hand_rows and t <= hand_rows[-1][1] + 4:
            hand_rows[-1] = (hand_rows[-1][0], max(hand_rows[-1][1], b))
        else:
            hand_rows.append((t, b))

    def _in_hand_row(e) -> tuple[float, float] | None:
        mid = (_top(e) + _bot(e)) / 2
        for y0, y1 in hand_rows:
            if y0 - 2 <= mid <= y1 + 2:
                return (y0, y1)
        return None

    left_blocks: list[tuple[float, dict]] = []    # (top, block)
    used_rows: set[tuple[float, float]] = set()
    for e in left:
        row = _in_hand_row(e)
        if row is not None:
            if row not in used_rows:
                used_rows.add(row)
                crop = _crop_page_image(dl_doc, page_no, l_min - 4, row[0] - 3,
                                        l_max + 6, row[1] + 8, pw, ph)
                if crop is not None:
                    left_blocks.append((row[0], {
                        "image": crop,
                        "image_width_inch": (l_max - l_min + 10) / 72,
                        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
                    }))
                    log.info("corner: рукописная строка y=%.0f..%.0f — кроп",
                             row[0], row[1])
            continue
        text = e[3]
        alpha = [c for c in text if c.isalpha()]
        bold = bool(alpha) and all(c.isupper() for c in alpha) and len(text) <= 60
        left_blocks.append((_top(e), {
            "text": text, "font_pt": 8.0, "bold": bold, "italic": False,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "space_before": 3.0,     # межстрочный воздух, как в оригинале
        }))

    right_blocks: list[dict] = []
    title_skip: set[int] = set()      # индексы, которые НЕ скипаем (рендерятся ниже)
    for e in right:
        text = e[3]
        if _DOC_TITLE_RE.match(text.strip()) and len(text.strip()) <= 20:
            title_skip.add(e[0])
            log.info("corner: заголовок тела в правой колонке — рендер ниже: %r", text[:30])
            continue
        if not _CYR3_RE.search(text) and not _CORNER_KEEP_RE.search(text):
            log.info("corner: мусор в правой колонке пропущен: %r", text[:30])
            continue
        right_blocks.append({
            "text": text, "font_pt": 10.0, "bold": False, "italic": False,
            "alignment": right_align,
            "space_before": 3.0,
        })
    if not right_blocks or not left_blocks:
        return set()

    margin = MARGIN_INCH * 72
    mid = (l_max + r_min) / 2
    left_w_pt = max(mid - margin, 72.0)
    right_w_pt = max((pw - margin) - mid, 72.0)

    cells = [
        {
            "kind": "text",
            "blocks": [b for _, b in sorted(left_blocks, key=lambda x: x[0])],
            "width_pt": left_w_pt,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "image": pil_img,
            "image_width_inch": min(img_w_inch, left_w_pt / 72 - 0.1),
        },
        {
            "kind": "text",
            "blocks": right_blocks,
            "width_pt": right_w_pt,
            "alignment": right_align,
        },
    ]
    text_w_inch = (pw - 2 * margin) / 72
    add_header_row(doc, cells, text_w_inch)
    return {pic_idx, *(e[0] for e in left),
            *(e[0] for e in right if e[0] not in title_skip)}


def _render_first_page_letterhead(
    doc,
    dl_doc,
    all_items: list,
    page_sizes: dict[int, tuple[float, float]],
    pdf_native: bool,
    ocr_blocks: tuple | None = None,   # (list[TextBlock], img_h) из word_order
    logo_only: bool = False,           # True: не рендерить таблицу, вернуть image-данные
) -> tuple:
    """Render the top first-page image/text row without assuming logo position."""
    # Первая страница — минимальный ключ в page_sizes
    first_page = min(page_sizes.keys()) if page_sizes else 1
    pw, ph = page_sizes.get(first_page, (595.0, 842.0))
    log.info("letterhead: first_page=%s pdf_native=%s pw=%.0f ph=%.0f",
             first_page, pdf_native, pw, ph)
    _empty = (set(), [], None, 0.0, 0.0) if logo_only else (set(), [])
    if ph <= 0 or pw <= 0:
        return _empty

    pictures: list[tuple[int, object, object, float, float]] = []
    for idx, (item, _level) in enumerate(all_items):
        if _label_str(item) not in {"picture", "figure", "image"}:
            continue
        bbox, page_no = _item_bbox_page(item, first_page)
        if page_no != first_page or bbox is None:
            continue
        # Принимаем картинку если она в верхних 40% в ЛЮБОЙ системе координат.
        # pdf_native может быть определён неверно — проверяем оба варианта.
        t_val = float(getattr(bbox, "t", 0.0))
        b_val = float(getattr(bbox, "b", 0.0))
        top_if_native = (ph - max(t_val, b_val)) / ph   # если y=0 снизу
        top_if_screen = min(t_val, b_val) / ph          # если y=0 сверху
        if top_if_native > 0.40 and top_if_screen > 0.40:
            continue  # картинка не в верхней части страницы
        top, bottom = _bbox_top_bottom(bbox, ph, pdf_native)
        pictures.append((idx, item, bbox, top, bottom))

    if not pictures:
        log.info("letterhead: картинки не найдены — используется старый код add_sidebyside")
        return _empty

    pic_idx, pic_item, pic_bbox, pic_top, pic_bottom = min(pictures, key=lambda x: x[3])
    pic_h = max(pic_bottom - pic_top, 1.0)
    row_top = max(0.0, pic_top - max(pic_h * 0.35, ph * 0.025))
    default_bottom = min(
        ph,
        max(pic_bottom + max(pic_h * 1.5, ph * 0.10), ph * 0.24),
    )
    stop_tops: list[float] = []
    for item, _level in all_items:
        lbl = _label_str(item)
        if lbl not in {"paragraph", "text", "list_item", "caption", "title", "section_header"}:
            continue
        bbox, page_no = _item_bbox_page(item, first_page)
        if page_no != first_page or bbox is None:
            continue
        top, _bottom = _bbox_top_bottom(bbox, ph, pdf_native)
        if top <= pic_bottom:
            continue
        text = postprocess((getattr(item, "text", None) or "").strip())
        if text and _is_letterhead_stop(text):
            stop_tops.append(top)
    if stop_tops:
        # Граница = первый стоп-маркер (без буфера): сам маркер уже исключён
        # через _is_letterhead_stop в цикле сбора text_blocks.
        # Буфер ph*0.01 (8pt) отрезал строки вплотную к маркеру (тел./факс).
        row_bottom = min(stop_tops)
        row_bottom = max(row_bottom, pic_bottom)
    else:
        # Стоп-маркер не найден — безопасный максимум: верхняя четверть страницы
        row_bottom = min(default_bottom, ph * 0.25)

    text_indices: list[int] = []
    text_blocks: list[tuple[int, object, object, str]] = []
    for idx, (item, _level) in enumerate(all_items):
        if idx == pic_idx:
            continue
        lbl = _label_str(item)
        if lbl not in {"paragraph", "text", "list_item", "caption"}:
            continue
        bbox, page_no = _item_bbox_page(item, first_page)
        if page_no != first_page or bbox is None:
            continue
        top, bottom = _bbox_top_bottom(bbox, ph, pdf_native)
        if top > row_bottom or bottom < row_top:
            continue
        text = postprocess((getattr(item, "text", None) or "").strip())
        if not text or _is_letterhead_stop(text):
            continue
        text_indices.append(idx)
        text_blocks.append((idx, item, bbox, text))

    pil_img = _get_item_image(pic_item, dl_doc)
    if pil_img is None:
        return _empty
    # Грязевая полоса/штамп, принятые за логотип шапки, — шапку не строим
    if is_junk_image(pil_img):
        log.info("letterhead: картинка — мусор скана (пятно/штамп), пропуск")
        return _empty

    text_w_inch = (pw - 2 * MARGIN_INCH * 72) / 72
    pic_l = float(getattr(pic_bbox, "l", 0.0))
    pic_r = float(getattr(pic_bbox, "r", pic_l))
    pic_w_pt = max(pic_r - pic_l, 1.0)
    img_w_inch = min(max(pic_w_pt / 72, 0.5), text_w_inch)
    log.info("letterhead: pic_l=%.1f pic_r=%.1f pic_w_pt=%.1f img_w_inch=%.3f",
             pic_l, pic_r, pic_w_pt, img_w_inch)

    # ── Угловой штамп (письма ФНС): герб над центрированной левой колонкой,
    # правая колонка адресата с выключкой вправо. Зона ограничивается
    # ГЕОМЕТРИЧЕСКИ (первый блок через обе половины страницы), а не
    # стоп-словами: «Должник:» в правой колонке — не конец шапки.
    _zone_all: list[tuple[int, object, object, str, float, float]] = []
    _span_tops: list[float] = []
    for idx, (item, _level) in enumerate(all_items):
        if idx == pic_idx:
            continue
        lbl = _label_str(item)
        if lbl not in {"paragraph", "text", "list_item", "caption",
                       "section_header", "title"}:
            continue
        bbox, page_no = _item_bbox_page(item, first_page)
        if page_no != first_page or bbox is None:
            continue
        top, bottom = _bbox_top_bottom(bbox, ph, pdf_native)
        text = postprocess((getattr(item, "text", None) or "").strip())
        if not text:
            continue
        _x0 = float(getattr(bbox, "l", 0)); _x1 = float(getattr(bbox, "r", 0))
        if _x0 < pw * 0.45 and _x1 > pw * 0.55 and top > pic_top:
            _span_tops.append(top)          # блок через обе половины — конец зоны
            continue
        # Центрированный заголовок («Заявление…») тоже завершает зону, даже
        # если он короткий и не пересекает обе половины страницы.
        if (lbl in ("section_header", "title") and top > pic_top
                and abs((_x0 + _x1) / 2 - pw / 2) < pw * 0.08):
            _span_tops.append(top)
            continue
        _zone_all.append((idx, item, bbox, text, top, bottom))
    _zone_bottom = min(_span_tops) if _span_tops else ph * 0.45
    _zone = [(i, it, bb, tx) for (i, it, bb, tx, t, b) in _zone_all
             if b <= _zone_bottom + 2.0]
    _corner = _detect_corner_letterhead(_zone, pic_bbox, pw)
    if _corner is not None:
        _cl, _cr, _right_align = _corner
        _skip = _render_corner_letterhead(
            doc, dl_doc, first_page, pw, ph, pdf_native,
            pil_img, pic_w_pt, img_w_inch, pic_idx, _cl, _cr, _right_align)
        if _skip:
            # Прочие картинки зоны шапки (Docling делает picture и из
            # рукописной строки) — тоже в skip, иначе кроп задвоится.
            for _pi, _pit, _pbb, _pt, _pb in pictures:
                if _pi != pic_idx and _pb <= _zone_bottom + 2.0:
                    _skip.add(_pi)
            log.info("letterhead: угловой штамп — %d блоков в 2-колоночной шапке",
                     len(_skip) - 1)
            if logo_only:
                # corner-рендер уже построил полную шапку — native_header не нужен
                return _skip, [], None, 0.0, 0.0
            return _skip, []

    left_blocks = []
    right_blocks = []
    # Сортировка в экранных координатах (top=0 сверху) для правильного порядка.
    # bbox_mid_y для PDF-native (y=0 снизу) даёт обратный порядок — используем
    # _bbox_top_bottom, который нормализует к screen coords (top возрастает вниз).
    for _idx, _item, bbox, text in sorted(
        text_blocks,
        key=lambda x: (_bbox_top_bottom(x[2], ph, pdf_native)[0], bbox_x0(x[2]))
    ):
        block = {"text": text, "font_pt": 8.5, "bold": False, "italic": False}
        if float(getattr(bbox, "r", 0.0)) <= pic_l:
            left_blocks.append(block)
        elif bbox_x0(bbox) >= pic_r:
            right_blocks.append(block)
        elif bbox_x0(bbox) < pic_l:
            left_blocks.append(block)
        else:
            right_blocks.append(block)

    cells: list[dict] = []
    if left_blocks:
        left_w = max(pic_l - min(float(getattr(x[2], "l", pic_l)) for x in text_blocks), 72.0)
        cells.append({
            "kind": "text",
            "blocks": left_blocks,
            "width_pt": left_w,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        })

    cells.append({
        "kind": "image",
        "image": pil_img,
        "width_pt": pic_w_pt,
        "image_width_inch": img_w_inch,
        "alignment": detect_alignment(pic_bbox, pw),
    })

    if right_blocks and not logo_only:
        # Ширина правой колонки = от правого края логотипа до правого поля страницы.
        right_w = max(pw - pic_r, 72.0)
        cells.append({
            "kind": "text",
            "blocks": right_blocks,
            "width_pt": right_w,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        })

    # OCR-дополнение шапки 
    # Docling иногда пропускает строки в шапке (напр. тел./факс) из-за слияния
    # нескольких визуальных строк в один элемент. EasyOCR из word_order работает
    # на уровне визуальных строк и может восстановить пропущенные.
    if ocr_blocks is not None and right_blocks:
        ocr_blk_list, img_h = ocr_blocks
        img_w = img_h * pw / ph if ph > 0 else img_h
        x_thresh = (pic_r / pw) * 0.75 if pw > 0 else 0
        y_top_frac = pic_bottom / ph if ph > 0 else 0
        y_bot_frac = row_bottom / ph if ph > 0 else 1

        ocr_lines: list[tuple[float, str]] = []
        for blk in sorted(ocr_blk_list, key=lambda b: b.mid_y):
            y_frac = blk.mid_y / img_h if img_h > 0 else 0
            x_frac = blk.x0 / img_w if img_w > 0 else 0
            if not (y_top_frac <= y_frac <= y_bot_frac):
                continue
            if x_frac < x_thresh:
                continue
            blk_text = postprocess(blk.text)
            if blk_text:
                ocr_lines.append((y_frac, blk_text))
                log.info("letterhead OCR-line y=%.2f: %r", y_frac, blk_text[:70])

        if ocr_lines:
            # Заменяем Docling-блоки EasyOCR-строками: более полный список строк
            right_blocks.clear()
            for _, line_text in ocr_lines:
                right_blocks.append({
                    "text": line_text, "font_pt": 8.5,
                    "bold": False, "italic": False,
                })
            # Пересчитываем ячейку правой колонки (right_w не меняется)
            for cell in cells:
                if cell.get("kind") == "text" and cell.get("alignment") == WD_ALIGN_PARAGRAPH.LEFT:
                    cell["blocks"] = right_blocks
                    break

    log.info("letterhead: left_blocks=%d right_blocks=%d",
             len(left_blocks), len(right_blocks))
    for i, rb in enumerate(right_blocks):
        full_text = rb.get("text", "")
        log.info("  right_block[%d] (%d chars): %r", i, len(full_text), full_text)

    skip_set = {pic_idx, *text_indices}

    if logo_only:
        # Не рендерим таблицу — native_header построит единую 2-колоночную таблицу.
        # Skip только лого: текстовые блоки оставляем для native_header чтобы он
        # мог собрать все блоки шапки включая те что letterhead отфильтровал.
        log.info("letterhead logo_only: пропуск рендера, возвращаем только лого для native_header")
        return {pic_idx}, [], pil_img, pic_w_pt, img_w_inch

    add_header_row(doc, cells, text_w_inch)
    log.info("letterhead: skip_indices добавлены — pic_idx=%d, text_idx=%s",
             pic_idx, sorted(text_indices))
    # Возвращаем right_blocks=[] когда не logo_only — native_header не используется
    return skip_set, []


def build_docx(
    dl_doc,
    page_sizes: dict[int, tuple[float, float]],
    ocr_reader=None,
    use_word_order: bool = True,
    doc_type: str | None = None,
    page_infos=None,
    ink_bold: bool = False,
) -> "_Document":
    """
    Конвертирует DoclingDocument в python-docx Document.

    Параметры:
        dl_doc         — Docling DoclingDocument
        page_sizes     — {page_no: (width_pt, height_pt)}
        ocr_reader     — EasyOCR Reader для word_order (None = отключить)
        use_word_order — True: пересортировать блоки через word_order
        doc_type       — тип документа из doc_type.detect_doc_type() (None = авто)
        page_infos     — PageInfo из page_analyser.analyse_pages() (None = не используется)
    """
    from .doc_type import detect_doc_type, DOC_LEGAL
    from .legal_fixes import apply_legal_fixes

    all_items = list(dl_doc.iterate_items())
    log.info("build_docx: %d элементов из Docling, %d страниц",
             len(all_items), len(page_sizes))

    # Базовая сортировка в порядке чтения по Docling-координатам.
    # Построчная кластеризация (sort_reading_order) вместо бинов 40pt:
    # бин покрывал 2-3 строки, и блоки соседних строк сортировались по X —
    # свап пунктов списка, перемешивание слов на фрагментированных сканах.
    pdf_native = detect_pdf_native(all_items)
    log.info("build_docx: pdf_native=%s (система координат Docling)", pdf_native)
    all_items = sort_reading_order(all_items, pdf_native)

    page_medians, page_left_min = analyse_pages(all_items)
    for pn in sorted(page_sizes):
        log.info("  стр.%d: %.0f×%.0f pt, median_h=%.1f px, left_min=%.1f pt",
                 pn, page_sizes[pn][0], page_sizes[pn][1],
                 page_medians.get(pn, 0.0), page_left_min.get(pn, 0.0))

    # Модель страницы: колонки (page_analyser.PageInfo) + уровни отступов.
    # Колонки: блоку в правой колонке multicolumn-страницы выравнивание
    # считается относительно СВОЕЙ колонки, а не всей страницы.
    # Уровни отступов: сырые x0 «дрожат» на 1-4pt (шум OCR) — прищёлкиваем
    # каждый блок к ближайшему уровню документа, вертикали текста ровные.
    page_infos = analyse_page_infos(all_items)
    for pn, _pi in sorted(page_infos.items()):
        if _pi.is_multicolumn:
            log.info("  стр.%d: MULTICOLUMN %s", pn,
                     [(round(c.x0), round(c.x1)) for c in _pi.columns])
    indent_levels = build_indent_levels(all_items, page_left_min)
    log.info("Уровни отступов документа (pt от левого поля): %s",
             [round(lv, 1) for lv in indent_levels])

    # word_order: строим карту блоков {page_no: (blocks, img_h_px)},
    # затем пересортировываем all_items по позициям этих блоков.
    # Текст Docling НЕ заменяем — только меняем порядок элементов.
    word_blocks_map: dict[int, tuple[list[TextBlock], float]] = {}
    if use_word_order and ocr_reader is not None:
        for page_no in page_sizes:
            try:
                blocks, img_h = blocks_for_page(dl_doc, page_no, ocr_reader)
                if blocks:
                    word_blocks_map[page_no] = (blocks, img_h)
            except Exception as exc:
                log.warning("word_order page %d failed: %s", page_no, exc)

        if word_blocks_map:
            all_items = _reorder_by_word_order(all_items, word_blocks_map, page_sizes)
            log.info("word_order: переупорядочено на %d стр.", len(word_blocks_map))

    # Определяем тип документа если не передан извне
    if doc_type is None:
        doc_type = detect_doc_type(all_items)
        log.info("build_docx: doc_type=%s (авто)", doc_type)
    else:
        log.info("build_docx: doc_type=%s (передан)", doc_type)

    # Пост-обработка порядка блоков: только для судебных документов
    _continuation_ids: set = set()
    if doc_type == DOC_LEGAL:
        all_items, _continuation_ids = apply_legal_fixes(all_items)
    else:
        # Для других типов — базовые исправления из старого кода
        all_items, _continuation_ids = _fix_reading_order(all_items)

    # Универсальная склейка фрагментов: для ЛЮБОГО типа документа.
    # Детектирует страницы с экстремальной фрагментацией (>50% коротких блоков)
    # и помечает их как continuation — рендер склеит соседние блоки в абзацы.
    _continuation_ids = _join_fragments(all_items, _continuation_ids)

    # iterate_items() ПРОПУСКАЕТ page_footer, но строку подписи OCR часто метит
    # именно так ("Представитель ... по доверенности" + ФИО). Возвращаем эти блоки
    # в обработку (рендерятся ниже как группа подписи).
    for _ft in getattr(dl_doc, "texts", []):
        if "page_footer" not in str(getattr(_ft, "label", "")).lower():
            continue
        _ftt = (getattr(_ft, "text", None) or "").strip()
        if _REPR_RE.search(_ftt) or _INITIALS_RE.match(_ftt):
            all_items.append((_ft, 0))
            log.info("signature footer возвращён в обработку: %r", _ftt[:50])

    doc               = init_document()
    last_content_page = -1
    current_page      = -1
    prev_midY: dict[int, float] = {}
    prev_h:    dict[int, float] = {}
    _gosposhlina_bold = False

    """
    Трекинг продолжений (orphan-блоки Docling) 
    Docling иногда разрывает абзац: конец одного блока + начало следующего
    теряют связь. Признак продолжения: блок начинается с союза/частицы
    строчными буквами И предыдущий text-абзац не завершён.
    Список консервативный — только слова, которые НЕ начинают новых предложений.
    """
    _CONT_RE = re.compile(
        r'^(как\b|а\s+также\b|в\s+том\s+числе\b|при\s+этом\b|при\s+условии\b'
        r'|которы[еийх]\b|котор\w{2,5}\b'
        r'|обеспеченн\w+|предусмотренн\w+)',
        re.IGNORECASE,
    )
    # Только настоящие концы предложений: . ! ?
    _SENT_END_RE = re.compile(r'[.!?]\s*$')

    _last_body_para: "Paragraph | None" = None  # последний Word-параграф тела документа
    _last_body_text:      str      = ""    # его текст (для проверки завершённости)
    _seen_text_pages:     set      = set() # страницы, на которых уже был text-блок
    _image_pages_rendered: set[int] = set() # страницы уже вставленные как картинки
    skip_indices: set[int] = set()

    _body_len_at_break = 0   # размер body на момент последнего разрыва

    def _page_break(target: int, force: bool = False) -> None:
        nonlocal last_content_page, _body_len_at_break
        # Разрыв страницы на границе исходной страницы скана (target > текущей):
        # число страниц и разбивка DOCX совпадают с оригиналом.
        # ЗАЩИТА ОТ ПУСТЫХ ЛИСТОВ: если после предыдущего разрыва в документ
        # не добавилось ни одного элемента (пустая/пропущенная страница скана,
        # двойной вызов на таблице+границе), второй разрыв не вставляем.
        if last_content_page >= 0 and target > last_content_page:
            cur_len = len(doc.element.body)
            if cur_len > _body_len_at_break:
                if _last_body_para is not None:
                    from docx.enum.text import WD_BREAK
                    _last_body_para.add_run().add_break(WD_BREAK.PAGE)
                else:
                    pb = doc.add_page_break()
                    pb.paragraph_format.space_before  = Pt(0)
                    pb.paragraph_format.space_after   = Pt(0)
                    pb.paragraph_format.widow_control = False
                _body_len_at_break = len(doc.element.body)
            else:
                log.info("[стр%d] пустой лист предотвращён (нет контента "
                         "после предыдущего разрыва)", target)
        last_content_page = target

    # Передаём OCR-блоки первой страницы для дополнения шапки
    first_page_no = min(page_sizes.keys()) if page_sizes else 1

    # Предварительно определяем: есть ли двухколоночная шапка на первой странице.
    # Если есть — letterhead рендерит только лого, а native_header строит полную
    # двухколоночную таблицу (включая блоки которые иначе ушли бы в letterhead).
    _native_preview = _render_native_two_column_header(
        None, all_items, page_sizes, pdf_native,  # doc=None — только анализ
        page_medians, page_left_min,
        already_skipped=set(),
    )
    _has_two_col_header = bool(_native_preview)

    _lh_result = _render_first_page_letterhead(
        doc, dl_doc, all_items, page_sizes, pdf_native,
        ocr_blocks=word_blocks_map.get(first_page_no),
        logo_only=_has_two_col_header,
    )
    if len(_lh_result) == 5:
        # logo_only=True - (skip_set, right_blocks, pil_img, pic_w_pt, img_w_inch)
        letterhead_indices, _lh_right_blocks, _lh_image, _lh_pic_w_pt, _lh_img_w_inch = _lh_result
    else:
        # logo_only=False - (skip_set, [])
        letterhead_indices, _lh_right_blocks = _lh_result
        _lh_image = _lh_img_w_inch = _lh_pic_w_pt = None

    if letterhead_indices:
        skip_indices.update(letterhead_indices)
        last_content_page = first_page_no
        log.info("letterhead: пропускаем %d элементов шапки", len(letterhead_indices))

    # Двухколоночная шапка с лого — letterhead передаёт image данные,
    # native_header строит единую 3-колоночную таблицу: [лого | label | content].
    _prepend = _lh_right_blocks if _has_two_col_header else None
    _logo_data = (_lh_image, _lh_pic_w_pt, _lh_img_w_inch) if _has_two_col_header else None
    native_indices = _render_native_two_column_header(
        doc, all_items, page_sizes, pdf_native,
        page_medians, page_left_min,
        already_skipped=skip_indices,
        prepend_right_blocks=_prepend,
        logo_data=_logo_data,
    )
    if native_indices:
        skip_indices.update(native_indices)
        last_content_page = first_page_no
        log.info("native_header: пропускаем %d элементов шапки", len(native_indices))

    # Слияние пословно-разорванной шапки (лесенка ФНС): склеивает фрагменты строки
    # в один блок ДО рендера. Выполняется после corner/native-скипов — их блоки
    # не трогает. Мутирует item.text/bbox.r и добавляет хвосты в skip_indices.
    _merge_header_lines(all_items, page_sizes, pdf_native, skip_indices)

    # Детальный дамп стр.1 для диагностики рендера
    _pg1 = min(page_sizes.keys()) if page_sizes else 1
    _pw1, _ph1 = page_sizes.get(_pg1, (595.0, 842.0))
    log.info("=== ДАМП СТРАНИЦЫ %d  pw=%.0f ph=%.0f pdf_native=%s ===",
             _pg1, _pw1, _ph1, pdf_native)
    for _di, (_ditem, _dlvl) in enumerate(all_items):
        _dprov = getattr(_ditem, "prov", None) or []
        if not _dprov:
            continue
        _dpg = int(getattr(_dprov[0], "page_no", -1))
        if _dpg != _pg1:
            continue
        _dbbox  = getattr(_dprov[0], "bbox", None)
        _dlbl   = _label_str(_ditem)
        _draw   = (getattr(_ditem, "text", None) or "")
        _dflags: list[str] = []
        if _di in skip_indices:
            _dflags.append("SKIP")
        if id(_ditem) in _continuation_ids:
            _dflags.append("CONT")
        if _dbbox is not None:
            _dl = float(getattr(_dbbox, "l", 0));  _dt = float(getattr(_dbbox, "t", 0))
            _dr = float(getattr(_dbbox, "r", 0));  _db = float(getattr(_dbbox, "b", 0))
            _dh = abs(_db - _dt)
            # Нормированная вертикальная позиция (0=верх, 1=низ) в обеих системах
            _top_screen  = min(_dt, _db) / _ph1
            _top_native  = 1.0 - max(_dt, _db) / _ph1
            log.info(
                "  [%2d] %-16s lvl=%d  l=%5.1f t=%5.1f r=%5.1f b=%5.1f  h=%5.1f"
                "  top_sc=%.3f top_nat=%.3f  %s | %r",
                _di, _dlbl, _dlvl,
                _dl, _dt, _dr, _db, _dh,
                _top_screen, _top_native,
                ",".join(_dflags) if _dflags else "-",
                _draw[:80],
            )
        else:
            log.info("  [%2d] %-16s lvl=%d  (no bbox)  %s | %r",
                     _di, _dlbl, _dlvl,
                     ",".join(_dflags) if _dflags else "-", _draw[:80])
    log.info("=== КОНЕЦ ДАМПА СТРАНИЦЫ %d ===", _pg1)

    _align_names = {
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        WD_ALIGN_PARAGRAPH.CENTER:  "CENTER",
        WD_ALIGN_PARAGRAPH.LEFT:    "LEFT",
        WD_ALIGN_PARAGRAPH.RIGHT:   "RIGHT",
    }

    _es_stamp_rendered = False

    """
    Адаптивный межстрочный интервал 
    Базовый интервал тела LINE_SPACING (1.5) заполняет страницу, но на ПЛОТНЫХ
    страницах выталкивает текст дальше. Оцениваем объём текста каждой исходной
    страницы и подбираем интервал: малозаполненные - 1.5, плотные - меньше
    (до 1.0), чтобы контент не вылезал на следующую страницу.
    """
    def _estimate_page_metrics():
        single_lh = BODY_PT * 1.15              # высота одинарной строки, pt
        margin_pt = MARGIN_INCH * 72
        lines: dict[int, float] = {}
        line_hs: dict[int, list[float]] = {}    # высоты СТРОК по странице (для кегля)
        _cur = 1
        for _it, _ in all_items:
            if _label_str(_it) not in ("text", "paragraph", "list_item"):
                continue
            _bb, _pg = _item_bbox_page(_it, _cur); _cur = _pg
            _t = postprocess((getattr(_it, "text", None) or "").strip())
            if not _t:
                continue
            _pw, _ph = page_sizes.get(_pg, (595.0, 842.0))
            _cpl = max(40, int((_pw - 2 * margin_pt) / (BODY_PT * 0.50)))
            _nl = max(1, -(-len(_t) // _cpl))   # оценка строк блока (ceil)
            lines[_pg] = lines.get(_pg, 0) + _nl
            _h = bbox_h(_bb) if _bb is not None else 0.0
            if _h > 2:
                line_hs.setdefault(_pg, []).append(_h / _nl)  # высота ОДНОЙ строки
        spacing: dict[int, float] = {}
        for _pg, _ln in lines.items():
            _pw, _ph = page_sizes.get(_pg, (595.0, 842.0))
            avail = (_ph - 2 * margin_pt) * 0.85    # запас на интервалы/красную строку/шапку
            base  = _ln * single_lh
            spacing[_pg] = (max(1.0, min(LINE_SPACING, avail / base))
                            if base > 0 else LINE_SPACING)
        line_median: dict[int, float] = {}      # ПОСТРАНИЧНО — для адаптивного интервала
        for _pg, _hs in line_hs.items():
            _s = sorted(_hs); line_median[_pg] = _s[len(_s) // 2]
            
        """
        КЕГЛЬ — ОДИН на весь документ (медиана высот строк ВСЕХ страниц), а не по
        странице. Постраничный кегль прыгал между страницами при нестабильной
        сегментации OCR (RapidOCR: на стр.1 мелкая шапка тянет медиану вниз - 10pt,
        тело - 12pt — визуально). Документная медиана устойчива:
        тело доминирует по числу строк, мелкая шапка/выбросы не перевешивают.
        font ≈ pitch/1.15 (одинарный TNR), ограничение [9.5 .. 12].
        """
        _all_hs = sorted(_h for _hs in line_hs.values() for _h in _hs)
        if _all_hs:
            _doc_med = _all_hs[len(_all_hs) // 2]
            _doc_pt  = max(9.5, min(12.0, round(_doc_med / 1.15 * 2) / 2))
        else:
            _doc_pt = BODY_PT
        body_pt = {_pg: _doc_pt for _pg in page_sizes}
        return spacing, line_median, body_pt

    page_line_spacing, page_line_median, page_body_pt = _estimate_page_metrics()
    log.info("Адаптивный интервал по страницам: %s",
             {p: round(s, 2) for p, s in sorted(page_line_spacing.items())})
    log.info("Адаптивный кегль тела по страницам: %s",
             {p: v for p, v in sorted(page_body_pt.items())})

    """
    Жирность по изображению (сканы)
    На сканах Docling formatting.bold пуст - оцениваем жирность по насыщенности
    штриха. Для body-блоков считаем stroke_density (ink.py), берём МЕДИАНУ тела
    по странице; в основном цикле блок жирный, если заметно темнее медианы.
    Кэшируем по id(item), чтобы не пересчитывать кроп дважды.
    """
    _INK_BOLD_K  = 1.30   # порог: stroke_density > median * K - жирный
    _INK_MIN_LEN = 15     # короткие блоки статистически шумны — пропускаем

    # "Скан" определяем по ОТСУТСТВИЮ Docling formatting у всех элементов
    # (OCR не даёт стиль). pdf_native — это лишь система координат (y снизу), а
    # НЕ признак текстового слоя: сканы тоже бывают с native-координатами.
    def _doc_has_formatting() -> bool:
        for _it, _ in all_items:
            _f = getattr(_it, "formatting", None)
            if _f is not None and (getattr(_f, "bold", False)
                                   or getattr(_f, "italic", False)):
                return True
        return False

    """
    ink-жирность ОПЦИОНАЛЬНА (по умолчанию выкл). Эксперименты показали: толщина
    штриха на скане НЕ разделяет жирный/обычный даже при 300 DPI — её определяет
    состав символов (заглавные имена толще обычного текста), а не насыщенность
    шрифта. Включение даёт ложную жирность, поэтому по умолчанию полагаемся на
    структурную жирность (заголовки/метки). Флаг --ink-bold — для высокого DPI/опытов.
    """
    _use_ink = ink_bold and not _doc_has_formatting()
    log.info("Жирность по изображению: %s",
             "ВКЛ (опц.)" if _use_ink else
             "выкл (структурная жирность по заголовкам/меткам)")

    def _estimate_ink_medians():
        if not _use_ink:
            return {}, {}          # нативный PDF — доверяем Docling formatting
        try:
            from .ink import block_ink_stats
        except Exception as _e:
            log.debug("ink: модуль недоступен: %s", _e)
            return {}, {}
        ratios: dict[int, list[float]] = {}
        cache: dict[int, float] = {}
        _cur = 1
        for _it, _ in all_items:
            if _label_str(_it) not in ("text", "paragraph", "list_item"):
                continue
            _t = (getattr(_it, "text", None) or "").strip()
            if len(_t) < _INK_MIN_LEN:
                continue
            _bb, _pg = _item_bbox_page(_it, _cur); _cur = _pg
            try:
                _img = _get_item_image(_it, dl_doc)
            except Exception:
                _img = None
            _st = block_ink_stats(_img) if _img is not None else None
            if _st is None:
                continue
            _run = _st["mean_run"]             # толщина штриха (среднее коротких пробегов)
            cache[id(_it)] = _run
            ratios.setdefault(_pg, []).append(_run)
        medians = {p: statistics.median(rs) for p, rs in ratios.items() if rs}
        if medians:
            log.info("Жирность по изображению: медианы mean_run(px): %s",
                     {p: round(m, 2) for p, m in sorted(medians.items())})
        return medians, cache

    page_ink_median, ink_cache = _estimate_ink_medians()

    """
    Группы строки подписи
    OCR метит подпись как page_footer + отдельная картинка, причём картинка в
    порядке чтения может идти РАНЬШЕ текста - собираем все компоненты подписи
    (должность "Представитель...по доверенности" + ФИО-инициалы + картинка) по
    странице ЗАРАНЕЕ и рендерим одной строкой при встрече первого компонента.
    """
    def _find_signature_groups() -> dict[int, dict]:
        by_page: dict[int, list[int]] = {}
        _cur = 1
        for _i, (_it, _) in enumerate(all_items):
            _, _pg = _item_bbox_page(_it, _cur); _cur = _pg
            by_page.setdefault(_pg, []).append(_i)
        groups: dict[int, dict] = {}
        for _pg, idxs in by_page.items():
            repr_i = init_i = pic_i = None
            for _i in idxs:
                _it = all_items[_i][0]
                _lbl = _label_str(_it)
                _t = postprocess((getattr(_it, "text", None) or "").strip())
                if repr_i is None and _REPR_RE.search(_t):
                    repr_i = _i
                elif init_i is None and _INITIALS_RE.match(_t):
                    init_i = _i
                if pic_i is None and _lbl in ("picture", "figure", "image"):
                    pic_i = _i
            # Группа подписи только если есть «должность» И (инициалы ИЛИ картинка).
            if repr_i is not None and (init_i is not None or pic_i is not None):
                members = {x for x in (repr_i, init_i, pic_i) if x is not None}
                groups[min(members)] = {"members": members, "repr": repr_i,
                                        "init": init_i, "pic": pic_i, "page": _pg}
        return groups

    _sig_groups = _find_signature_groups()
    if _sig_groups:
        log.info("Группы подписи: %s", {a: g["members"] for a, g in _sig_groups.items()})


    """
    ПРЕД-ПРОХОД: координатное спаривание "значение РАНЬШЕ метки"
    OCR иногда инвертирует порядок чтения: значение идёт в
    списке РАНЬШЕ своей метки. Тогда inline coplanar-поиск
    оставляет метку пустой, а значение рендерится одиночным
    абзацем. Здесь находим такие пары координатно
    и: метим значение в skip_indices (не рендерить отдельно) +
    запоминаем пару в _coplanar_back, чтобы метка взяла значение из карты.
    """
    # Метка перед «:» должна ВЫГЛЯДЕТЬ меткой: >=4 букв и заглавная первая.
    # Иначе обрывок переноса («…, из / них: налог …») считается меткой и
    # утаскивает своё «значение» из середины предложения.
    def _labelish(t: str) -> bool:
        head = t.rstrip().rstrip(":").strip()
        alpha = [c for c in head if c.isalpha()]
        return len(alpha) >= 4 and bool(alpha) and alpha[0].isupper()

    _coplanar_back: dict[int, int] = {}
    for _li, (_litem, _) in enumerate(all_items):
        if _li in skip_indices:
            continue
        if _label_str(_litem) not in ("text", "paragraph"):
            continue
        _ltext = (getattr(_litem, "text", None) or "").rstrip()
        if not _ltext.endswith(":") or not _labelish(_ltext):
            continue
        _lprov = getattr(_litem, "prov", None) or []
        if not _lprov:
            continue
        _lpg = int(getattr(_lprov[0], "page_no", -1))
        _lbb = getattr(_lprov[0], "bbox", None)
        if _lbb is None:
            continue
        _lpw, _lph = page_sizes.get(_lpg, (595.0, 842.0))
        _ll = float(getattr(_lbb, "l", 0)); _lr = float(getattr(_lbb, "r", 0))
        # До 0.75/0.85 ширины: метки правовыровненных строк («Адрес:» в шапке
        # заявления, x0≈0.55-0.65pw) тоже пары — решение «склеить в строку или
        # таблица» принимает зазор в рендере.
        if not (_ll < _lpw * 0.75 and _lr < _lpw * 0.85):
            continue
        _ltop = max(float(getattr(_lbb, "t", 0)), float(getattr(_lbb, "b", 0)))
        # значение ищем СРЕДИ ПРЕДШЕСТВУЮЩИХ блоков (тех, что inline-поиск не видит)
        for _vi in range(max(0, _li - 12), _li):
            if _vi in skip_indices or _vi in _coplanar_back.values():
                continue
            _vitem, _ = all_items[_vi]
            if _label_str(_vitem) not in ("text", "paragraph"):
                continue
            _vtext = (getattr(_vitem, "text", None) or "").strip()
            if not _vtext or _vtext.rstrip().endswith(":"):
                continue
            _vprov = getattr(_vitem, "prov", None) or []
            if not _vprov or int(getattr(_vprov[0], "page_no", -1)) != _lpg:
                continue
            _vbb = getattr(_vprov[0], "bbox", None)
            if _vbb is None:
                continue
            _vx0 = float(getattr(_vbb, "l", 0))
            _vtop = max(float(getattr(_vbb, "t", 0)), float(getattr(_vbb, "b", 0)))
            if _vx0 > _lr and abs(_vtop - _ltop) <= 8.0:
                _coplanar_back[_li] = _vi
                skip_indices.add(_vi)
                log.info("coplanar-back: метка[%d] %r ← значение[%d] %r",
                         _li, _ltext[:30], _vi, _vtext[:30])
                break

    for idx, (item, _level) in enumerate(all_items):
        if idx in skip_indices:
            raw = (getattr(item, "text", None) or "").strip()[:50]
            log.debug("[%d] пропуск (skip_indices): lbl=%s  %r",
                      idx, _label_str(item), raw)
            continue

        lbl = _label_str(item)
        # Исключение: строку подписи OCR часто метит как page_footer
        # ("Представитель ... по доверенности" + ФИО). Её НЕ пропускаем — она
        # обрабатывается ниже как строка подписи.
        _is_sign_footer = (lbl == "page_footer"
                           and _REPR_RE.search(getattr(item, "text", "") or ""))
        if not lbl or (lbl in SKIP_LABELS and not _is_sign_footer):
            log.debug("[%d] пропуск (SKIP_LABELS): lbl=%r", idx, lbl)
            continue

        bbox, page_no = _item_bbox_page(item, max(current_page, 1))
        current_page  = page_no

        pw, ph    = page_sizes.get(page_no, (595.0, 842.0))

        # Строка подписи (группа: должность + ФИО + картинка)
        # Рендерим при встрече ПЕРВОГО компонента группы (порядок чтения может
        # ставить картинку раньше текста), остальные компоненты пропускаем.
        if idx in _sig_groups:
            _g = _sig_groups[idx]
            _repr_txt = postprocess(
                (getattr(all_items[_g["repr"]][0], "text", None) or "").strip())
            _r_txt = ""
            if _g["init"] is not None:
                _r_txt = postprocess(
                    (getattr(all_items[_g["init"]][0], "text", None) or "").strip())
            _s_img = None
            if _g["pic"] is not None:
                try:
                    _s_img = _get_item_image(all_items[_g["pic"]][0], dl_doc)
                except Exception:
                    pass
            _sig_tw = (pw - 2 * MARGIN_INCH * 72) / 72
            add_signature_row(doc, _repr_txt, _r_txt, _sig_tw,
                              space_before=8.0, sig_image=_s_img)
            skip_indices.update(_g["members"])
            _last_body_para = None
            log.info("[стр%d] signature group: должность=%r инициалы=%r pic=%s",
                     page_no, _repr_txt[:40], _r_txt, _s_img is not None)
            continue

        """
        Штамп электронной подписи (ЭП)
        Маркер «Электронная подпись действительна» открывает штамп ЭП. Собираем
        его короткие строки (Данные ЭП / Удостоверяющий центр / Дата / Кому
        выдана) и рендерим в обрамлённой рамке. Строку-метку, оканчивающуюся
        на "":"", склеиваем со следующей (значением).
        """
        if (not _es_stamp_rendered
                and lbl in ("text", "paragraph")
                and _ES_STAMP_MARKER_RE.search(getattr(item, "text", "") or "")):
            _stamp_raw: list[str] = []
            for _sj in range(idx, len(all_items)):
                if _sj in skip_indices:
                    continue
                _sit, _ = all_items[_sj]
                if _label_str(_sit) not in ("text", "paragraph", "list_item"):
                    continue
                _sbbox, _sp = _item_bbox_page(_sit, page_no)
                if _sp != page_no:
                    break
                _stext = postprocess((getattr(_sit, "text", None) or "").strip())
                if not _stext:
                    continue
                if len(_stext) > 130:    # длинный абзац — не часть штампа
                    break
                _stamp_raw.append(_stext)
                skip_indices.add(_sj)
            # Склеиваем «метка:» + следующая строка
            _stamp_lines: list[str] = []
            for _ln in _stamp_raw:
                if _stamp_lines and _stamp_lines[-1].rstrip().endswith(":"):
                    _stamp_lines[-1] = _stamp_lines[-1].rstrip() + " " + _ln
                else:
                    _stamp_lines.append(_ln)
            _es_tw = (pw - 2 * MARGIN_INCH * 72) / 72
            _page_break(page_no)
            add_es_stamp(doc, _stamp_lines, _es_tw)
            _es_stamp_rendered = True
            log.info("[стр%d] ЭП-штамп: %d строк в рамке", page_no, len(_stamp_lines))
            continue
        median_h  = page_medians.get(page_no, 0.0)
        item_h    = bbox_h(bbox) if bbox is not None else 0.0

        # Страница-картинка
        # Вставляем страницу целиком как изображение через picture-элемент.
        # Остальные элементы этой страницы пропускаем.
        if (median_h == 0.0 and page_no not in _image_pages_rendered
                and lbl in ("picture", "figure", "image")):
            _image_pages_rendered.add(page_no)
            try:
                pil_img = _get_item_image(item, dl_doc)
                if pil_img is not None:
                    from io import BytesIO as _BytesIO
                    _buf = _BytesIO()
                    pil_img.save(_buf, format="PNG")
                    _buf.seek(0)
                    text_w_inch = (pw - 2 * MARGIN_INCH * 72) / 72
                    _page_break(page_no, force=True)   # страница-картинка — на свой лист
                    _para = doc.add_paragraph()
                    _para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _para.paragraph_format.space_before = Pt(0)
                    _para.paragraph_format.space_after  = Pt(0)
                    _iw, _ih = pil_img.size
                    _aspect  = _ih / _iw if _iw > 0 else 1.0
                    _tw      = min(text_w_inch, (ph - 2 * MARGIN_INCH * 72) / 72 / _aspect)
                    _para.add_run().add_picture(_buf, width=Inches(max(_tw, 0.5)))
                    _last_body_para = None   # разрыв след. страницы — после картинки
                    log.info("[стр%d] PAGE_IMAGE: вставлена страница как картинка %dx%d px",
                             page_no, pil_img.width, pil_img.height)
            except Exception as _exc:
                log.debug("[стр%d] PAGE_IMAGE: ошибка вставки: %s", page_no, _exc)
            continue

        font_pt = LABEL_PT.get(lbl, BODY_PT)
        # Тело: АБСОЛЮТНЫЙ кегль по измеренному шагу строк оригинала,
        # чтобы число страниц совпадало со сканом.
        _base_pt = page_body_pt.get(page_no, BODY_PT)
        if lbl in ("paragraph", "text", "list_item"):
            font_pt = _base_pt
        # section_header в правой части первой страницы
        # — используем размер тела (11pt), не 13pt.
        # ЗАЯВЛЕНИЕ (x0=270.7 < pw*0.48=285.6) и подзаголовки остаются 13pt.
        if lbl == "section_header" and page_no == first_page_no and bbox is not None:
            _sh_x0 = float(getattr(bbox, "l", 0))
            if _sh_x0 > pw * 0.48:
                font_pt = BODY_PT

        if lbl in ("paragraph", "text") and item_h > 2:
            # Кегль оцениваем по высоте ОДНОЙ СТРОКИ (item_h / число строк), а не
            # по высоте блока: иначе однострочные нормальные блоки ошибочно сжимались -
            # "текст разного размера". Сравниваем с медианной высотой строки страницы.
            raw_x0 = float(getattr(bbox, "l", 0)) if bbox is not None else 0.0
            is_potential_label = raw_x0 <= pw * LABEL_MARGIN_THRESHOLD
            _pre_align = (detect_alignment(bbox, pw)
                          if bbox is not None and pw > 0
                          else WD_ALIGN_PARAGRAPH.JUSTIFY)
            _lm = page_line_median.get(page_no, 0.0)
            if _lm > 0 and not is_potential_label \
                    and _pre_align != WD_ALIGN_PARAGRAPH.CENTER \
                    and _pre_align != WD_ALIGN_PARAGRAPH.RIGHT:
                _txt_est = (getattr(item, "text", None) or "").strip()
                _cpl_p = max(40, int((pw - 2 * MARGIN_INCH * 72) / (BODY_PT * 0.50)))
                _nl    = max(1, -(-len(_txt_est) // _cpl_p))
                _line_ratio = (item_h / _nl) / _lm
                """
                масштабируем кегль по высоте строки относительно медианы тела:
                < 0.80 - уменьшаем (сноска/мелкий шрифт); > 1.3 - увеличиваем
                (подзаголовок/акцент в теле); около 1 — оставляем базовый кегль.
                Базой служит адаптивный _base_pt (кегль страницы), не фикс 11pt.
                """
                if _line_ratio < 0.80:
                    font_pt = max(round(_base_pt * _line_ratio * 2) / 2, 8.0)
                elif _line_ratio > 1.30:
                    font_pt = min(round(_base_pt * _line_ratio * 2) / 2, 16.0)
                    log.debug("[стр%d] крупная строка ratio=%.2f → font=%.1f %r",
                              page_no, _line_ratio, font_pt, _txt_est[:40])

        # Таблицы
        if lbl == "table":
            try:
                data = getattr(item, "data", None)
                if data is None:
                    log.debug("[стр%d] table: нет data — пропуск", page_no)
                    continue
                _cells = list(getattr(data, "table_cells", []) or [])
                nr = getattr(data, "num_rows", len(getattr(data, "grid", [])))
                nc = getattr(data, "num_cols",
                             max((len(r) for r in getattr(data, "grid", [[]])), default=0))
                log.info("[стр%d] table: %d строк × %d столбцов, ячеек %d",
                         page_no, nr, nc, len(_cells))
                # «Таблица» 1x1 — ложная сработка детектора на обычном тексте:
                # рендерим содержимое обычным абзацем, не рамкой.
                if nr <= 1 and nc <= 1:
                    _t11 = postprocess(" ".join(
                        (getattr(c, "text", "") or "").strip() for c in _cells).strip())
                    if _t11:
                        _page_break(page_no)
                        _p11 = doc.add_paragraph()
                        _p11.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        _p11.paragraph_format.space_before = Pt(0)
                        _p11.paragraph_format.space_after  = Pt(0)
                        _r11 = _p11.add_run(_t11)
                        _r11.font.name = FONT_NAME
                        _r11.font.size = Pt(BODY_PT)
                        _last_body_para = _p11
                        _last_body_text = _t11
                        log.info("[стр%d] table 1x1 → абзац: %r", page_no, _t11[:60])
                    continue
                _page_break(page_no, force=True)   # таблица на новой стр. скана — сохраняем разрыв
                _tw_in = (pw - 2 * MARGIN_INCH * 72) / 72
                # Ячейки TableFormer в приоритете: сохраняют span'ы и ширины.
                # grid дублирует текст объединённых ячеек в каждую позицию.
                if _cells and nr > 0:
                    add_table_from_cells(doc, item, text_w_inch=_tw_in)
                elif hasattr(data, "grid") and data.grid:
                    add_table_from_grid(doc, data.grid)
                # Иначе разрыв СЛЕДУЮЩЕЙ страницы уйдёт в абзац ПЕРЕД таблицей
                # (два разрыва в одном абзаце = пустой лист, таблицы после них)
                _last_body_para = None
            except Exception as exc:
                log.warning("[стр%d] table: пропущена — %s", page_no, exc)
            continue

        # Картинки / логотипы
        if lbl in ("picture", "figure", "image"):
            try:
                pil_img = _get_item_image(item, dl_doc)
                if pil_img is None:
                    log.debug("[стр%d] %s: изображение недоступно — пропуск", page_no, lbl)
                    continue
                # Мусорные «картинки»: layout-модель принимает грязевые полосы
                # и канцелярские штампы за рисунки — в DOCX их не вставляем.
                if is_junk_image(pil_img):
                    log.info("[стр%d] %s: мусорная картинка (пятно/штамп) — пропуск",
                             page_no, lbl)
                    continue
                log.info("[стр%d] %s: %dx%d px", page_no, lbl, pil_img.width, pil_img.height)
                text_w_inch = (pw - 2 * MARGIN_INCH * 72) / 72
                text_h_inch = (ph - 2 * MARGIN_INCH * 72) / 72

                if bbox is not None:
                    bbox_w_pt     = float(getattr(bbox, "r", 0)) - float(getattr(bbox, "l", 0))
                    target_w_inch = min(bbox_w_pt / 72, text_w_inch) if bbox_w_pt > 10 else text_w_inch
                    # На первой странице — выравниваем по bbox; на остальных — центр
                    img_align = (detect_alignment(bbox, pw)
                                 if page_no == first_page_no
                                 else WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    target_w_inch = text_w_inch
                    img_align     = WD_ALIGN_PARAGRAPH.CENTER

                pil_w, pil_h_px = pil_img.size
                aspect = pil_h_px / pil_w if pil_w > 0 else 1.0
                if target_w_inch * aspect > text_h_inch:
                    target_w_inch = text_h_inch / aspect
                target_w_inch = max(target_w_inch, 0.5)

                pic_right   = float(getattr(bbox, "r", 0)) if bbox is not None else pw
                right_blocks: list[dict] = []
                right_skip:   list[int]  = []
                _stop = {"table", "picture", "figure", "image"}

                for j in range(idx + 1, min(idx + 15, len(all_items))):
                    nxt_item, _ = all_items[j]
                    nxt_lbl     = _label_str(nxt_item)
                    if not nxt_lbl or nxt_lbl in SKIP_LABELS:
                        right_skip.append(j)
                        continue
                    if nxt_lbl in _stop:
                        break
                    nxt_prov = getattr(nxt_item, "prov", None) or []
                    if not nxt_prov:
                        continue
                    if int(getattr(nxt_prov[0], "page_no", -1)) != page_no:
                        break
                    nxt_bbox = getattr(nxt_prov[0], "bbox", None)
                    nxt_x0   = bbox_x0(nxt_bbox) if nxt_bbox else 0.0
                    if not (coplanar(bbox, nxt_bbox, tolerance=250.0)
                            and nxt_x0 >= pic_right * 0.80):
                        break
                    nxt_text = postprocess(
                        (getattr(nxt_item, "text", None) or "").strip()
                    )
                    if _is_letterhead_stop(nxt_text):
                        break
                    right_skip.append(j)
                    if not nxt_text:
                        continue
                    nxt_h  = bbox_h(nxt_bbox) if nxt_bbox else 0.0
                    if nxt_h > 2 and median_h > 0 and nxt_h / median_h < 0.85:
                        nxt_pt = max(round(BODY_PT * (nxt_h / median_h) * 2) / 2, 9.0)
                    else:
                        nxt_pt = 9.0
                    _ha = [c for c in nxt_text if c.isalpha()]
                    right_blocks.append({
                        "text":      nxt_text,
                        "font_pt":   nxt_pt,
                        "bold":      bool(_ha) and all(c.isupper() for c in _ha) and len(nxt_text) <= 60,
                        "italic":    nxt_lbl in ("caption", "footnote"),
                        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
                    })

                _page_break(page_no, force=True)   # картинка/логотип на новой стр. — сохраняем разрыв
                if right_blocks:
                    skip_indices.update(right_skip)
                    add_sidebyside(doc, pil_img, target_w_inch, right_blocks, text_w_inch)
                else:
                    buf = BytesIO()
                    pil_img.save(buf, format="PNG")
                    buf.seek(0)
                    para = doc.add_paragraph()
                    para.alignment                     = img_align
                    para.paragraph_format.space_before = Pt(4)
                    para.paragraph_format.space_after  = Pt(4)
                    para.add_run().add_picture(buf, width=Inches(target_w_inch))
                _last_body_para = None   # разрыв след. страницы — после картинки
            except Exception as exc:
                log.debug("Picture skipped: %s", exc)
            continue

        # ── Получаем текст из Docling (word_order меняет только порядок блоков) ──
        raw_text = (getattr(item, "text", None) or "").strip()
        text = postprocess(raw_text)
        # Инлайн-чистка тела/заголовков: непарные скобки/кавычки и шум-токены.
        # Таблицы (add_table_from_*) сюда не попадают — их текст чистится отдельно.
        if lbl in ("text", "paragraph", "list_item", "section_header",
                   "title", "caption"):
            text = clean_body_text(text)
        if not text:
            log.debug("[стр%d] idx=%d lbl=%s: пустой текст — пропуск", page_no, idx, lbl)
            continue
        # Блок без единой буквы/цифры («`», «;», «?» …) — OCR-шум от остатков
        # пятен/штампов: отдельный абзац из знаков препинания смысла не несёт.
        if lbl in ("text", "paragraph", "list_item") \
                and not re.search(r"[A-Za-zА-Яа-яЁё0-9]", text):
            log.info("[стр%d] пропуск шум-блока (нет букв/цифр): %r", page_no, text[:20])
            continue
        # Микро-блок: 1-2 символа в bbox высотой < 4pt («ы» из обрывка пятна).
        # Настоящий текст такого размера в юр-документах не встречается.
        if lbl in ("text", "paragraph") and bbox is not None \
                and bbox_h(bbox) < 4.0 and len(text) <= 2:
            log.info("[стр%d] пропуск микро-блока (h=%.1fpt): %r",
                     page_no, bbox_h(bbox), text)
            continue
        # Короткий блок-мусор (обрывок пятна/штампа: «t», «ч», «theme cow»,
        # «19/2 t aad Ger») — ни кир-слова, ни реквизита, ни email/URL. Только
        # для тела; заголовки/подписи и длинные абзацы не трогаем.
        if lbl in ("text", "paragraph", "list_item") and is_junk_text(text):
            log.info("[стр%d] пропуск мусор-блока: %r", page_no, text[:30])
            continue

        # Логируем значимые исправления OCR (изменение >= 3 символов или 5% текста)
        if raw_text != text:
            n_changed = sum(a != b for a, b in zip(raw_text, text)) + abs(len(raw_text) - len(text))
            if n_changed >= 3 or n_changed / max(len(raw_text), 1) >= 0.05:
                log.info("  ocr_fix [стр%d]: %r → %r", page_no, raw_text[:70], text[:70])

        # space_before, indent, alignment
        space_before = 0.0
        if bbox is not None and page_no in prev_midY and median_h > 0:
            gap   = abs(bbox_mid_y(bbox) - prev_midY[page_no]) \
                    - (item_h / 2 + prev_h.get(page_no, item_h) / 2)
            extra = gap - median_h * 1.2
            if extra > 2:
                space_before = min(extra * 0.5, 8.0)

        if bbox is not None:
            prev_midY[page_no] = bbox_mid_y(bbox)
            prev_h[page_no]    = item_h

        
        """
        Левый отступ блока относительно левого края контента страницы.
        Сырой отступ прищёлкивается к ближайшему уровню документа
        (build_indent_levels): OCR-шум в 1-4pt не разъезжает вертикали,
        блоки одного отступа встают ровно в одну линию. Порог шума 4pt,
        верхний предел 45% ширины (защита от мусора) — как раньше.
        """
        indent_pt = 0.0
        if bbox is not None:
            raw_indent = float(getattr(bbox, "l", 0)) - page_left_min.get(page_no, 0.0)
            snapped = snap_indent(raw_indent, indent_levels)
            if snapped > 0.0:
                indent_pt = round(min(snapped, pw * 0.45), 1)

        alignment = (
            detect_alignment(bbox, pw)
            if bbox is not None and pw > 0
            else WD_ALIGN_PARAGRAPH.JUSTIFY
        )

        # Корректировка ложного выравнивания для body text:
        # – блок в правой колонке multicolumn-страницы — относительно КОЛОНКИ
        # – CENTER для блоков шире 50% страницы - JUSTIFY
        # – RIGHT для блоков, начинающихся в левых 30% страницы - JUSTIFY
        if lbl in ("paragraph", "text") and bbox is not None:
            bw    = float(getattr(bbox, "r", pw)) - float(getattr(bbox, "l", 0))
            raw_x0 = float(getattr(bbox, "l", 0))
            # Модель колонок: блок правой колонки живёт в системе координат
            # СВОЕЙ колонки — CENTER/RIGHT относительно страницы для него ложны.
            # Заполняет колонку и длинный — JUSTIFY; иначе LEFT, позицию задаёт
            # left_indent (край колонки, уже прищёлкнутый к уровню отступов).
            _pinfo = page_infos.get(page_no)
            _col_idx = (_pinfo.column_for_x(raw_x0)
                        if _pinfo is not None and _pinfo.is_multicolumn else -1)
            if _col_idx > 0:
                _col   = _pinfo.columns[_col_idx]
                _col_w = max(_col.x1 - _col.x0, 1.0)
                _col_align = (WD_ALIGN_PARAGRAPH.JUSTIFY
                              if bw / _col_w >= 0.85 and len(text) > 80
                              else WD_ALIGN_PARAGRAPH.LEFT)
                if _col_align != alignment:
                    log.debug("[стр%d] колонка %d: align %s → %s  %r",
                              page_no, _col_idx,
                              _align_names.get(alignment, "?"),
                              _align_names.get(_col_align, "?"), text[:40])
                alignment = _col_align
            elif alignment == WD_ALIGN_PARAGRAPH.CENTER and bw / pw >= 0.50:
                alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT and raw_x0 < pw * 0.30:
                alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # ДЛИННЫЙ "правый" блок — это колоночный текст (адрес, реквизиты),
            # читаемый слева направо, а не реально правовыключенный (даты/подписи
            # коротки). Делаем LEFT — позицию задаст left_indent (отступ колонки).
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT and len(text) > 60:
                alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Длинный абзац У ЛЕВОГО ПОЛЯ (indent≈0) — это тело документа:
            # делаем выключку по ширине (даст и красную строку). НЕ трогаем
            # колоночный текст (адреса/реквизиты с отступом) — у него indent>10.
            elif (alignment == WD_ALIGN_PARAGRAPH.LEFT
                  and len(text) > 80 and indent_pt < 10.0):
                alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Строка, прижатая к ПРАВОМУ краю (r≈правое поле, не на всю ширину,
            # начало заметно правее левого поля), — правовыровненная строка
            # шапки заявления. Колоночный текст сюда не попадает: он либо в
            # ветке колонок выше, либо уже стал LEFT по правилу len>60.
            elif (alignment == WD_ALIGN_PARAGRAPH.LEFT
                  and bw / pw <= 0.75
                  and float(getattr(bbox, "r", 0)) >= pw * 0.92
                  and raw_x0 >= pw * 0.25):
                alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Короткая строка в ПРАВОЙ КОЛОНКЕ, которую detect_alignment принял
            # за CENTER/RIGHT — на деле левовыключенная строка колонки-шапки заявления.
            # Возвращаем LEFT, чтобы применился отступ колонки и строка встала в
            # один край с длинными соседями того же блока. НЕ трогаем строки,
            # реально прижатые к правому полю (r>=0.92pw) — они правовыровнены.
            _bbox_r = float(getattr(bbox, "r", 0))
            if (alignment in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT)
                    and indent_pt > pw * 0.18 and _bbox_r < pw * 0.92):
                alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Реквизитные строки (Р/с, ИНН, КПП, Кор/сч, БИК, Получатель, ОГРН) —
            # LEFT (маркированные данные), кроме прижатых к правому полю строк
            # правовыровненной шапки. Применяем последним как override.
            if _REQ_LABEL_RE.match(text) and not (
                    _bbox_r >= pw * 0.92 and bw / pw <= 0.75):
                alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Жирный/курсив: Docling formatting (нативные PDF) - приоритет; для сканов
        # — только СТРУКТУРНЫЕ заголовки и Госпошлина-merge. НЕ жирним ALL-CAPS
        #  имена/реквизиты. Метки полей жирнятся в label:content-таблице.
        _fmt      = getattr(item, "formatting", None)
        _fmt_bold = bool(getattr(_fmt, "bold", False))   if _fmt is not None else False
        _fmt_ital = bool(getattr(_fmt, "italic", False)) if _fmt is not None else False
        # Жирность по изображению (сканы): stroke_density блока заметно выше медианы
        # тела страницы. Защиты: достаточная длина и НЕ полностью ALL-CAPS (caps
        # искажает насыщенность - ложная жирность; такие строки жирнятся по метке).
        _ink_bold = False
        _ink_run  = None
        if (_use_ink and lbl in ("text", "paragraph", "list_item")
                and len(text) >= _INK_MIN_LEN):
            _ink_run = ink_cache.get(id(item))
            _ink_med = page_ink_median.get(page_no, 0.0)
            _alpha = [c for c in text if c.isalpha()]
            _all_caps = bool(_alpha) and all(c.isupper() for c in _alpha)
            if (_ink_run is not None and _ink_med > 0 and not _all_caps
                    and _ink_run > _ink_med * _INK_BOLD_K):
                _ink_bold = True
        bold   = _fmt_bold or lbl in ("title", "section_header") or _gosposhlina_bold or _ink_bold
        italic = _fmt_ital or lbl in ("caption", "footnote")

        # Предупреждение при нетипичном размере шрифта
        if font_pt < 8.0:
            log.warning("[стр%d] lbl=%s: font=%.1fpt СЛИШКОМ МАЛО — возможна ошибка bbox",
                        page_no, lbl, font_pt)
        elif font_pt > 20.0 and lbl not in ("title", "section_header"):
            log.warning("[стр%d] lbl=%s: font=%.1fpt СЛИШКОМ ВЕЛИК для тела — возможна ошибка",
                        page_no, lbl, font_pt)

        _bold_src = ("fmt" if _fmt_bold else "ink" if _ink_bold
                     else "lbl" if lbl in ("title", "section_header") else "-")
        log.info("[стр%d] lbl=%-16s align=%-8s bold=%s(%s) font=%.1f run=%s  %r",
                 page_no, lbl,
                 _align_names.get(alignment, str(alignment)),
                 "Y" if bold else "N", _bold_src,
                 font_pt,
                 ("%.2f" % _ink_run) if _ink_run is not None else "-",
                 text[:80])

        _page_break(page_no)

        # Госпошлина: два соседних bbox на одной строке - объединяем
        _gosposhlina_bold = False
        if lbl in ("text", "paragraph") and _GOSPOSHLINA_RE.match(text):
            for _gj in range(idx + 1, min(idx + 3, len(all_items))):
                if _gj in skip_indices:
                    continue
                _gj_item, _ = all_items[_gj]
                _gj_prov = getattr(_gj_item, "prov", None) or []
                if not _gj_prov or int(getattr(_gj_prov[0], "page_no", -1)) != page_no:
                    break
                _gj_bbox = getattr(_gj_prov[0], "bbox", None)
                if _gj_bbox and bbox and coplanar(bbox, _gj_bbox, tolerance=20.0):
                    _gj_text = postprocess(
                        (getattr(_gj_item, "text", None) or "").strip()
                    )
                    if _gj_text:
                        text = f"{text}: {_gj_text}"
                        skip_indices.add(_gj)
                        _gosposhlina_bold = True
                        alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        log.info("[стр%d] Госпошлина merge: %r", page_no, text)
                    break

        """
        Coplanar-пары: "Метка:" + значение
        Паттерн: текущий блок заканчивается на "":" и находится в левой части
        страницы, а следующий непропущенный блок — coplanar и правее.
        Пример: "Место рождения:" (x0=172) + "ГОР. РОСТОВ-НА-ДОНУ" (x0=265)
        """
        _coplanar_rendered = False
        # Значение могло стоять РАНЬШЕ метки в порядке чтения — взято пред-проходом.
        _cp_value_idx = _coplanar_back.get(idx)
        if _cp_value_idx is None and (lbl in ("text", "paragraph") and bbox is not None
                and text.rstrip().endswith(":")
                and _labelish(text)
                and float(getattr(bbox, "l", 0)) < pw * 0.75
                and float(getattr(bbox, "r", pw)) < pw * 0.85):
            _label_right = float(getattr(bbox, "r", 0))
            # Верх метки: в pdf_native t>b, поэтому берём max(t,b)
            _label_top = max(float(getattr(bbox, "t", 0)), float(getattr(bbox, "b", 0)))
            # Ищем до 10 следующих элементов: партнёр должен быть правее метки
            # и начинаться на той же вертикальной позиции. Используем top а не mid_y:
            # для многострочных значений mid_y смещается вниз, top совпадает с меткой.
            for _cj in range(idx + 1, min(idx + 10, len(all_items))):
                if _cj in skip_indices:
                    continue
                _cj_item, _ = all_items[_cj]
                _cj_prov = getattr(_cj_item, "prov", None) or []
                if not _cj_prov or int(getattr(_cj_prov[0], "page_no", -1)) != page_no:
                    break
                _cj_bbox = getattr(_cj_prov[0], "bbox", None)
                if _cj_bbox is None:
                    continue
                _cj_x0 = float(getattr(_cj_bbox, "l", 0))
                _cj_top = max(float(getattr(_cj_bbox, "t", 0)), float(getattr(_cj_bbox, "b", 0)))
                # Значение должно быть правее метки и на той же строке (top ±8pt).
                # 14pt было слишком широко: Должник:(253.3) захватывал ГОР.РОСТОВ(239.6) diff=13.7
                if _cj_x0 > _label_right and abs(_cj_top - _label_top) <= 8.0:
                    _cp_value_idx = _cj
                    break
                # Если встретили метку в ПРАВОЙ части страницы — это значение другого поля,
                # дальше не ищем (значение нашего поля было бы левее этого блока).
                # Метки в ЛЕВОЙ части (x0 < label_right) просто пропускаем.
                _cj_lbl = _label_str(_cj_item)
                if (_cj_lbl in ("text", "paragraph")
                        and (getattr(_cj_item, "text", None) or "").rstrip().endswith(":")
                        and _cj_x0 >= _label_right):
                    break
        if _cp_value_idx is not None:
                _cj_item, _ = all_items[_cp_value_idx]
                _cj_text = postprocess((getattr(_cj_item, "text", None) or "").strip())
                if _cj_text:
                    _full_tw = (pw - 2 * MARGIN_INCH * 72) / 72
                    _cj_bbox = getattr(
                        (getattr(_cj_item, "prov", None) or [None])[0], "bbox", None)
                    _cp_x0 = (float(getattr(_cj_bbox, "l", pw * 0.45))
                              if _cj_bbox is not None else pw * 0.45)
                    _cp_r  = (float(getattr(_cj_bbox, "r", pw))
                              if _cj_bbox is not None else pw)
                    _label_r = float(getattr(bbox, "r", 0)) if bbox is not None else 0.0
                    _gap = _cp_x0 - _label_r
                    if _gap <= 8.0:
                        # Зазор в одну пробельную ширину — Tesseract разрезал ОДНУ
                        # строку («Адрес: 344002, …») на два блока. Склеиваем в один
                        # абзац; выравнивание — по объединённому bbox: строка,
                        # прижатая к правому краю, остаётся правовыровненной.
                        _u_l = float(getattr(bbox, "l", 0)) if bbox is not None else 0.0
                        _u_ratio = (_cp_r - _u_l) / pw if pw > 0 else 1.0
                        if _cp_r >= pw * 0.90 and _u_ratio <= 0.80:
                            _cp_align = WD_ALIGN_PARAGRAPH.RIGHT
                        elif _u_ratio > 0.75:
                            _cp_align = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            _cp_align = WD_ALIGN_PARAGRAPH.LEFT
                        para = doc.add_paragraph()
                        para.alignment                      = _cp_align
                        para.paragraph_format.space_before  = Pt(space_before)
                        para.paragraph_format.space_after   = Pt(0)
                        para.paragraph_format.widow_control = False
                        run = para.add_run(f"{text} {_cj_text}")
                        run.font.name = FONT_NAME
                        run.font.size = Pt(BODY_PT)
                        log.info("[стр%d] coplanar-строка (gap=%.1fpt, %s): %r + %r",
                                 page_no, _gap,
                                 _align_names.get(_cp_align, "?"),
                                 text[:30], _cj_text[:40])
                    else:
                        _cp_col  = max((_cp_x0 - MARGIN_INCH * 72) / 72, 0.5)
                        _cp_ratio = min(_cp_col / _full_tw, 0.65)
                        add_label_content_table(
                            doc,
                            text.rstrip(":").rstrip() + ":",
                            [{"text": _cj_text, "font_pt": BODY_PT, "bold": False, "italic": False}],
                            _full_tw,
                            space_before,
                            indent_inch=0.0,
                            col_ratio=_cp_ratio,
                        )
                        log.info("[стр%d] coplanar-pair: %r → %r",
                                 page_no, text[:40], _cj_text[:40])
                    skip_indices.add(_cp_value_idx)
                    _coplanar_rendered = True
        if _coplanar_rendered:
            continue

        # Блоки МЕТКА:содержимое
        _item_x0 = float(getattr(bbox, "l", 0)) if bbox is not None else 0.0
        # Якорные метки сторон (Кредитор:/Должник:/Финансовый управляющий:...) собираем
        # в ЛЮБОЙ части страницы: в ряде документов RapidOCR кладёт шапку сторон в
        # правую колонку (x0 > порога), и обычное ограничение «левые 45%» их отсекало.
        _is_side_label = SIDE_LABEL_RE.match(text) is not None
        lc = (split_label_content(text)
              if (_item_x0 <= pw * LABEL_MARGIN_THRESHOLD or _is_side_label)
              else None)
        # Одиночная правовыровненная строка с инлайн-значением («Адрес: 630055,
        # г. Новосибирск…» одним блоком в правой зоне) — это просто строка
        # оригинала, а не колонка «метка: содержимое». Рендерим обычным абзацем.
        if (lc is not None and lc[1] and bbox is not None
                and float(getattr(bbox, "r", 0)) >= pw * 0.85
                and _item_x0 >= pw * 0.30):
            log.info("[стр%d] label-строка правой зоны → обычный абзац: %r",
                     page_no, text[:50])
            lc = None
        if lc is not None:
            _full_tw      = (pw - 2 * MARGIN_INCH * 72) / 72
            _lc_indent    = 0.0
            _lc_col_ratio = None
            # Определяем ширину колонки метки динамически по x0 первого контент-блока.
            # Это выравнивает КРЕДИТОР:/ДОЛЖНИК:/АРБИТРАЖНЫЙ УПРАВЛЯЮЩИЙ в одну колонку:
            # все три имеют контент около x=298pt — col_ratio вычисляется одинаково.
            _cont_x0 = None
            for _pk in range(idx + 1, min(idx + 4, len(all_items))):
                if _pk in skip_indices:
                    continue
                _pk_item, _ = all_items[_pk]
                _pk_lbl = _label_str(_pk_item)
                if _pk_lbl in (SKIP_LABELS | {"table", "picture", "figure", "image",
                                               "list_item", "section_header", "title"}):
                    break
                _pk_prov = getattr(_pk_item, "prov", None) or []
                if not _pk_prov or int(getattr(_pk_prov[0], "page_no", -1)) != page_no:
                    break
                _pk_bbox = getattr(_pk_prov[0], "bbox", None)
                if _pk_bbox is not None:
                    _cont_x0 = float(getattr(_pk_bbox, "l", 0))
                break

            _lm_pt = MARGIN_INCH * 72
            if _cont_x0 is not None and _cont_x0 > pw * 0.42:
                # Таблица начинается от левого поля (indent=0).
                # Ширина колонки метки = расстояние от левого поля до x0 контента.
                # контент выравнивается по одной вертикали, метки начинаются с левого поля.
                _lc_indent    = 0.0
                _label_w_in   = max((_cont_x0 - _lm_pt) / 72, 0.5)
                _lc_col_ratio = min(_label_w_in / _full_tw, 0.65)

            text_w_inch = _full_tw - _lc_indent

            log.info("  → label:content  label=%r  content=%r  x0=%.1fpt (thr=%.1fpt) "
                     "indent=%.2fin col_ratio=%s",
                     lc[0], lc[1][:50], _item_x0, pw * LABEL_MARGIN_THRESHOLD,
                     _lc_indent, f"{_lc_col_ratio:.2f}" if _lc_col_ratio else "default")

            label_txt, first_content = lc
            content_items: list[dict] = []
            if first_content:
                content_items.append({"text": first_content, "font_pt": BODY_PT, "bold": True, "italic": False})
            last_y = bbox_mid_y(bbox) if bbox is not None else 0.0
            for j in range(idx + 1, min(idx + 8, len(all_items))):
                if j in skip_indices:
                    continue
                nxt_item, _ = all_items[j]
                nxt_lbl     = _label_str(nxt_item)
                # section_header и title — самостоятельные блоки, не сливаем в label:content
                if nxt_lbl in (SKIP_LABELS | {"table", "picture", "figure", "image",
                                               "list_item", "section_header", "title"}):
                    log.debug("  label:content остановлен на lbl=%s %r", nxt_lbl,
                              (getattr(nxt_item, "text", "") or "")[:40])
                    break
                nxt_prov = getattr(nxt_item, "prov", None) or []
                if not nxt_prov or int(getattr(nxt_prov[0], "page_no", -1)) != page_no:
                    break
                nxt_bbox = getattr(nxt_prov[0], "bbox", None)
                if nxt_bbox and median_h > 0:
                    nxt_y = bbox_mid_y(nxt_bbox)
                    if abs(nxt_y - last_y) > median_h * 4:
                        break
                    last_y = nxt_y
                nxt_text = postprocess((getattr(nxt_item, "text", None) or "").strip())
                if not nxt_text:
                    skip_indices.add(j)
                    continue
                if split_label_content(nxt_text) is not None:
                    break
                # Стоп-слова: самостоятельные секции документа (не продолжение блока)
                if re.match(r'^(Госпошлин|ПРОСИТ\s+СУД|Приложени)', nxt_text, re.IGNORECASE):
                    log.debug("  label:content стоп (standalone) %r", nxt_text[:30])
                    break
                # Содержимое label:content всегда рендерится BODY_PT — не масштабируем
                # по высоте bbox (высота блоков персданных меньше медианы,
                # что даёт ложный результат 9pt для Потапова/адреса).
                nxt_pt = BODY_PT
                content_items.append({
                    "text": nxt_text, "font_pt": nxt_pt,
                    "bold":   (len(content_items) == 0),
                    "italic": False,
                })
                log.info("    content[%d] font=%.1fpt bold=%s %r",
                         len(content_items) - 1, nxt_pt,
                         "Y" if len(content_items) == 1 else "N", nxt_text[:60])
                skip_indices.add(j)
            if bbox is not None:
                prev_midY[page_no] = bbox_mid_y(bbox)
                prev_h[page_no]    = item_h
            add_label_content_table(doc, label_txt, content_items, text_w_inch, space_before,
                                    indent_inch=_lc_indent, col_ratio=_lc_col_ratio)
            continue

        # Заголовки
        if lbl in LABEL_HEADING:
            # Подзаголовки со строчной буквы центрируем ТОЛЬКО если блок
            # расположен в левой половине страницы.
            alpha = [c for c in text if c.isalpha()]
            _h_x0 = float(getattr(bbox, "l", 0)) if bbox is not None else 0.0
            if alpha and alpha[0].islower() and _h_x0 < pw * 0.48:
                alignment = WD_ALIGN_PARAGRAPH.CENTER
            para = doc.add_heading("", level=min(LABEL_HEADING[lbl], 9))
            para.alignment                       = alignment
            para.paragraph_format.space_before   = Pt(max(space_before, 4.0))
            para.paragraph_format.space_after    = Pt(2.0)
            para.paragraph_format.widow_control  = False
            para.paragraph_format.keep_with_next = False
            para.paragraph_format.keep_together  = False
            _h_lm = page_left_min.get(page_no, 0.0)
            # Правовыровненный заголовок шапки заявления и ЕЁ ПРОДОЛЖЕНИЯ на
            # след. странице (Кредиторы:/«N. ООО …»/Уполномоченный орган:) —
            # выключка ВПРАВО к полю (align=RIGHT, дотягивается до края) в правой
            # части листа. Флашим без left_indent: иначе большой отступ (из x0 в
            # правой половине) зажимает короткий заголовок в узкую полосу →
            # перенос по слову/слогу. Признак — сам блок (RIGHT+x0), не флаг
            # страницы: продолжение списка на стр.2 детект колонок не помечает rj.
            if alignment == WD_ALIGN_PARAGRAPH.RIGHT and _h_x0 > pw * 0.42:
                para.paragraph_format.left_indent = Pt(0)
            # Метки-заголовки ЛЕВОвыровненной правой КОЛОНКИ (ФНС/ПСБ: Кредитор:/
            # Должник:/Финансовый управляющий:) — кончаются на ":", НЕ дотянуты до
            # поля (align != RIGHT). Выравниваем по левому краю колонки с отступом
            # соседей блока. alignment != CENTER: центрированные — по центру.
            elif (text.rstrip().endswith(":") and _h_x0 > pw * 0.42
                    and alignment != WD_ALIGN_PARAGRAPH.CENTER):
                alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.alignment = alignment
                # Отступ прищёлкиваем к уровню документа — метка встаёт в одну
                # вертикаль с текстовыми блоками той же колонки.
                para.paragraph_format.left_indent = Pt(
                    snap_indent(max(_h_x0 - _h_lm, 0.0), indent_levels))
            else:
                para.paragraph_format.left_indent = Pt(0)
            # Заголовок "тип документа + подзаголовок/адрес" : Docling склеивает в один
            # section_header, но в оригинале ЗАГЛАВНЫЙ тип/название стоит ОТДЕЛЬНОЙ
            # строкой, а подзаголовок/адрес — следующей, обычным шрифтом (не жирным,
            # мельче). Выносим тип на свою строку; вторую часть — телесным стилем.
            # За ALL-CAPS названием следует строчное слово (подзаголовок) ИЛИ цифра
            # почтового индекса (адрес).
            _title_m = (re.match(r'^([А-ЯЁ][А-ЯЁ]+(?:\s+[А-ЯЁ]+)*)\s+([а-яё\d].*)$',
                                 text, re.DOTALL)
                        if lbl == "section_header"
                           and alignment == WD_ALIGN_PARAGRAPH.CENTER
                        else None)

            def _style_run(_r, bold=True, size=None):
                _r.font.name      = FONT_NAME
                _r.font.size      = Pt(size if size is not None else font_pt)
                _r.font.bold      = bold
                _r.font.color.rgb = RGBColor(0, 0, 0)

            if _title_m:
                run = para.add_run(_title_m.group(1))
                _style_run(run)
                run.add_break()
                # Подзаголовок/адрес — обычным телесным шрифтом (как в оригинале).
                run2 = para.add_run(_title_m.group(2))
                _style_run(run2, bold=False, size=BODY_PT)
            else:
                run = para.add_run(text)
                _style_run(run)
            continue

        # Списки
        # Нумерованные пункты («N текст») 
        # Подпункты (первая буква строчная)
        # Остальные list_item     
        if lbl == "list_item":
            if alignment in (WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER):
                alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _li_alpha  = [c for c in text if c.isalpha()]
            num_match  = _NUM_ITEM_RE.match(text)
            # Подпункт: явный дефис/буллет в начале (любой регистр далее) ИЛИ
            # строчная первая буква (продолжение-перечисление).
            # \s* (не \s+): на сканах дефис приклеен к слову («-задолженность»)
            # — иначе он не срезался и к нему добавлялся второй: «- -задолженность».
            _dash_sub  = bool(re.match(r'^\s*[-–—•·]\s*[^\W\d]', text))
            is_sub_item = (not num_match) and (
                _dash_sub or (bool(_li_alpha) and _li_alpha[0].islower()))

            if num_match:
                num_str   = num_match.group(1).strip() + "."
                text_body = text[num_match.end():]
                para = doc.add_paragraph()
                para.alignment                       = alignment
                para.paragraph_format.space_before   = Pt(space_before)
                para.paragraph_format.space_after    = Pt(0)
                para.paragraph_format.widow_control  = False
                # Красная строка: «2. текст» на первой строке с отступом,
                # продолжение — от левого края (по просьбе пользователя).
                para.paragraph_format.left_indent        = Pt(0)
                para.paragraph_format.first_line_indent  = Pt(35.4)
                run           = para.add_run(f"{num_str} {text_body}")
                run.font.name = FONT_NAME
                run.font.size = Pt(font_pt)
                log.info("[стр%d] list_item → numbered  font=%.1fpt align=%-8s %r",
                         page_no, font_pt,
                         _align_names.get(alignment, str(alignment)), text[:60])
            elif is_sub_item:
                # Подпункт: та же красная строка что у нумерованных (35.4pt),
                # продолжение — от левого края. Дефис и текст на одной строке.
                para = doc.add_paragraph()
                para.alignment                      = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before  = Pt(space_before)
                para.paragraph_format.space_after   = Pt(0)
                para.paragraph_format.widow_control = False
                para.paragraph_format.left_indent        = Pt(0)
                para.paragraph_format.first_line_indent  = Pt(35.4)
                # Убираем уже имеющийся дефис/буллет, чтобы не задвоить.
                # Пробел после дефиса опционален («-задолженность» на сканах),
                # но дефис перед цифрой не трогаем (минус числа — не буллет).
                _sub_body = re.sub(r'^\s*[-–—•·]\s*(?=[^\W\d])', '', text)
                run           = para.add_run("- " + _sub_body)
                run.font.name = FONT_NAME
                run.font.size = Pt(font_pt)
                log.info("[стр%d] list_item → dash   font=%.1fpt %r",
                         page_no, font_pt, text[:60])
            else:
                para = doc.add_paragraph()
                para.alignment                      = alignment
                para.paragraph_format.space_before  = Pt(space_before)
                para.paragraph_format.space_after   = Pt(0)
                para.paragraph_format.widow_control = False
                para.paragraph_format.first_line_indent = Pt(35.4)
                para.paragraph_format.left_indent       = Pt(0)
                run           = para.add_run(text)
                run.font.name = FONT_NAME
                run.font.size = Pt(font_pt)
                log.info("[стр%d] list_item → красная строка font=%.1fpt align=%-8s %r",
                         page_no, font_pt,
                         _align_names.get(alignment, str(alignment)), text[:60])
            para.paragraph_format.line_spacing = page_line_spacing.get(
                page_no, LINE_SPACING)                          # адаптивный интервал тела
            # Обновляем трекинг для последующей проверки continuation
            _last_body_para = para
            _last_body_text = text
            _seen_text_pages.add(page_no)
            continue

        # Строка подписи
        # page_footer — OCR часто метит строку подписи как колонтитул
        if lbl in ("text", "paragraph", "page_footer") and _REPR_RE.search(text):
            _right_txt = ""
            _sig_pic   = None
            # Просматриваем ближайшие 10 элементов: ищем картинку подписи
            # И все вхождения инициалов (могут дублироваться OCR).
            for _j in range(idx + 1, min(idx + 10, len(all_items))):
                if _j in skip_indices:
                    continue
                _nj_item, _ = all_items[_j]
                _nj_lbl  = _label_str(_nj_item)
                _nj_prov = getattr(_nj_item, "prov", None) or []
                _nj_pg   = int(getattr(_nj_prov[0], "page_no", -1)) if _nj_prov else -1
                if _nj_pg != page_no:
                    break
                if _nj_lbl in ("picture", "figure", "image"):
                    if _sig_pic is None:
                        _sig_pic = _get_item_image(_nj_item, dl_doc)
                    skip_indices.add(_j)
                    continue
                _nj_text = postprocess(
                    (getattr(_nj_item, "text", None) or "").strip()
                )
                if _INITIALS_RE.match(_nj_text):
                    if not _right_txt:
                        _right_txt = _nj_text
                    skip_indices.add(_j)   # скипаем ВСЕ дубли инициалов
                    log.info("[стр%d] signature initials skip %d: %r", page_no, _j, _nj_text)
            # Рендерим только если нашли картинку ИЛИ инициалы
            if _right_txt or _sig_pic is not None:
                _sig_tw = (pw - 2 * MARGIN_INCH * 72) / 72
                add_signature_row(doc, text, _right_txt, _sig_tw,
                                  space_before=max(space_before, 6.0),
                                  sig_image=_sig_pic)
                _last_body_para = None
                _last_body_text = text
                _seen_text_pages.add(page_no)
                log.info("[стр%d] signature row: pic=%s right=%r",
                         page_no, _sig_pic is not None, _right_txt)
                continue

        # Обычные параграфы

        # Детектирование продолжения (orphan-блок): блок начинается с союза/
        # частицы строчными буквами, а предыдущий параграф не завершён.
        # Две ситуации:
        #   1) Межстраничное: первый text-блок страницы с маленькой буквы
        #   2) Внутри страницы: блок начинается с явного союза
        _text_alpha = [c for c in text if c.isalpha()]
        _is_lowercase_start = bool(_text_alpha) and _text_alpha[0].islower()
        _prev_unfinished = (
            _last_body_para is not None and
            not _SENT_END_RE.search(_last_body_text)
        )

        _is_crosspage = (
            _is_lowercase_start and
            page_no not in _seen_text_pages and   # первый text-блок этой страницы
            _prev_unfinished
        )
        _is_conjunction = (
            _is_lowercase_start and
            bool(_CONT_RE.match(text)) and
            _prev_unfinished
        )
        # JUSTIFY-блок с маленькой буквы после блока, заканчивающегося
        # на двоеточие ":" — явный признак начатого перечисления/продолжения.
        # Ограничения делают правило безопасным:
        #   предыдущий блок ДОЛЖЕН заканчиваться на ':'
        #   текущий блок НЕ является датой DD.MM.YYYY (дата всегда начинает новый абзац)
        #   текущий блок НЕ начинается с известного открывающего слова нового абзаца
        _is_justify_cont = (
            _is_lowercase_start and
            not bool(_CONT_RE.match(text)) and
            not bool(_DATE_START_RE.match(text)) and
            not bool(_PARA_STARTERS_RE.match(text)) and
            lbl in ("text", "paragraph") and
            alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and
            _prev_unfinished and
            _last_body_text.rstrip().endswith(':')
        )

        # блок помечен в Fix 2 (date-start swap) как прямое
        # продолжение вставленного датой блока — сливаем независимо от регистра.
        _is_marked_cont = (
            id(item) in _continuation_ids and
            lbl in ("text", "paragraph") and
            _prev_unfinished
        )

        if lbl in ("text", "paragraph") and (_is_crosspage or _is_conjunction
                                              or _is_justify_cont or _is_marked_cont):
            # Присоединяем к предыдущему параграфу вместо создания нового
            kind = ("cross-page"   if _is_crosspage   else
                    "conjunction"  if _is_conjunction  else
                    "marked-cont"  if _is_marked_cont  else "justify-cont")
            log.info("[стр%d] continuation (%s): %r → merged into previous",
                     page_no, kind, text[:60])
            assert _last_body_para is not None  # гарантировано _prev_unfinished
            run           = _last_body_para.add_run(" " + text)
            run.font.name = FONT_NAME
            run.font.size = Pt(font_pt)
            _last_body_text += " " + text
            if page_no not in _seen_text_pages:
                _seen_text_pages.add(page_no)
        else:
            para = doc.add_paragraph()
            para.alignment                      = alignment
            para.paragraph_format.space_before  = Pt(space_before)
            para.paragraph_format.space_after   = Pt(0)
            para.paragraph_format.widow_control = False
            # Межстрочный интервал тела (адаптивный по странице: плотные — меньше,
            # чтобы текст не вылезал). Шапку (LEFT-колонка) оставляем плотной.
            if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                para.paragraph_format.line_spacing = page_line_spacing.get(
                    page_no, LINE_SPACING)

            # Красная строка: ВСЕ полноширинные (JUSTIFY) абзацы тела получают
            # first-line indent — как в исходном юр-документе (у каждого абзаца
            # красная строка). left_indent тут не нужен (текст на всю ширину).
            _is_red_line = (
                lbl in ("paragraph", "text")
                and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
            )
            if _is_red_line:
                para.paragraph_format.first_line_indent = Pt(35.4)
                para.paragraph_format.left_indent       = Pt(0)
            elif _gosposhlina_bold and bbox is not None and alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                # Госпошлина: left_indent совпадает с left_indent heading-блоков той же страницы
                _gp_lm = page_left_min.get(page_no, 0.0)
                _gp_x0 = float(getattr(bbox, "l", 0))
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.left_indent       = Pt(max(_gp_x0 - _gp_lm, 0.0))
            else:
                para.paragraph_format.first_line_indent = Pt(0)
                # left_indent применяем ТОЛЬКО к LEFT/JUSTIFY: для RIGHT/CENTER
                # позицию задаёт выравнивание, добавочный отступ уводит текст вбок.
                if alignment in (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY):
                    para.paragraph_format.left_indent = Pt(indent_pt)
                else:
                    para.paragraph_format.left_indent = Pt(0)

            # Судебная шапка определения. Разбиваем на строки
            if (lbl in ("text", "paragraph")
                    and re.search(r'Дело\s*[№N]', text)
                    and re.search(r'\bСудья\b', text)):
                text = re.sub(r'\s+(Дело\s*[№N])', r'\n\1', text)
                text = re.sub(r'\s+(Судья\b)', r'\n\1', text)
            # Поддержка переносов строк внутри блока (\n): Docling иногда склеивает
            # визуальные строки шапки в один блок — разбиваем их обратно на строки.
            for _k, _part in enumerate(text.split("\n")):
                if _k > 0:
                    para.add_run().add_break()
                run           = para.add_run(_part)
                run.font.name = FONT_NAME
                run.font.size = Pt(font_pt)
                run.bold      = bold
                run.italic    = italic

            if lbl in ("text", "paragraph"):
                _last_body_para  = para
                _last_body_text  = text
                _seen_text_pages.add(page_no)

    # Итоговая статистика
    total   = len(all_items)
    skipped = len(skip_indices)
    labels_seen: dict[str, int] = {}
    for item, _ in all_items:
        lbl = _label_str(item)
        if lbl:
            labels_seen[lbl] = labels_seen.get(lbl, 0) + 1
    log.info("build_docx итог: обработано %d из %d элементов (пропущено %d)",
             total - skipped, total, skipped)
    for lbl, cnt in sorted(labels_seen.items()):
        log.info("  %-20s %d", lbl, cnt)

    return doc


#  Публичные функции

def _is_native_pdf(pdf_path: Path) -> bool:
    """Возвращает True если PDF содержит текстовый слой (не чистый скан).

    Проверяет первые 3 страницы: если хотя бы на одной есть текстовые блоки
    с реальным содержимым — PDF нативный и word_order не нужен.
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(pdf_path))
        for i in range(min(3, doc.page_count)):
            blocks = doc[i].get_text("blocks")
            for b in blocks:
                if b[6] == 0:  # текстовый блок
                    text = b[4].strip()
                    if len(text) > 10:  # есть реальный текст
                        doc.close()
                        return True
        doc.close()
    except Exception:
        pass
    return False


def convert_pdf(
    pdf_path: Path,
    docx_path: Path,
    converter,
    ocr_reader=None,
    use_word_order: bool = True,
    highlight: bool = True,
    llm: bool = False,
    llm_model: str = "",
    ink_bold: bool = False,
) -> bool:
    log.info("  Конвертация: %s", pdf_path.name)
    try:
        # Нативные PDF не нуждаются в EasyOCR для word_order —
        # Docling сам правильно читает порядок блоков из текстового слоя.
        # Это ускоряет конвертацию с ~4 мин до ~10 сек на файл.
        effective_word_order = use_word_order
        if use_word_order and ocr_reader is not None:
            if _is_native_pdf(pdf_path):
                effective_word_order = False
                log.info("  native PDF — word_order отключён (EasyOCR не нужен)")
            else:
                log.info("  scan PDF — word_order включён (Tesseract)")

        result = converter.convert(str(pdf_path))
        dl_doc = result.document

        page_sizes: dict[int, tuple[float, float]] = {}
        for page_no, page in dl_doc.pages.items():
            size = getattr(page, "size", None)
            if size:
                page_sizes[int(page_no)] = (
                    float(getattr(size, "width",  0.0)),
                    float(getattr(size, "height", 0.0)),
                )

        doc = build_docx(dl_doc, page_sizes, ocr_reader=ocr_reader,
                         use_word_order=effective_word_order, ink_bold=ink_bold)
        if highlight:
            # Автоочистка детерминируемых OCR-ошибок в помеченных словах +
            # подсветка остатка жёлтым (для быстрой ручной вычитки).
            from .highlight import highlight_suspicious
            highlight_suspicious(doc)
            if llm:
                # локальная LLM (llama-cpp-python + .gguf, в процессе)
                # добивает оставшиеся подсвеченные слова — только их, не весь
                # документ. No-op, если пакета/модели нет.
                from .llm_correct import correct_highlighted
                correct_highlighted(doc, llm_model or None)
        doc.save(str(docx_path))
        log.info("  ✓ %s", docx_path.name)
        return True
    except Exception as exc:
        log.error("  ✗ %s: %s", pdf_path.name, exc, exc_info=True)
        return False


class DoclingBatchConverter:
    def __init__(
        self,
        input_folder: str,
        output_folder: str,
        backup_folder: str | None = None,
        langs: list[str] | None = None,
        use_word_order: bool = True,
        highlight: bool = True,
        llm: bool = False,
        llm_model: str = "",
        ink_bold: bool = False,
        ocr_preprocess: bool = True,
    ) -> None:
        self.input_folder   = Path(input_folder)
        self.output_folder  = Path(output_folder)
        self.backup_folder  = Path(backup_folder) if backup_folder else None
        self.langs          = langs or ["ru", "en"]
        self.use_word_order = use_word_order
        self.highlight      = highlight
        self.llm            = llm
        self.llm_model      = llm_model
        self.ink_bold       = ink_bold
        self.ocr_preprocess = ocr_preprocess
        self.output_folder.mkdir(parents=True, exist_ok=True)
        if self.backup_folder:
            self.backup_folder.mkdir(parents=True, exist_ok=True)
        self._converter  = None
        self._ocr_reader = None

    @property
    def converter(self):
        if self._converter is None:
            log.info("Инициализация Docling pipeline (DocLayNet + TableFormer + Tesseract)...")
            self._converter = build_converter(self.langs, ocr_preprocess=self.ocr_preprocess)
            log.info("Pipeline готов.")
        return self._converter

    @property
    def ocr_reader(self):
        if self._ocr_reader is None and self.use_word_order:
            try:
                from .word_order import TesseractWordReader
                log.info("Инициализация Tesseract reader для word_order...")
                self._ocr_reader = TesseractWordReader(self.langs)
                log.info("Tesseract готов.")
            except Exception as exc:
                log.warning("Tesseract недоступен, word_order отключён: %s", exc)
        return self._ocr_reader

    def process(self, move_to_backup: bool = True) -> None:
        pdf_files = sorted(self.input_folder.glob("*.pdf"))
        total     = len(pdf_files)
        log.info("Найдено %d PDF-файлов.", total)
        stats = {"ok": 0, "fail": 0, "skip": 0}
        t0    = time.monotonic()

        for i, pdf_path in enumerate(pdf_files, 1):
            docx_path = self.output_folder / f"{pdf_path.stem}.docx"
            if docx_path.exists():
                log.info("[%d/%d] Пропуск: %s", i, total, pdf_path.name)
                stats["skip"] += 1
                continue
            log.info("[%d/%d] %s (%.2f MB)", i, total, pdf_path.name,
                     pdf_path.stat().st_size / 1_048_576)
            ok = convert_pdf(pdf_path, docx_path, self.converter,
                             ocr_reader=self.ocr_reader,
                             use_word_order=self.use_word_order,
                             highlight=self.highlight,
                             llm=self.llm, llm_model=self.llm_model,
                             ink_bold=self.ink_bold)
            if ok:
                stats["ok"] += 1
                if self.backup_folder and move_to_backup:
                    shutil.move(str(pdf_path), str(self.backup_folder / pdf_path.name))
            else:
                stats["fail"] += 1
            if i % 10 == 0:
                gc.collect()

        elapsed = time.monotonic() - t0
        sep = "=" * 55
        log.info("\n%s", sep)
        log.info("Готово за %.1f мин | OK: %d | Ошибок: %d | Пропущено: %d",
                 elapsed / 60, stats["ok"], stats["fail"], stats["skip"])
        log.info(sep)
